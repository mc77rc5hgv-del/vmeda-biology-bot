import asyncio
import copy
import html
import io
import json
import logging
import re
import os
import sys
import time
import urllib.parse
from concurrent.futures import ThreadPoolExecutor
from datetime import date, datetime, timedelta, timezone
from aiogram import Bot, Dispatcher, F
from aiogram.types import (
    Message, CallbackQuery, InlineKeyboardButton, FSInputFile, BufferedInputFile, Update,
    BotCommand, BotCommandScopeDefault, BotCommandScopeChat, LabeledPrice,
    ReplyKeyboardMarkup, KeyboardButton, ReplyKeyboardRemove,
)
from aiogram.filters import CommandStart, Command
from aiogram.utils.keyboard import InlineKeyboardBuilder
from aiogram.enums import ChatMemberStatus
from aiogram.exceptions import TelegramBadRequest
from aiogram.dispatcher.event.bases import SkipHandler
from docx import Document as DocxDocument
from docx.shared import Pt

from ai import prompts as ai_prompts
from ai import rag as ai_rag
from ai import router as ai_router
from ai import service as ai_service
from ai.providers import gemini as ai_gemini
from ai.providers import openai as ai_openai
from ai.providers import xai as ai_xai
from ai import confidence as ai_confidence
from ai import math_verifier as ai_math_verifier
from ai import mcq_verifier as ai_mcq_verifier
from ai import reference_bank as ai_reference_bank
from ai import validator as ai_validator
from ai import vision_parser as ai_vision_parser
from ai.router import AIRefusalError
from ai.service import solve as solve_ai_request
from ai.task import TaskRepresentation
from ai.vision import resize_image as resize_image_for_ai
from repositories import knowledge

# Когда файл запущен напрямую (python3 telegram_bot.py — так его стартует Railway, см.
# railway.json), Python грузит его как модуль "__main__", а не "telegram_bot". Но handlers/*.py
# и services/access.py делают `import telegram_bot as tb` — без этой строки такой импорт не
# находит "telegram_bot" в sys.modules и запускает ВТОРОЕ, вложенное выполнение этого же файла
# с нуля (под именем "telegram_bot"), которое тут же падает с AttributeError на первом же
# собственном `from handlers import ...`/`from services import ...` (модуль ещё не доинициализирован
# в момент обращения к его атрибутам). Алиас ниже говорит импорт-системе, что "telegram_bot" —
# это тот же самый уже выполняющийся модуль, а не что-то, что надо импортировать заново. Должен
# стоять до первого `from services import ...`/`from handlers import ...` ниже по файлу — но не
# обязательно раньше ai/repositories (они telegram_bot не импортируют).
sys.modules.setdefault("telegram_bot", sys.modules[__name__])

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

BOT_TOKEN = os.getenv("BOT_TOKEN")
if not BOT_TOKEN:
    raise RuntimeError(
        "BOT_TOKEN is not set. Set it as an environment variable (e.g. Railway → Variables) — "
        "never hardcode the token in source code."
    )
CHANNEL_ID = "@Vmeda_examen"
ADMIN_IDS = {1326779223, 8601892147}
STATS_DIR = os.getenv("STATS_DIR", ".")
STATS_FILE = os.path.join(STATS_DIR, "stats.json")

# Все студенты и админы бота — в России, но контейнер (Railway) по умолчанию работает в UTC,
# так что "новый день"/"новый месяц" по системному времени наступает на 3 часа позже реального
# московского. Фиксированный offset, а не zoneinfo("Europe/Moscow") — Россия с 2014 года не
# переходит на летнее/зимнее время, МСК = UTC+3 круглый год, и фиксированный offset не зависит
# от наличия системной tzdata на минимальном Docker-образе (zoneinfo мог бы упасть рантайм-
# ошибкой, если tzdata не установлена). Используется для ВСЕХ повторяющихся суточных/месячных
# периодов (дневной лимит AI, реферальный месяц, окна cost circuit breaker и т.п.) — но НЕ трогает
# уже сохранённые sub["expires"] существующих подписок, только то, как будущие грант-моменты и
# константы-дедлайны (services/access.py) вычисляются с этого момента.
APP_TIMEZONE = timezone(timedelta(hours=3), name="MSK")

def local_now() -> datetime:
    """Текущее время в APP_TIMEZONE (МСК) — единая точка отсчёта для всех дневных/месячных
    период-ключей вместо часового пояса контейнера."""
    return datetime.now(APP_TIMEZONE)

def local_today() -> date:
    return local_now().date()
# Ключи провайдеров AI живут в ai/providers/*.py (каждый модуль сам читает свою переменную
# окружения) — эти два имени просто ре-экспортированы, потому что UI-уровень бота (кнопка
# "Отправить фото", текст меню AI) должен знать, доступен ли AI, не заглядывая внутрь пакета ai.
OPENAI_API_KEY = ai_openai.OPENAI_API_KEY  # без него AI-раздел показывает "временно недоступен"

def ai_provider_available() -> bool:
    """True, если доступен хотя бы ОДИН провайдер, способный реально провести пользователя через
    конвейер VMedA AI целиком — не просто "какой-то ключ где-то задан". И vision-парсинг (см.
    ai.vision_parser.parse_task), и quick-ответ (см. ai.router.build_attempts — даже когда
    primary="openai", список попыток становится ["openai", "gemini"], если Gemini настроен) уже
    умеют падать с OpenAI на Gemini автоматически, поэтому UI не должен требовать именно
    OPENAI_API_KEY — бот полностью в состоянии обслужить AI-запрос на одном Gemini, если OpenAI не
    настроен вообще. xAI/Grok сюда сознательно не входит: он никогда не используется ни для
    vision-парсинга, ни как единственный провайдер quick-ответа (только как один из вариантов
    ПОДРОБНОГО разбора для theory_complex) — если настроен только он, конвейер всё равно не
    сможет даже разобрать первое сообщение сессии."""
    return bool(OPENAI_API_KEY) or bool(ai_gemini.GEMINI_API_KEY)

DIVIDER = "━━━━━━━━━━━━━━"
IMAGES_DIR = "images"
ANATOMY_IMAGES_DIR = os.path.join(IMAGES_DIR, "anatomy")
HISTOLOGY_IMAGES_DIR = os.path.join(IMAGES_DIR, "histology")

# ==================== ЗАГРУЗКА ДАННЫХ ====================
# Контент (билеты/вопросы/теория/анатомия/гистология) грузится из JSON в repositories/knowledge.py
# (импортирован выше вместе с остальными модулями) — здесь только реэкспорт тех же имён под теми
# же названиями, чтобы все обращения по всему файлу (QUESTIONS[...], ANATOMY[...] и т.д.) остались
# без изменений.
TICKETS = knowledge.TICKETS
TICKETS_DICT = knowledge.TICKETS_DICT
QUESTIONS = knowledge.QUESTIONS
PHYSICS_QUESTIONS = knowledge.PHYSICS_QUESTIONS
PHYSICS_GRADE45_QUESTIONS = knowledge.PHYSICS_GRADE45_QUESTIONS
PHYSICS_EXTRA_QUESTIONS = knowledge.PHYSICS_EXTRA_QUESTIONS
CHEMISTRY_LABS = knowledge.CHEMISTRY_LABS
CHEMISTRY_THEORY = knowledge.CHEMISTRY_THEORY
CHEMISTRY_TASKS = knowledge.CHEMISTRY_TASKS
CHEMISTRY_THEORY_TICKETS = knowledge.CHEMISTRY_THEORY_TICKETS
CHEMISTRY_PRACTICE_TICKETS = knowledge.CHEMISTRY_PRACTICE_TICKETS
PHYSICS_TASKS = knowledge.PHYSICS_TASKS
PHYSICS_TEST_TICKETS = knowledge.PHYSICS_TEST_TICKETS
PHYSICS_TASK_TICKETS = knowledge.PHYSICS_TASK_TICKETS
PHYSICS_THEORY_TICKETS = knowledge.PHYSICS_THEORY_TICKETS
ANATOMY = knowledge.ANATOMY
ANATOMY_EXAM_TEST_PARTS = knowledge.ANATOMY_EXAM_TEST_PARTS
ANATOMY_EXAM_THEORY_SECTIONS = knowledge.ANATOMY_EXAM_THEORY_SECTIONS
ANATOMY_EXAM_PRACTICE_SECTIONS = knowledge.ANATOMY_EXAM_PRACTICE_SECTIONS
HISTOLOGY = knowledge.HISTOLOGY
OPERATIVE_SURGERY = knowledge.OPERATIVE_SURGERY
PHYSIOLOGY = knowledge.PHYSIOLOGY

ai_rag.configure(
    questions=QUESTIONS, physics_questions=PHYSICS_QUESTIONS, chemistry_theory=CHEMISTRY_THEORY,
    chemistry_theory_tickets=CHEMISTRY_THEORY_TICKETS, chemistry_practice_tickets=CHEMISTRY_PRACTICE_TICKETS,
    anatomy=ANATOMY, operative_surgery=OPERATIVE_SURGERY, physiology=PHYSIOLOGY,
)
ai_reference_bank.configure(ANATOMY_EXAM_TEST_PARTS)

# ==================== СТАТИСТИКА (СОХРАНЯЕТСЯ НА ДИСК) ====================
def load_stats() -> dict:
    os.makedirs(STATS_DIR, exist_ok=True)
    if os.path.exists(STATS_FILE):
        try:
            with open(STATS_FILE, "r", encoding="utf-8") as f:
                data = json.load(f)
            data["total_users"] = set(data.get("total_users", []))
            data.setdefault("start_count", 0)
            data.setdefault("random_ticket_used", 0)
            data.setdefault("random_question_used", 0)
            data.setdefault("question_opened", {})
            data.setdefault("broadcast_count", 0)
            data.setdefault("helperchat_promo_seen", {})
            data.setdefault("referrals", {})
            data.setdefault("referred_by", {})
            data.setdefault("referral_warnings", {})
            data.setdefault("referral_monthly", {})
            data.setdefault("user_names", {})
            data.setdefault("user_username", {})
            data.setdefault("usernames", {})
            data.setdefault("manual_access_granted", [])
            data.setdefault("manual_anatomy_demo_granted", [])
            data.setdefault("assistant_admins", [])
            data.setdefault("payment_admins", [])
            data.setdefault("referral_battle", None)
            data.setdefault("donations_stars_total", 0)
            data.setdefault("donations_stars_count", 0)
            data.setdefault("donor_stars", {})
            data.setdefault("donor_rubles", {})
            data.setdefault("donor_hide_name", {})
            data.setdefault("temporary_access", {})
            data.setdefault("subscriptions", {})
            data.setdefault("section_promos", {})
            data.setdefault("histology_warnings", {})
            data.setdefault("histology_temp_access", {})
            data.setdefault("rollcall_confirmed", {})
            data.setdefault("anatomy_latin_scores", {})
            data.setdefault("anatomy_exam_test_scores", {})
            data.setdefault("anatomy_exam_test_mode", {})
            data.setdefault("anatomy_exam_flash_scores", {})
            data.setdefault("ai_usage", {})
            data.setdefault("ai_cost_totals", {"requests": 0, "input_tokens": 0, "output_tokens": 0, "cost_usd": 0.0})
            data.setdefault("ai_answer_cache", {})
            data.setdefault("ai_cost_windows", {
                "hour_key": None, "hour_cost_usd": 0.0, "day_key": None, "day_cost_usd": 0.0,
                "breaker_tripped": False, "breaker_alerted": False,
            })
            data.setdefault("ai_raw_text_aliases", {})
            data.setdefault("processed_payment_charge_ids", {})
            data.setdefault("physiology_progress", {})
            data.setdefault("physiology_favorites", {})
            return data
        except (json.JSONDecodeError, OSError):
            logger.exception("Не удалось прочитать %s, статистика будет создана заново", STATS_FILE)
    return {
        "total_users": set(),
        "start_count": 0,
        "random_ticket_used": 0,
        "random_question_used": 0,
        "question_opened": {},
        "broadcast_count": 0,
        "helperchat_promo_seen": {},
        "referrals": {},
        "referred_by": {},
        "referral_warnings": {},
        "referral_monthly": {},
        "user_names": {},
        "user_username": {},
        "usernames": {},
        "manual_access_granted": [],
        "manual_anatomy_demo_granted": [],
        "assistant_admins": [],
        "payment_admins": [],
        "referral_battle": None,
        "donations_stars_total": 0,
        "donations_stars_count": 0,
        "donor_stars": {},
        "donor_rubles": {},
        "donor_hide_name": {},
        "temporary_access": {},
        "subscriptions": {},
        "section_promos": {},
        "histology_warnings": {},
        "histology_temp_access": {},
        "rollcall_confirmed": {},
        "anatomy_latin_scores": {},
        "anatomy_exam_test_scores": {},
        "anatomy_exam_test_mode": {},
        "anatomy_exam_flash_scores": {},
        "ai_usage": {},
        "ai_cost_totals": {"requests": 0, "input_tokens": 0, "output_tokens": 0, "cost_usd": 0.0},
        "ai_answer_cache": {},
        "ai_cost_windows": {
            "hour_key": None, "hour_cost_usd": 0.0, "day_key": None, "day_cost_usd": 0.0,
            "breaker_tripped": False, "breaker_alerted": False,
        },
        "ai_raw_text_aliases": {},
        "processed_payment_charge_ids": {},
        "physiology_progress": {},
        "physiology_favorites": {},
    }

# Один воркер сериализует записи на диск и не даёт им блокировать event loop бота.
_stats_executor = ThreadPoolExecutor(max_workers=1, thread_name_prefix="stats-writer")

def _write_stats_file(data: dict) -> None:
    tmp_path = f"{STATS_FILE}.tmp"
    with open(tmp_path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    os.replace(tmp_path, STATS_FILE)

def _log_stats_write_result(future) -> None:
    exc = future.exception()
    if exc is not None:
        logger.error("Не удалось сохранить статистику: %s", exc)

def save_stats() -> None:
    # Снимок делаем сразу (deepcopy — быстро), сама запись на диск уходит в отдельный поток.
    data = copy.deepcopy(stats)
    data["total_users"] = list(data["total_users"])
    future = _stats_executor.submit(_write_stats_file, data)
    future.add_done_callback(_log_stats_write_result)

stats = load_stats()

# ==================== ДОСТУП: РЕФЕРАЛЫ, ПОДПИСКИ, ГЕЙТЫ ====================
# Реферальная система, платная подписка, промо-окна разделов и гейты по предметам вынесены
# в services/access.py (чистая предикатная/учётная логика, без UI и без aiogram-хендлеров).
# Ниже — импорт и плоский реэкспорт: весь остальной код файла (включая уже вынесенные
# handlers/*.py) продолжает обращаться к этим именам как раньше, без префикса access.
from services import access  # noqa: E402 — mid-file by design, see above

is_admin = access.is_admin
is_assistant_admin = access.is_assistant_admin
is_admin_or_assistant = access.is_admin_or_assistant
is_payment_admin = access.is_payment_admin
start_section_promo = access.start_section_promo
is_section_promo_active = access.is_section_promo_active
BOT_USERNAME = access.BOT_USERNAME
REFERRAL_FULL_ACCESS_THRESHOLD = access.REFERRAL_FULL_ACCESS_THRESHOLD
REFERRAL_WARNING_THRESHOLD = access.REFERRAL_WARNING_THRESHOLD
REFERRAL_WARNING_COOLDOWN_SECONDS = access.REFERRAL_WARNING_COOLDOWN_SECONDS
TEMP_ACCESS_GRANT_SECONDS = access.TEMP_ACCESS_GRANT_SECONDS
GLOBAL_PROMO_SECONDS = access.GLOBAL_PROMO_SECONDS
GLOBAL_PROMO_12H_SECONDS = access.GLOBAL_PROMO_12H_SECONDS
get_referral_link = access.get_referral_link
get_referral_count = access.get_referral_count
get_referral_count_this_month = access.get_referral_count_this_month
_current_referral_month_key = access._current_referral_month_key
get_temp_access_expiry = access.get_temp_access_expiry
has_temp_access = access.has_temp_access
TIER1_HISTOLOGY_DEADLINE = access.TIER1_HISTOLOGY_DEADLINE
JULY_END_2026 = access.JULY_END_2026
OCT_2026_CUTOFF = access.OCT_2026_CUTOFF
NOV_END_2026_CUTOFF = access.NOV_END_2026_CUTOFF
FEB_2027_CUTOFF = access.FEB_2027_CUTOFF
SECOND_YEAR_END_2027 = access.SECOND_YEAR_END_2027
NOV_1_2026_CUTOFF = access.NOV_1_2026_CUTOFF
JAN_1_2027_CUTOFF = access.JAN_1_2027_CUTOFF
MAR_1_2027_CUTOFF = access.MAR_1_2027_CUTOFF
FIRST_YEAR_END_2027 = access.FIRST_YEAR_END_2027
LEGACY_PAID_AI_MONTHLY_BONUS = access.LEGACY_PAID_AI_MONTHLY_BONUS
SUBSCRIPTION_TIERS = access.SUBSCRIPTION_TIERS
ACTIVE_SUBSCRIPTION_TIERS = access.ACTIVE_SUBSCRIPTION_TIERS
SEPTEMBER_PRICE_INCREASE = access.SEPTEMBER_PRICE_INCREASE
september_price = access.september_price
DISCOUNT_RATE = access.DISCOUNT_RATE
discount_price = access.discount_price
get_tier_price_line = access.get_tier_price_line
sorted_active_tiers = access.sorted_active_tiers
cheapest_active_tier = access.cheapest_active_tier
cheapest_gated3_tier = access.cheapest_gated3_tier
cheapest_histology_tier = access.cheapest_histology_tier
cheapest_anatomy_tier = access.cheapest_anatomy_tier
cheapest_biology_download_tier = access.cheapest_biology_download_tier
get_subscription = access.get_subscription
has_active_subscription = access.has_active_subscription
has_subject_access = access.has_subject_access
_sub_has_histology = access._sub_has_histology
_sub_has_anatomy = access._sub_has_anatomy
_sub_has_biology_download = access._sub_has_biology_download
has_subscription_scope_all = access.has_subscription_scope_all
has_subscription_histology_access = access.has_subscription_histology_access
has_subscription_anatomy_access = access.has_subscription_anatomy_access
biology_tickets_download_ok = access.biology_tickets_download_ok
grant_subscription = access.grant_subscription
has_free_access = access.has_free_access
get_exhausted_users = access.get_exhausted_users
get_below_threshold_users = access.get_below_threshold_users
register_referral = access.register_referral
GATED_CALLBACKS_BIOLOGY = access.GATED_CALLBACKS_BIOLOGY
GATED_PREFIXES_BIOLOGY = access.GATED_PREFIXES_BIOLOGY
GATED_CALLBACKS_PHYSICS = access.GATED_CALLBACKS_PHYSICS
GATED_PREFIXES_PHYSICS = access.GATED_PREFIXES_PHYSICS
GATED_CALLBACKS_CHEMISTRY = access.GATED_CALLBACKS_CHEMISTRY
GATED_PREFIXES_CHEMISTRY = access.GATED_PREFIXES_CHEMISTRY
GATED_CALLBACKS = access.GATED_CALLBACKS
GATED_PREFIXES = access.GATED_PREFIXES
get_gated_subject = access.get_gated_subject
is_gated_callback = access.is_gated_callback
chemistry_tickets_access_ok = access.chemistry_tickets_access_ok


bot = Bot(token=BOT_TOKEN)
dp = Dispatcher()

# ==================== ЕЖЕДНЕВНОЕ НАПОМИНАНИЕ ПРО HELPERCHAT_BOT ====================
HELPERCHAT_PROMO_ENABLED = False  # временно отключено по просьбе — включить обратно, поставив True
HELPERCHAT_URL = "https://t.me/Helperchat_bot?start=vmeda"

def get_helperchat_promo_keyboard():
    builder = InlineKeyboardBuilder()
    builder.button(text="🚀 Запустить Helperchat_bot", url=HELPERCHAT_URL)
    return builder.as_markup()

async def send_helperchat_promo_if_new_day(user_id: int) -> None:
    today = local_today().isoformat()
    seen = stats["helperchat_promo_seen"]
    if seen.get(str(user_id)) == today:
        return
    seen[str(user_id)] = today
    save_stats()
    try:
        await bot.send_message(
            user_id,
            "🚀 <b>Не забудь запустить нашего бота-помощника</b>\n\n"
            "Он тоже пригодится для подготовки — жми и запускай в один тап:\n"
            f"👉 {HELPERCHAT_URL}",
            parse_mode="HTML",
            reply_markup=get_helperchat_promo_keyboard()
        )
    except Exception:
        logger.exception("Не удалось отправить напоминание про Helperchat_bot пользователю %s", user_id)

@dp.update.outer_middleware()
async def helperchat_promo_middleware(handler, event: Update, data):
    user = None
    if event.message:
        user = event.message.from_user
    elif event.callback_query:
        user = event.callback_query.from_user
    if HELPERCHAT_PROMO_ENABLED and user and not user.is_bot:
        await send_helperchat_promo_if_new_day(user.id)
    return await handler(event, data)

# Реферальная система и платная подписка (SUBSCRIPTION_TIERS и вся связанная логика)
# перенесены в services/access.py — см. импорт и реэкспорт в начале файла.

SUBJECT_TITLES = {"biology": "Биологии", "physics": "Физике", "chemistry": "Химии"}
ADMIN_SUBJECT_LABELS_RU = {"Биология": "biology", "Физика": "physics", "Химия": "chemistry"}

def get_admin_tier_reply_keyboard() -> ReplyKeyboardMarkup:
    rows = [[KeyboardButton(text=f"{t} — {cfg['short']} — {cfg['price_rub']}₽")] for t, cfg in ACTIVE_SUBSCRIPTION_TIERS.items()]
    rows.append([KeyboardButton(text="❌ Отмена")])
    return ReplyKeyboardMarkup(keyboard=rows, resize_keyboard=True, one_time_keyboard=True)

def get_admin_subject_reply_keyboard() -> ReplyKeyboardMarkup:
    return ReplyKeyboardMarkup(
        keyboard=[
            [KeyboardButton(text="Биология")],
            [KeyboardButton(text="Физика")],
            [KeyboardButton(text="Химия")],
            [KeyboardButton(text="❌ Отмена")],
        ],
        resize_keyboard=True, one_time_keyboard=True
    )

def get_subscription_scope_label(sub: dict) -> str:
    restricted = sub.get("restricted_subject")
    if restricted:
        return f"только к {SUBJECT_TITLES[restricted]}"
    if sub.get("scope") == "all":
        return "ко всем разделам бота"
    parts = ["Биологии, Физике и Химии"]
    if _sub_has_histology(sub):
        parts.append("Гистологии")
    if _sub_has_anatomy(sub):
        parts.append("Анатомии")
    if len(parts) == 3:
        return "ко всем разделам бота"
    return "к " + ", ".join(parts)

def get_referral_status_text(user_id: int) -> str:
    # count — рефералы ИМЕННО этого календарного месяца (то, что реально сравнивается с порогом,
    # см. get_referral_count_this_month/services/access.py: доступ по рефералам с этой версии не
    # разовый навсегда, а ЕЖЕМЕСЯЧНО обновляемый); total — рефералы за всё время, только для
    # контекста ("сколько всего людей ты когда-либо привёл"), на сам гейт не влияет.
    count = get_referral_count_this_month(user_id)
    total = get_referral_count(user_id)
    link = get_referral_link(user_id)
    if has_active_subscription(user_id):
        sub = get_subscription(user_id)
        cfg = SUBSCRIPTION_TIERS.get(sub["tier"], {})
        scope_label = get_subscription_scope_label(sub)
        return (
            f"👥 <b>Твои приглашения</b>\n{DIVIDER}\n\n"
            f"💎 У тебя активна подписка «{cfg.get('title', '')}» — доступ {scope_label}, "
            f"{format_subscription_expiry(sub['expires'])}.\n\n"
            "Рефералы тебе не нужны, но можно продолжать приглашать друзей и участвовать "
            "в <b>битве рефералов</b> за призы!\n\n"
            f"Твоя ссылка:\n{link}"
        )
    manual = user_id in stats["manual_access_granted"]
    if count >= REFERRAL_FULL_ACCESS_THRESHOLD or manual:
        extra = f"Приглашено в этом месяце: <b>{count}</b>\n" if count > 0 else ""
        reset_note = "" if manual else (
            f"\n⚠️ Условие обновляется каждый месяц: чтобы доступ не закрылся, в следующем "
            f"месяце снова понадобится пригласить {REFERRAL_FULL_ACCESS_THRESHOLD} новых друзей.\n"
        )
        return (
            f"👥 <b>Твои приглашения</b>\n{DIVIDER}\n\n"
            f"{extra}"
            "Доступ ко всем разделам бота открыт. Спасибо! 🎉\n"
            f"{reset_note}\n"
            "⚔️ А ещё сейчас можно побороться за призы в <b>битве рефералов</b> — "
            "приглашай друзей дальше и попади в топ-5!\n\n"
            f"Всего приглашено за всё время: <b>{total}</b>\n\n"
            f"Твоя ссылка (можно приглашать ещё):\n{link}"
        )
    if has_temp_access(user_id):
        remaining = format_time_left(get_temp_access_expiry(user_id) - time.time())
        return (
            f"👥 <b>Твои приглашения</b>\n{DIVIDER}\n\n"
            f"🎁 Тебе временно открыт полный доступ ко всем разделам бота — осталось "
            f"<b>{remaining}</b>.\n\n"
            f"Приглашено в этом месяце: <b>{count}</b> из {REFERRAL_FULL_ACCESS_THRESHOLD}\n\n"
            "Пригласи друзей уже сейчас, чтобы доступ остался открытым и после окончания "
            f"временного периода:\n{link}"
        )
    warn_count = stats["referral_warnings"].get(str(user_id), {}).get("count", 0)
    remaining_free = max(REFERRAL_WARNING_THRESHOLD - warn_count, 0)
    remaining_refs = max(REFERRAL_FULL_ACCESS_THRESHOLD - count, 0)
    if remaining_refs <= 1:
        invite_line = (
            "Отправь эту ссылку ещё одному другу — как только он нажмёт /start, "
            "у тебя откроется полный доступ ко всем разделам бота:"
        )
    else:
        friends_word = "двум друзьям" if remaining_refs == 2 else f"{remaining_refs} друзьям"
        invite_line = (
            f"Отправь эту ссылку ещё {friends_word} — как только они нажмут /start, "
            "у тебя откроется полный доступ ко всем разделам бота:"
        )
    return (
        f"👥 <b>Пригласи друзей</b>\n{DIVIDER}\n\n"
        f"{invite_line}\n\n"
        f"{link}\n\n"
        f"Приглашено в этом месяце: <b>{count}</b> из {REFERRAL_FULL_ACCESS_THRESHOLD}\n"
        f"Осталось бесплатных заходов без рефералов: <b>{remaining_free}</b>\n\n"
        f"💎 Не хочешь ждать друзей? Открой доступ сразу оплатой!\n\n"
        f"🔥 Самые выгодные варианты — «{SUBSCRIPTION_TIERS[21]['short']}» за "
        f"{SUBSCRIPTION_TIERS[21]['price_rub']}₽ или «{SUBSCRIPTION_TIERS[26]['short']}» за "
        f"{SUBSCRIPTION_TIERS[26]['price_rub']}₽ ({SUBSCRIPTION_TIERS[26]['badge']}).\n\n"
        f"Также доступна подписка от {cheapest_gated3_tier()['price_rub']}₽. "
        "Жми «💎 Открыть доступ без рефералов» ниже.\n\n"
        "🌐 А ещё рекомендуем пользоваться @Medical_vpn_bot — жми «🚀 Запустить Medical_vpn_bot» ниже."
    )

RANK_MEDALS = ["🥇", "🥈", "🥉"]

def get_referral_leaderboard_keyboard():
    builder = InlineKeyboardBuilder()
    builder.button(text="🔄 Обновить", callback_data="referral_leaderboard")
    builder.button(text="🔙 Назад в меню", callback_data="back_to_main")
    builder.adjust(1)
    return builder.as_markup()

def get_referral_leaderboard_text(user_id: int = None) -> str:
    referrals = stats["referrals"]
    names = stats["user_names"]
    ranked = sorted(referrals.items(), key=lambda kv: len(kv[1]), reverse=True)
    ranked = [(uid, refs) for uid, refs in ranked if len(refs) > 0]
    if not ranked:
        return f"🏆 <b>Рейтинг приглашений</b>\n{DIVIDER}\n\nПока никто никого не пригласил — стань первым!"
    top = ranked[:10]
    lines = [f"🏆 <b>Рейтинг приглашений</b>\n{DIVIDER}\n"]
    uid_str = str(user_id) if user_id is not None else None
    for i, (uid, refs) in enumerate(top):
        rank_icon = RANK_MEDALS[i] if i < 3 else f"{i+1}."
        name = names.get(uid, f"Пользователь {uid}")
        you = " 👈 ты" if uid == uid_str else ""
        lines.append(f"{rank_icon} {name} — <b>{len(refs)}</b>{you}")
    if uid_str and uid_str not in dict(top):
        pos = next((i for i, (uid, _) in enumerate(ranked) if uid == uid_str), None)
        if pos is not None:
            lines.append("…")
            name = names.get(uid_str, "Ты")
            lines.append(f"{pos + 1}. {name} — <b>{len(referrals[uid_str])}</b> 👈 ты")
    lines.append("")
    lines.append(f"👤 Всего участников: <b>{len(ranked)}</b>")
    return "\n".join(lines)

# ==================== БИТВА РЕФЕРАЛОВ (ЛИМИТИРОВАННОЕ СОРЕВНОВАНИЕ) ====================
BATTLE_DURATION_SECONDS = 7 * 24 * 60 * 60
BATTLE_PLACE_COUNT = 5
BATTLE_PLACE_ICONS = ["🥇", "🥈", "🥉", "🏅", "🎖"]

# 1-3 место разыгрываются только среди тех, кто пригласил от этого числа друзей за битву.
# 4-5 место — без минимума, отдаются следующим по рейтингу.
BATTLE_TOP3_MIN_REFERRALS = 30

MEDICAL_VPN_URL = "https://t.me/Medical_vpn_bot?start=vmeda"

BATTLE_PRIZE_LABELS = [
    'подписка «6 лет — абсолютно всё» (<b>6 лет</b>, все предметы + Анатомия/Гистология) в '
    '<a href="https://t.me/VMEDA_examen_bot">VMEDA_examen_bot</a> + ВПН на <b>год</b> в '
    f'<a href="{MEDICAL_VPN_URL}">Medical_vpn_bot</a>',
    'полный доступ ко всем разделам <a href="https://t.me/VMEDA_examen_bot">VMEDA_examen_bot</a> '
    'на <b>год</b> + ВПН на <b>полгода</b> в '
    f'<a href="{MEDICAL_VPN_URL}">Medical_vpn_bot</a>',
    'Химия, Физика, Биология и ранний доступ к Гистологии <b>навсегда</b> в '
    '<a href="https://t.me/VMEDA_examen_bot">VMEDA_examen_bot</a> + ВПН на <b>месяц</b> в '
    f'<a href="{MEDICAL_VPN_URL}">Medical_vpn_bot</a>',
    'Химия, Физика, Биология и ранний доступ к Гистологии <b>навсегда</b> в '
    '<a href="https://t.me/VMEDA_examen_bot">VMEDA_examen_bot</a> + доступ к '
    '<a href="https://t.me/Helperchat_bot">Helperchat_bot</a>',
    'Химия, Физика, Биология и ранний доступ к Гистологии <b>навсегда</b> в '
    '<a href="https://t.me/VMEDA_examen_bot">VMEDA_examen_bot</a> + доступ к '
    '<a href="https://t.me/Helperchat_bot">Helperchat_bot</a>',
]
BATTLE_CHANNEL_POSTING_NOTICE = "📢 <b>ПОСТИНГ В TELEGRAM-КАНАЛЫ РАЗРЕШЁН 🤝</b>"

# Денежная оценка приза за место (задана вручную, не выводится из цен тарифов подписки).
BATTLE_PRIZE_VALUES_RUB = [5400, 2000, 1800, 1599, 1599]

def format_rub(amount: int) -> str:
    return f"{amount:,}".replace(",", " ")

def format_battle_savings_line(place_index: int) -> str:
    return f"💰🔥 Экономия: <b>~{format_rub(BATTLE_PRIZE_VALUES_RUB[place_index])}₽</b> 🔥💰"

def battle_place_icon(i: int) -> str:
    return BATTLE_PLACE_ICONS[i] if i < len(BATTLE_PLACE_ICONS) else f"{i + 1}."

def format_battle_prizes_block() -> str:
    lines = [
        f"{BATTLE_PLACE_ICONS[i]} <b>{i + 1} место</b> — {BATTLE_PRIZE_LABELS[i]}\n{format_battle_savings_line(i)}"
        for i in range(BATTLE_PLACE_COUNT)
    ]
    lines.append(
        f"\n🔒 1-3 место — только для тех, кто пригласит от <b>{BATTLE_TOP3_MIN_REFERRALS}</b> друзей за битву. "
        "4-5 место — без минимума, по числу приглашений."
    )
    return "\n".join(lines)

def is_battle_active() -> bool:
    battle = stats.get("referral_battle")
    return bool(battle and battle.get("active") and time.time() < battle.get("end_ts", 0))

def get_battle_gained(user_id: int) -> int:
    battle = stats.get("referral_battle")
    if not battle:
        return 0
    uid_str = str(user_id)
    current = len(stats["referrals"].get(uid_str, []))
    start = battle.get("snapshot", {}).get(uid_str, 0)
    return max(current - start, 0)

def get_battle_leaderboard(limit: int | None = 10):
    battle = stats.get("referral_battle")
    if not battle:
        return []
    snapshot = battle.get("snapshot", {})
    gained = []
    for uid_str, refs in stats["referrals"].items():
        diff = len(refs) - snapshot.get(uid_str, 0)
        if diff > 0:
            gained.append((uid_str, diff))
    gained.sort(key=lambda kv: kv[1], reverse=True)
    return gained if limit is None else gained[:limit]

def resolve_battle_winners() -> list:
    """5 призовых мест. 1-3 место — только для участников с BATTLE_TOP3_MIN_REFERRALS+
    рефералами за битву. 4-5 место — без минимума, отдаются следующим по рейтингу.
    Возвращает список из BATTLE_PLACE_COUNT элементов: (uid_str, diff) или None, если
    место не разыграно."""
    full = get_battle_leaderboard(limit=None)
    winners = [None] * BATTLE_PLACE_COUNT
    used_uids = set()
    qualifying = [e for e in full if e[1] >= BATTLE_TOP3_MIN_REFERRALS]
    for i in range(3):
        if i < len(qualifying):
            winners[i] = qualifying[i]
            used_uids.add(qualifying[i][0])
    remaining = [e for e in full if e[0] not in used_uids]
    for i, entry in enumerate(remaining[:2]):
        winners[3 + i] = entry
    return winners

def format_time_left(seconds: float) -> str:
    seconds = max(int(seconds), 0)
    d, rem = divmod(seconds, 86400)
    h, rem = divmod(rem, 3600)
    m = rem // 60
    if d > 0:
        return f"{d}д {h}ч"
    return f"{h}ч {m}мин"

def format_battle_duration() -> str:
    days = BATTLE_DURATION_SECONDS // 86400
    if days == 7:
        return "неделю"
    if days > 0:
        return f"{days} дней"
    hours = BATTLE_DURATION_SECONDS // 3600
    return f"{hours} часов"

def start_referral_battle() -> None:
    now = time.time()
    snapshot = {uid: len(refs) for uid, refs in stats["referrals"].items()}
    stats["referral_battle"] = {
        "active": True,
        "start_ts": now,
        "end_ts": now + BATTLE_DURATION_SECONDS,
        "snapshot": snapshot,
        "results": None,
    }
    save_stats()

def get_battle_keyboard():
    builder = InlineKeyboardBuilder()
    builder.button(text="🔄 Обновить", callback_data="referral_battle")
    builder.button(text="👥 Пригласить друзей", callback_data="referral_info")
    builder.button(text="🔙 Назад в меню", callback_data="back_to_main")
    builder.adjust(1)
    return builder.as_markup()

def get_battle_text(user_id: int) -> str:
    if not is_battle_active():
        battle = stats.get("referral_battle")
        results = battle.get("results") if battle else None
        if results and any(w is not None for w in results):
            medals_lines = []
            for i, w in enumerate(results):
                if w is None:
                    continue
                uid_str, diff = w
                name = stats["user_names"].get(uid_str, f"Пользователь {uid_str}")
                medals_lines.append(f"{battle_place_icon(i)} {i + 1} место — {name} — <b>{diff}</b>")
            results_block = "\n".join(medals_lines)
            return (
                f"⚔️ <b>Битва рефералов</b>\n{DIVIDER}\n\n"
                "Сейчас битва не идёт. Результаты последней битвы:\n\n"
                f"{results_block}\n\n"
                "Следи за объявлениями — как только стартует новая битва, "
                f"у тебя будет {format_battle_duration()}, чтобы побороться за призы:\n\n"
                f"{format_battle_prizes_block()}"
            )
        return (
            f"⚔️ <b>Битва рефералов</b>\n{DIVIDER}\n\n"
            "Сейчас битва не идёт. Следи за объявлениями — как только стартует новая, "
            f"у тебя будет {format_battle_duration()}, чтобы побороться за призы:\n\n"
            f"{format_battle_prizes_block()}"
        )
    battle = stats["referral_battle"]
    remaining = format_time_left(battle["end_ts"] - time.time())
    my_gained = get_battle_gained(user_id)
    leaderboard = get_battle_leaderboard()
    uid_str = str(user_id)
    lines = [
        f"⚔️ <b>Битва рефералов — идёт!</b>\n{DIVIDER}\n",
        BATTLE_CHANNEL_POSTING_NOTICE,
        "",
        f"⏳ Осталось: <b>{remaining}</b>",
        "🎁 Призы для топ-5:",
        format_battle_prizes_block(),
        "",
        f"🙋 Твой результат за битву: <b>{my_gained}</b>",
        "",
    ]
    if leaderboard:
        lines.append("<b>Текущий рейтинг битвы:</b>")
        for i, (uid, diff) in enumerate(leaderboard):
            name = stats["user_names"].get(uid, f"Пользователь {uid}")
            you = " 👈 ты" if uid == uid_str else ""
            lines.append(f"{battle_place_icon(i)} {name} — <b>{diff}</b>{you}")
    else:
        lines.append("Пока никто не пригласил друзей в рамках битвы — стань первым!")
    lines.append("")
    lines.append(f"Твоя ссылка:\n{get_referral_link(user_id)}")
    return "\n".join(lines)

async def _broadcast_to(user_ids, text: str, keyboard=None) -> None:
    for user_id in list(user_ids):
        try:
            await bot.send_message(user_id, text, parse_mode="HTML", reply_markup=keyboard, disable_web_page_preview=True)
        except Exception:
            logger.exception("Не удалось отправить рассылку пользователю %s", user_id)
        await asyncio.sleep(0.05)

async def _broadcast(text: str, keyboard=None) -> None:
    await _broadcast_to(stats["total_users"], text, keyboard)

async def announce_global_promo_start() -> None:
    text = (
        "🎉🚀 <b>ВСЕ ОГРАНИЧЕНИЯ СНЯТЫ НА 24 ЧАСА!</b> 🚀🎉\n"
        f"{DIVIDER}\n\n"
        "Абсолютно все разделы — Биология, Физика, Химия, Гистология — сейчас "
        "полностью бесплатны для всех, без рефералов и подписки.\n\n"
        "Успей позаниматься, пока открыто — доступ вернётся к обычным правилам ровно через сутки! ⏳"
    )
    await _broadcast(text)

async def announce_global_promo_12h_start() -> None:
    text = (
        "🎉🚀 <b>ВСЕ ОГРАНИЧЕНИЯ СНЯТЫ НА 12 ЧАСОВ!</b> 🚀🎉\n"
        f"{DIVIDER}\n\n"
        "Абсолютно все разделы — Биология, Физика, Химия, Гистология — сейчас "
        "полностью бесплатны для всех, без рефералов и подписки.\n\n"
        "Успей позаниматься, пока открыто — доступ вернётся к обычным правилам через 12 часов! ⏳"
    )
    await _broadcast(text)

async def announce_battle_start() -> None:
    builder = InlineKeyboardBuilder()
    builder.button(text="⚔️ Битва рефералов", callback_data="referral_battle")
    text = (
        "⚔️🔥 <b>СТАРТУЕТ БИТВА РЕФЕРАЛОВ!</b> 🔥⚔️\n"
        f"{DIVIDER}\n\n"
        f"У тебя есть <b>{format_battle_duration()}</b>, чтобы пригласить в бота как можно больше друзей "
        "и забрать один из пяти эксклюзивных призов:\n\n"
        f"{format_battle_prizes_block()}\n\n"
        f"{DIVIDER}\n\n"
        "Считаются только друзья, приглашённые с этого момента.\n"
        "Следи за живым рейтингом на кнопке «⚔️ Битва рефералов» в главном меню.\n\n"
        f"{BATTLE_CHANNEL_POSTING_NOTICE}\n\n"
        "Погнали! 🚀"
    )
    await _broadcast(text, builder.as_markup())

def get_battle_remind_broadcast_text() -> str:
    battle = stats["referral_battle"]
    remaining = format_time_left(battle["end_ts"] - time.time())
    leaderboard = get_battle_leaderboard()
    lines = [
        "⚔️🔥 <b>БИТВА РЕФЕРАЛОВ ПРОДОЛЖАЕТСЯ!</b> 🔥⚔️\n",
        f"{DIVIDER}\n",
        f"⏳ Осталось: <b>{remaining}</b>\n",
        "🎁 Призы для топ-5:",
        format_battle_prizes_block(),
        "",
    ]
    if leaderboard:
        lines.append("<b>Текущий рейтинг битвы:</b>")
        for i, (uid, diff) in enumerate(leaderboard):
            name = stats["user_names"].get(uid, f"Пользователь {uid}")
            lines.append(f"{battle_place_icon(i)} {name} — <b>{diff}</b>")
        lines.append("")
    lines.append("Успей попасть в топ — жми «👥 Пригласить друзей» в главном меню и забирай свою ссылку!")
    return "\n".join(lines)

def get_access_restored_broadcast_text() -> str:
    return (
        "🎁 <b>Тебе восстановлен доступ!</b>\n"
        f"{DIVIDER}\n\n"
        "Мы заметили, что у тебя закончились бесплатные заходы в разделы Биология, Физика и Химия "
        "без приглашения друзей.\n\n"
        "Специально для тебя доступ ко всем разделам бота открыт заново на <b>7 дней</b> — "
        "взамен, пожалуйста, включи уведомления от бота (в Telegram: настройки чата с ботом → "
        "уведомления), чтобы не пропустить важные новости и новые материалы.\n\n"
        "⏳ Через 7 дней временный доступ закончится, и снова понадобится "
        f"{REFERRAL_FULL_ACCESS_THRESHOLD} приглашённых друга — это правило теперь обновляется "
        "каждый месяц: чтобы доступ оставался открытым, каждый месяц нужны новые приглашённые "
        "друзья, старые в счёт следующего месяца не идут.\n\n"
        "👥 Открыть доступ можно в любой момент — кнопка «Пригласить друзей» в главном меню."
    )

def get_referral_reminder_broadcast_text() -> str:
    cheapest = cheapest_gated3_tier()
    return (
        f"👋 <b>Напоминание</b>\n{DIVIDER}\n\n"
        f"Чтобы бесплатно пользоваться разделами Биология, Физика и Химия, нужно приглашать "
        f"{REFERRAL_FULL_ACCESS_THRESHOLD} новых друзей КАЖДЫЙ МЕСЯЦ — открой «👥 Пригласить "
        "друзей» в главном меню, посмотри свой прогресс за этот месяц и отправь ссылку.\n\n"
        f"💎 Не хочешь ждать друзей? Открой доступ сразу оплатой — подписки от "
        f"{cheapest['price_rub']}₽/{cheapest['price_stars']}⭐. Жми «💎 Подписка» в главном меню."
    )

def get_referral_reminder_broadcast_keyboard():
    builder = InlineKeyboardBuilder()
    builder.button(text="👥 Пригласить друзей", callback_data="referral_info")
    builder.button(text="💎 Подписка", callback_data="subscription_menu")
    builder.adjust(1)
    return builder.as_markup()

DISCOUNT_PROMO_TIER_IDS = (21, 26)  # какие тарифы предлагаются со скидкой в этой рассылке

def get_discount_promo_broadcast_text() -> str:
    t21, t26 = (SUBSCRIPTION_TIERS[t] for t in DISCOUNT_PROMO_TIER_IDS)
    return (
        f"🔥 <b>Скидка {int(DISCOUNT_RATE * 100)}% специально для тебя!</b>\n{DIVIDER}\n\n"
        "Ты ещё не пригласил друзей и пока не открыл доступ к боту — специально для тебя разовая "
        f"скидка {int(DISCOUNT_RATE * 100)}% на два самых выгодных тарифа:\n\n"
        f"{t21['emoji']} «{t21['title']}» — <s>{t21['price_rub']}₽</s> <b>{discount_price(t21['price_rub'])}₽</b>\n"
        f"{t26['emoji']} «{t26['title']}» — <s>{t26['price_rub']}₽</s> <b>{discount_price(t26['price_rub'])}₽</b>\n\n"
        "Жми на кнопку ниже, чтобы забрать скидку — предложение разовое!"
    )

def get_discount_promo_broadcast_keyboard():
    builder = InlineKeyboardBuilder()
    for t in DISCOUNT_PROMO_TIER_IDS:
        cfg = SUBSCRIPTION_TIERS[t]
        builder.button(
            text=f"{cfg['emoji']} {cfg['short']} — {discount_price(cfg['price_rub'])}₽ со скидкой",
            callback_data=f"sub_discount:{t}"
        )
    builder.adjust(1)
    return builder.as_markup()

def get_battle_results_announcement_text(winners: list) -> str:
    awarded = [(i, w) for i, w in enumerate(winners) if w is not None]
    if not awarded:
        return (
            f"🏁 <b>Битва рефералов завершена!</b>\n{DIVIDER}\n\n"
            "За время битвы никто не пригласил новых друзей — приз не разыгран в этот раз."
        )
    lines = [f"🏁 <b>Битва рефералов завершена!</b>\n{DIVIDER}\n", "Победители:"]
    for i, (uid_str, diff) in awarded:
        name = stats["user_names"].get(uid_str, f"Пользователь {uid_str}")
        lines.append(f"{BATTLE_PLACE_ICONS[i]} <b>{i + 1} место</b> — {name} — <b>{diff}</b> приглашённых")
        lines.append(f"🎁 {BATTLE_PRIZE_LABELS[i]}")
        lines.append(format_battle_savings_line(i))
        lines.append("")
    if any(winners[i] is None for i in range(3)):
        lines.append(
            f"⚠️ Часть мест в топ-3 не разыграна — не набралось участников от "
            f"{BATTLE_TOP3_MIN_REFERRALS} приглашённых.\n"
        )
    lines.append("Администратор свяжется с победителями лично 🤝")
    return "\n".join(lines)

async def resolve_referral_battle() -> None:
    battle = stats.get("referral_battle")
    if not battle or not battle.get("active"):
        return
    battle["active"] = False
    winners = resolve_battle_winners()
    battle["results"] = winners
    save_stats()

    result_text = get_battle_results_announcement_text(winners)
    await _broadcast(result_text)

    awarded = [(i, w) for i, w in enumerate(winners) if w is not None]
    if awarded:
        admin_lines = ["🏁 <b>Битва рефералов завершена.</b> Победители (для выдачи приза):"]
        for i, (uid_str, diff) in awarded:
            username = stats["user_username"].get(uid_str)
            handle = f"@{username}" if username else "(нет username)"
            admin_lines.append(f"{BATTLE_PLACE_ICONS[i]} {i + 1} место — ID <code>{uid_str}</code> {handle} — {diff} рефералов")
        admin_text = "\n".join(admin_lines)
        for admin_id in ADMIN_IDS:
            try:
                await bot.send_message(admin_id, admin_text, parse_mode="HTML")
            except Exception:
                logger.exception("Не удалось уведомить админа %s об итогах битвы", admin_id)

async def _battle_timer(end_ts: float) -> None:
    await asyncio.sleep(max(end_ts - time.time(), 0))
    battle = stats.get("referral_battle")
    if battle and battle.get("active") and time.time() >= battle.get("end_ts", 0):
        await resolve_referral_battle()

def resume_battle_timer_if_needed() -> None:
    battle = stats.get("referral_battle")
    if not battle or not battle.get("active"):
        return
    if time.time() >= battle.get("end_ts", 0):
        asyncio.create_task(resolve_referral_battle())
    else:
        asyncio.create_task(_battle_timer(battle["end_ts"]))

# register_referral перенесена в services/access.py — см. импорт и реэкспорт в начале файла.

def track_user_identity(user) -> None:
    """Обновляет карты имя/username <-> id, чтобы админ мог находить пользователей по @username."""
    uid_str = str(user.id)
    changed = False
    new_name = html.escape(user.full_name) if user.full_name else f"Пользователь {user.id}"
    if stats["user_names"].get(uid_str) != new_name:
        stats["user_names"][uid_str] = new_name
        changed = True
    new_username = (user.username or "").strip().lower() or None
    if stats["user_username"].get(uid_str) != new_username:
        stats["user_username"][uid_str] = new_username
        changed = True
    if new_username and stats["usernames"].get(new_username) != user.id:
        stats["usernames"][new_username] = user.id
        changed = True
    if changed:
        save_stats()

# Константы GATED_CALLBACKS_*/GATED_PREFIXES_* и get_gated_subject()/is_gated_callback()
# перенесены в services/access.py — см. импорт и реэкспорт в начале файла.

@dp.update.outer_middleware()
async def referral_gate_middleware(handler, event: Update, data):
    user = None
    if event.message:
        user = event.message.from_user
    elif event.callback_query:
        user = event.callback_query.from_user
    if not user or user.is_bot:
        return await handler(event, data)

    track_user_identity(user)

    # команды (/start, /stats, /broadcast и т.д.) не блокируем — гейт касается только контента
    if event.message and event.message.text and event.message.text.startswith("/"):
        return await handler(event, data)

    # поддержку автора (пожертвования) не блокируем никому, даже без рефералов
    if event.message and event.message.successful_payment:
        return await handler(event, data)
    if user.id in DONATION_PENDING:
        return await handler(event, data)

    # гейт касается только разделов Биология/Физика/Химия — остальные кнопки
    # (админка, рефералы, битва, поддержка автора, анатомия) доступны всегда
    subject = get_gated_subject(event.callback_query.data or "") if event.callback_query else None
    if event.callback_query and subject is None:
        return await handler(event, data)

    # у callback'ов проверяем доступ именно к ЭТОМУ предмету (тариф «3 дня, 1 предмет» открывает
    # только один) — у обычных сообщений (не привязаны к конкретному предмету) достаточно
    # любого действующего доступа вообще.
    if event.callback_query:
        if has_subject_access(user.id, subject):
            return await handler(event, data)
    elif has_free_access(user.id):
        return await handler(event, data)

    user_id_str = str(user.id)
    entry = stats["referral_warnings"].get(user_id_str, {"count": 0, "last_warn_at": 0})

    if entry["count"] >= REFERRAL_WARNING_THRESHOLD:
        block_text = (
            "🚨❗️ <b>ДОСТУП ЗАКРЫТ!</b> ❗️🚨\n\n"
            "Чтобы продолжить пользоваться ботом бесплатно — <b>пригласи друзей</b>! "
            "Это займёт меньше минуты! ⏱️\n\n"
            f"{get_referral_status_text(user.id)}\n\n"
            "⚡️ Как только твои друзья нажмут /start по этой ссылке — бот <b>сразу</b> станет доступен!"
        )
        try:
            if event.callback_query:
                await event.callback_query.answer("🚨 Доступ закрыт — пригласи друзей! ‼️", show_alert=True)
                await event.callback_query.message.answer(block_text, parse_mode="HTML", reply_markup=get_subscription_teaser_keyboard())
            elif event.message:
                await event.message.answer(block_text, parse_mode="HTML", reply_markup=get_subscription_teaser_keyboard())
        except Exception:
            logger.exception("Не удалось отправить сообщение о блокировке пользователю %s", user.id)
        return  # обработчик НЕ вызываем — доступ закрыт

    now = time.time()
    if now - entry.get("last_warn_at", 0) >= REFERRAL_WARNING_COOLDOWN_SECONDS:
        entry["count"] += 1
        entry["last_warn_at"] = now
        stats["referral_warnings"][user_id_str] = entry
        save_stats()
        remaining = REFERRAL_WARNING_THRESHOLD - entry["count"]
        warn_text = (
            "⚠️❗️ <b>ВНИМАНИЕ! Пригласи друзей!</b> ❗️⚠️\n\n"
            f"{get_referral_status_text(user.id)}"
            if remaining > 0 else
            "🚨‼️ <b>ПОСЛЕДНЕЕ ПРЕДУПРЕЖДЕНИЕ!</b> ‼️🚨\n\n"
            "В следующий раз доступ будет <b>полностью закрыт</b>, пока не пригласишь друзей!\n\n"
            f"{get_referral_status_text(user.id)}"
        )
        try:
            if event.callback_query:
                await event.callback_query.message.answer(warn_text, parse_mode="HTML", reply_markup=get_subscription_teaser_keyboard())
            elif event.message:
                await event.message.answer(warn_text, parse_mode="HTML", reply_markup=get_subscription_teaser_keyboard())
        except Exception:
            logger.exception("Не удалось отправить предупреждение о реферале пользователю %s", user.id)

    return await handler(event, data)

# ==================== ПЕРЕКЛИЧКА ГРУПП ====================
# Собираем по одному представителю от каждой группы, чтобы быть с ними на связи. Список
# групп генерируется по шаблону, не хранится в JSON — сами группы никогда не меняются местами.
ROLLCALL_GROUP_COUNT = 45

def rollcall_group_name(n: int) -> str:
    return f"25-ЛД/СТ-{n}"

def get_rollcall_menu_text() -> str:
    confirmed = len(stats["rollcall_confirmed"])
    return (
        f"📋 <b>Перекличка групп</b>\n{DIVIDER}\n\n"
        "Собираем по одному представителю от каждой группы, чтобы быть на связи.\n\n"
        "⚡ <b>ПЕРВЫМ ВЫБЕРИ СВОЮ ГРУППУ</b> и получи бонусную подписку на неделю!\n\n"
        f"Подтверждено представителей: <b>{confirmed}</b> из {ROLLCALL_GROUP_COUNT}\n\n"
        "Выбери номер своей группы:"
    )

def get_rollcall_menu_keyboard():
    builder = InlineKeyboardBuilder()
    for n in range(1, ROLLCALL_GROUP_COUNT + 1):
        group = rollcall_group_name(n)
        if group in stats["rollcall_confirmed"]:
            builder.button(text=f"✅ {group}", callback_data="rollcall_taken")
        else:
            builder.button(text=group, callback_data=f"rollcall_group:{n}")
    builder.adjust(3)
    builder.row(InlineKeyboardButton(text="🔙 Назад в меню", callback_data="back_to_main"))
    return builder.as_markup()

def get_rollcall_group_text(n: int) -> str:
    group = rollcall_group_name(n)
    return (
        f"👥 <b>Группа {group}</b>\n{DIVIDER}\n\n"
        "Нажми на кнопку ниже — откроется чат с @vmeda_helper, сообщение с номером твоей "
        "группы уже будет готово. Напиши его и подтверди, что ты из этой группы — как только "
        "это проверят, тебе включат бонусную подписку на неделю.\n\n"
        "Спасибо, что будешь на связи! 🙏"
    )

def get_rollcall_group_keyboard(n: int):
    group = rollcall_group_name(n)
    template = f"Привет! Я представитель группы {group}, хочу подтвердить участие в перекличке."
    url = f"{HELPER_ACCOUNT_URL}?text={urllib.parse.quote(template)}"
    builder = InlineKeyboardBuilder()
    builder.row(InlineKeyboardButton(text="💬 Написать @vmeda_helper", url=url))
    builder.row(InlineKeyboardButton(text="🔙 Назад", callback_data="rollcall_menu"))
    return builder.as_markup()

async def notify_admins_of_rollcall_request(n: int, user) -> None:
    group = rollcall_group_name(n)
    text = (
        f"📋 <b>Отклик на перекличку</b>\n{DIVIDER}\n\n"
        f"Группа: <b>{group}</b>\n"
        f"От: {html.escape(user.full_name)} "
        f"({f'@{user.username} ' if user.username else ''}ID <code>{user.id}</code>)\n\n"
        "Проверь в чате с @vmeda_helper, что это правда представитель этой группы, и подтверди:"
    )
    builder = InlineKeyboardBuilder()
    builder.button(text="✅ Подтвердить — выдать бонус", callback_data=f"rollcall_confirm:{n}:{user.id}")
    keyboard = builder.as_markup()
    for admin_id in ADMIN_IDS:
        try:
            await bot.send_message(admin_id, text, parse_mode="HTML", reply_markup=keyboard)
        except Exception:
            logger.exception("Не удалось уведомить админа %s о перекличке", admin_id)

def get_rollcall_announcement_text() -> str:
    return (
        f"📋 <b>Перекличка групп!</b>\n{DIVIDER}\n\n"
        "Собираем по одному представителю от каждой группы, чтобы быть на связи с курсом.\n\n"
        "⚡ <b>ПЕРВЫМ ВЫБЕРИ СВОЮ ГРУППУ</b> и получи бонусную подписку на неделю!\n\n"
        "Жми «📋 Перекличка» в главном меню и выбирай номер своей группы."
    )

def get_rollcall_announcement_keyboard():
    builder = InlineKeyboardBuilder()
    builder.button(text="📋 Перекличка", callback_data="rollcall_menu")
    return builder.as_markup()

@dp.callback_query(F.data == "rollcall_menu")
async def cb_rollcall_menu(callback: CallbackQuery):
    await callback.answer()
    await safe_edit_text(
        callback.message,
        get_rollcall_menu_text(),
        parse_mode="HTML",
        reply_markup=get_rollcall_menu_keyboard()
    )

@dp.callback_query(F.data == "rollcall_taken")
async def cb_rollcall_taken(callback: CallbackQuery):
    await callback.answer("У этой группы уже есть подтверждённый представитель. Спасибо! 🙏", show_alert=True)

@dp.callback_query(F.data.startswith("rollcall_group:"))
async def cb_rollcall_group(callback: CallbackQuery):
    n = int(callback.data.split(":")[1])
    if not (1 <= n <= ROLLCALL_GROUP_COUNT):
        await callback.answer("Группа не найдена", show_alert=True)
        return
    group = rollcall_group_name(n)
    if group in stats["rollcall_confirmed"]:
        await callback.answer("У этой группы уже есть подтверждённый представитель. Спасибо! 🙏", show_alert=True)
        await safe_edit_text(
            callback.message, get_rollcall_menu_text(), parse_mode="HTML", reply_markup=get_rollcall_menu_keyboard()
        )
        return
    await callback.answer()
    await safe_edit_text(
        callback.message,
        get_rollcall_group_text(n),
        parse_mode="HTML",
        reply_markup=get_rollcall_group_keyboard(n),
        disable_web_page_preview=True,
    )
    await notify_admins_of_rollcall_request(n, callback.from_user)

@dp.callback_query(F.data.startswith("rollcall_confirm:"))
async def cb_rollcall_confirm(callback: CallbackQuery):
    if not is_admin(callback.from_user.id):
        await callback.answer()
        return
    _, n_raw, target_id_raw = callback.data.split(":")
    n = int(n_raw)
    target_id = int(target_id_raw)
    if not (1 <= n <= ROLLCALL_GROUP_COUNT):
        await callback.answer("Группа не найдена", show_alert=True)
        return
    group = rollcall_group_name(n)
    existing = stats["rollcall_confirmed"].get(group)
    if existing:
        await callback.answer("У этой группы уже есть подтверждённый представитель", show_alert=True)
        await safe_edit_text(
            callback.message,
            f"✅ Уже подтверждено — представитель группы {group}: "
            f"{format_admin_target_label(None, existing['user_id'])}.",
            parse_mode="HTML"
        )
        return

    stats["rollcall_confirmed"][group] = {"user_id": target_id, "confirmed_at": time.time()}
    stats["temporary_access"][str(target_id)] = time.time() + TEMP_ACCESS_GRANT_SECONDS
    save_stats()
    await callback.answer("Подтверждено ✅", show_alert=True)
    await safe_edit_text(
        callback.message,
        f"✅ Подтверждено — {format_admin_target_label(None, target_id)} представитель группы {group}. Бонус выдан.",
        parse_mode="HTML"
    )
    try:
        await bot.send_message(
            target_id,
            f"🎉 <b>Ты подтверждён как представитель группы {group}!</b>\n\n"
            "Бонусная подписка на неделю активирована — доступ к Биологии, Физике и Химии "
            "открыт на 7 дней без рефералов. Спасибо, что будешь на связи! 🙏",
            parse_mode="HTML"
        )
    except Exception:
        logger.exception("Не удалось уведомить пользователя %s о подтверждении переклички", target_id)

@dp.callback_query(F.data == "admin_announce_rollcall_confirm")
async def cb_admin_announce_rollcall_confirm(callback: CallbackQuery):
    if not (is_admin(callback.from_user.id) or is_payment_admin(callback.from_user.id)):
        await callback.answer()
        return
    await callback.answer()
    builder = InlineKeyboardBuilder()
    builder.button(text="✅ Отправить всем", callback_data="admin_announce_rollcall_go")
    builder.button(text="❌ Отмена", callback_data="admin_announcements_menu")
    builder.adjust(1)
    preview = (
        f"👀 <b>Предпросмотр анонса</b>\n{DIVIDER}\n\n"
        f"{get_rollcall_announcement_text()}\n\n{DIVIDER}\n"
        f"Отправить это всем {len(stats['total_users'])} пользователям?"
    )
    await safe_edit_text(callback.message, preview, parse_mode="HTML", reply_markup=builder.as_markup())

@dp.callback_query(F.data == "admin_announce_rollcall_go")
async def cb_admin_announce_rollcall_go(callback: CallbackQuery):
    if not (is_admin(callback.from_user.id) or is_payment_admin(callback.from_user.id)):
        await callback.answer()
        return
    await callback.answer("📣 Рассылка запущена!", show_alert=True)
    recipients = len(stats["total_users"])
    stats["broadcast_count"] = stats.get("broadcast_count", 0) + 1
    save_stats()
    await _broadcast(get_rollcall_announcement_text(), get_rollcall_announcement_keyboard())
    await safe_edit_text(
        callback.message,
        f"✅ Анонс переклички отправлен (попытка охватить {recipients} пользователей).",
        parse_mode="HTML",
        reply_markup=get_admin_announcements_keyboard(
            "admin_panel" if is_admin(callback.from_user.id) else "payment_admin_panel"
        )
    )

# ==================== ПОДДЕРЖКА АВТОРА ====================
DONATION_PENDING: dict[int, dict] = {}
STARS_MIN, STARS_MAX = 1, 2500
RUBLES_MIN, RUBLES_MAX = 10, 1_000_000
STARS_PRESETS = [25, 50, 100, 250, 500]
RUBLES_PRESETS = [100, 300, 500, 1000, 2000]
HELPER_ACCOUNT_URL = "https://t.me/vmeda_helper"

def get_support_text() -> str:
    return (
        f"😇💰 <b>Поддержка автора</b>\n{DIVIDER}\n\n"
        "Бот без рекламы, а основные разделы всегда можно открыть бесплатно за рефералов.\n\n"
        "На разработку и организацию бота (хостинг, домен, работа над контентом) "
        "потрачено уже около <b>5000₽</b>, а получено с бота — <b>0₽</b>.\n\n"
        "Здесь — просто пожертвование без каких-либо условий, любая сумма. "
        "Если хочешь вместо этого открыть доступ без рефералов — "
        "загляни в «💎 Подписка» в главном меню.\n\n"
        "Можно звёздами Telegram или переводом в рублях — выбери, как удобнее 👇"
    )

def get_support_keyboard(user_id: int):
    hidden = stats.get("donor_hide_name", {}).get(str(user_id), False)
    visibility_label = "🙋 Показывать меня в рейтинге" if hidden else "🙈 Скрыть меня в рейтинге"
    builder = InlineKeyboardBuilder()
    builder.button(text="⭐ Пожертвовать звёзды", callback_data="donate_stars_menu")
    builder.button(text="💵 Перевести рубли", callback_data="donate_rubles_menu")
    builder.button(text="🏆 Лучшие донатеры", callback_data="donors_leaderboard")
    builder.button(text=visibility_label, callback_data="toggle_donor_visibility")
    builder.adjust(1)
    builder.row(InlineKeyboardButton(text="🔙 Назад в меню", callback_data="back_to_main"))
    return builder.as_markup()

def get_support_announcement_text() -> str:
    return (
        f"📣 <b>Новое в боте — раздел «Поддержка автора»!</b>\n{DIVIDER}\n\n"
        "Бот без рекламы, основные разделы всегда можно открыть бесплатно за рефералов — "
        "но на разработку и хостинг уже "
        "потрачено около <b>5000₽</b>, а получено с бота — <b>0₽</b>.\n\n"
        "Теперь его можно поддержать:\n"
        "⭐ звёздами Telegram — сумму выбираешь сам\n"
        "💵 переводом в рублях — тоже любая сумма, реквизиты пришлют в чате с "
        '<a href="https://t.me/vmeda_helper">@vmeda_helper</a>\n\n'
        "А ещё есть рейтинг «🏆 Лучшие донатеры» — топ по звёздам и топ по рублям! "
        "Можно засветить свой ник или остаться анонимом — выбираешь сам.\n\n"
        "Заходи в «😇 Поддержать автора 💰» в главном меню, жертвуй любую сумму — "
        "и попади в топ! 🙏"
    )

def get_support_announcement_keyboard():
    builder = InlineKeyboardBuilder()
    builder.button(text="😇 Поддержать автора 💰", callback_data="support_menu")
    return builder.as_markup()

def get_subscription_announcement_text() -> str:
    tier_lines = "\n".join(
        f"{cfg['emoji']} <b>{cfg['price_rub']}₽ / {cfg['price_stars']}⭐</b> — {cfg['short']}"
        for cfg in ACTIVE_SUBSCRIPTION_TIERS.values()
    )
    return (
        f"💎 <b>Новое в боте — платная подписка без рефералов!</b>\n{DIVIDER}\n\n"
        "Разработка и содержание бота требуют серьёзных затрат — поэтому в дополнение "
        f"к бесплатному доступу за {REFERRAL_FULL_ACCESS_THRESHOLD} рефералов в месяц теперь можно "
        "открыть доступ сразу оплатой:\n\n"
        f"{tier_lines}\n\n"
        "После оплаты правило с рефералами для тебя больше не действует — доступ "
        "открывается сразу и держится всё оплаченное время.\n\n"
        "Загляни в «💎 Подписка» в главном меню, чтобы посмотреть плюсы каждого тарифа 👇"
    )

def get_subscription_announcement_keyboard():
    builder = InlineKeyboardBuilder()
    builder.button(text="💎 Подписка без рефералов", callback_data="subscription_menu")
    return builder.as_markup()

def get_ai_announcement_text() -> str:
    # Каждый активный тариф (20-29) уже несёт свой ai_limit_type/ai_limit — см. CLAUDE.md,
    # "VMedA AI" — поэтому здесь достаточно общей фразы про подписку, без переноса конкретных
    # цифр по тарифам в текст (см. пункт "Values duplicated out of SUBSCRIPTION_TIERS").
    return (
        f"🤖 <b>Новое в боте — VMedA AI!</b>\n{DIVIDER}\n\n"
        "AI-помощник, который разбирает задание по фото или тексту и сразу выдаёт решение: "
        "чёткий ответ и объяснение по шагам. Работает по биологии, физике, химии, анатомии и оперативной хирургии — "
        "тесты, билеты, контрольные, летучки. Просто присылаешь фото — получаешь разбор.\n\n"
        f"Бесплатно — до {AI_FREE_DAILY_LIMIT} запросов в день. Любая подписка добавляет к этому "
        "лимиту дополнительные запросы VMedA AI поверх обычного доступа к разделам бота.\n\n"
        "Жми на кнопку ниже, чтобы попробовать 👇"
    )

def get_ai_announcement_keyboard():
    builder = InlineKeyboardBuilder()
    builder.button(text="🤖 Попробовать VMedA AI", callback_data="ai_menu")
    builder.button(text="💎 Подписка", callback_data="subscription_menu")
    builder.adjust(1)
    return builder.as_markup()

def get_anatomy_announcement_text() -> str:
    tier_lines = " или ".join(
        f"«{cfg['emoji']} {cfg['title']}» ({cfg['price_rub']}₽)"
        for cfg in ACTIVE_SUBSCRIPTION_TIERS.values()
        if cfg.get("anatomy")
    )
    return (
        f"🦴 <b>Открываем раздел «Анатомия»!</b>\n{DIVIDER}\n\n"
        "Самый подробный раздел бота — теперь доступен по подписке 👇\n\n"
        "📚 <b>Что внутри:</b>\n"
        "• 10 модулей по программе Кафарова — 107 тем: от остеологии и миологии до "
        "нервной, сердечно-сосудистой и клинической анатомии\n"
        "• Материал написан по учебнику Гайворонского и методичкам кафедры, с латинской "
        "номенклатурой в каждой теме\n"
        "• Кости черепа, туловища и конечностей разобраны отдельно по каждой кости — "
        "с фото, флэш-карточками, сопоставлением терминов и мнемониками\n"
        "• Атлас с иллюстрациями Неттера и Гайворонского плюс учебные фотопрезентации "
        "кафедры\n"
        "• Тренажёр латинских терминов по каждой кости и теме, а также общий тест по "
        "всей номенклатуре с рейтингом лучших\n"
        "• Видео-разборы по темам и разделам\n\n"
        f"💎 Доступна в тарифах: {tier_lines}\n\n"
        "Загляни в «💎 Подписка» в главном меню, чтобы выбрать тариф 👇"
    )

def get_anatomy_announcement_keyboard():
    builder = InlineKeyboardBuilder()
    builder.button(text="💎 Подписка", callback_data="subscription_menu")
    builder.button(text="🦴 Анатомия", callback_data="anatomy_root")
    builder.adjust(1)
    return builder.as_markup()

def get_anatomy_exam_announcement_text() -> str:
    return (
        f"🎓 <b>Открыт раздел «Экзамен» по Анатомии!</b>\n{DIVIDER}\n\n"
        "Готовься к экзамену прицельно — три инструмента, и все бесплатно, без рефералов "
        "и подписки:\n\n"
        "✅ <b>ТЕСТ</b> — официальный банк вопросов кафедры нормальной анатомии ВМедА "
        "(Гайворонский и др.), 1040 вопросов, 10 частей\n"
        "📖 <b>Вопросы теории</b> — разбор экзаменационных вопросов с полными ответами\n"
        "🖐 <b>Вопросы практики</b> — «покажите и назовите» с фото атласа к каждому ответу\n\n"
        "🏆 Включи <b>рейтинговый режим</b> в ТЕСТе — результаты идут в общий рейтинг лучших. "
        "Самые активные и точные получат призы от нас!\n\n"
        "Жми «🎓 Экзамен» в разделе Анатомии и начинай готовиться 👇"
    )

def get_anatomy_exam_announcement_keyboard():
    builder = InlineKeyboardBuilder()
    builder.button(text="✅ ТЕСТ", callback_data="anatomy_exam_test_menu")
    builder.button(text="📖 Вопросы теории", callback_data="anatomy_exam_theory")
    builder.button(text="🖐 Вопросы практики", callback_data="anatomy_exam_practice")
    builder.button(text="🏆 Рейтинг ТЕСТа", callback_data="anatomy_exam_test_leaderboard")
    builder.adjust(1)
    return builder.as_markup()

def get_anatomy_latin_announcement_text() -> str:
    return (
        f"🏛 <b>Тест по латинским терминам — по всему курсу анатомии!</b>\n{DIVIDER}\n\n"
        "Проверь себя: общий тест собирает вопросы по латинской номенклатуре сразу со всего "
        "курса анатомии, а не по одной теме — и совершенно бесплатно, без рефералов и "
        "подписки.\n\n"
        "🏆 Лучшие результаты попадают в общий рейтинг — посмотри, на каком ты месте, "
        "и постарайся его улучшить.\n\n"
        "Жми кнопку ниже и проходи тест 👇"
    )

def get_anatomy_latin_announcement_keyboard():
    builder = InlineKeyboardBuilder()
    builder.button(text="🏛 Пройти тест", callback_data="anatomy_latin_all_start")
    builder.button(text="🏆 Рейтинг", callback_data="anatomy_latin_leaderboard")
    builder.adjust(1)
    return builder.as_markup()

def donor_display_name(uid_str: str) -> str:
    if stats.get("donor_hide_name", {}).get(uid_str):
        return "🙈 Аноним"
    username = stats["user_username"].get(uid_str)
    if username:
        return f"@{username}"
    return stats["user_names"].get(uid_str, f"Пользователь {uid_str}")

def get_donors_leaderboard_text() -> str:
    star_ranked = sorted(stats.get("donor_stars", {}).items(), key=lambda kv: kv[1], reverse=True)[:10]
    ruble_ranked = sorted(stats.get("donor_rubles", {}).items(), key=lambda kv: kv[1], reverse=True)[:10]

    lines = [f"🏆 <b>Лучшие донатеры</b>\n{DIVIDER}"]
    if star_ranked:
        lines.append("\n⭐ <b>По звёздам:</b>")
        for i, (uid, total) in enumerate(star_ranked):
            icon = RANK_MEDALS[i] if i < 3 else f"{i + 1}."
            lines.append(f"{icon} {donor_display_name(uid)} — <b>{total}</b> ⭐")
    if ruble_ranked:
        lines.append("\n💵 <b>По рублям:</b>")
        for i, (uid, total) in enumerate(ruble_ranked):
            icon = RANK_MEDALS[i] if i < 3 else f"{i + 1}."
            lines.append(f"{icon} {donor_display_name(uid)} — <b>{total}</b>₽")
    if not star_ranked and not ruble_ranked:
        lines.append("\nПока никто не пожертвовал — стань первым! 🙏")
    return "\n".join(lines)

def get_donors_leaderboard_keyboard():
    builder = InlineKeyboardBuilder()
    builder.button(text="🔄 Обновить", callback_data="donors_leaderboard")
    builder.button(text="🔙 Назад", callback_data="support_menu")
    builder.adjust(1)
    return builder.as_markup()

def get_visibility_choice_text(amount: int, unit: str) -> str:
    return (
        f"👀 <b>Показывать тебя в рейтинге?</b>\n{DIVIDER}\n\n"
        f"Сумма: <b>{amount}{unit}</b>\n\n"
        "Можно пожертвовать открыто — твой ник появится в «🏆 Лучшие донатеры» — "
        "или анонимно, тогда в рейтинге будет просто «Аноним»."
    )

def get_stars_visibility_keyboard(amount: int):
    builder = InlineKeyboardBuilder()
    builder.button(text="🙋 Показывать мой ник", callback_data=f"donate_stars_confirm:{amount}:pub")
    builder.button(text="🙈 Анонимно", callback_data=f"donate_stars_confirm:{amount}:anon")
    builder.adjust(1)
    builder.row(InlineKeyboardButton(text="🔙 Назад", callback_data="donate_stars_menu"))
    return builder.as_markup()

def get_rubles_visibility_keyboard(amount: int):
    builder = InlineKeyboardBuilder()
    builder.button(text="🙋 Показывать мой ник", callback_data=f"donate_rubles_confirm:{amount}:pub")
    builder.button(text="🙈 Анонимно", callback_data=f"donate_rubles_confirm:{amount}:anon")
    builder.adjust(1)
    builder.row(InlineKeyboardButton(text="🔙 Назад", callback_data="donate_rubles_menu"))
    return builder.as_markup()

def get_stars_menu_text() -> str:
    return (
        f"⭐ <b>Пожертвовать звёзды</b>\n{DIVIDER}\n\n"
        "Выбери количество звёзд Telegram, либо укажи своё:"
    )

def get_stars_menu_keyboard():
    builder = InlineKeyboardBuilder()
    for n in STARS_PRESETS:
        builder.button(text=f"⭐ {n}", callback_data=f"donate_stars_amount:{n}")
    builder.adjust(3)
    builder.row(InlineKeyboardButton(text="✏️ Своя сумма", callback_data="donate_stars_custom"))
    builder.row(InlineKeyboardButton(text="🔙 Назад", callback_data="support_menu"))
    return builder.as_markup()

def get_rubles_menu_text() -> str:
    return (
        f"💵 <b>Перевести рубли</b>\n{DIVIDER}\n\n"
        "Выбери сумму в рублях, либо укажи свою:"
    )

def get_rubles_menu_keyboard():
    builder = InlineKeyboardBuilder()
    for n in RUBLES_PRESETS:
        builder.button(text=f"{n}₽", callback_data=f"donate_rubles_amount:{n}")
    builder.adjust(3)
    builder.row(InlineKeyboardButton(text="✏️ Своя сумма", callback_data="donate_rubles_custom"))
    builder.row(InlineKeyboardButton(text="🔙 Назад", callback_data="support_menu"))
    return builder.as_markup()

def get_rubles_donation_message_text(amount: int) -> str:
    return (
        f"💵 <b>Перевод {amount}₽</b>\n{DIVIDER}\n\n"
        f'Нажми на кнопку ниже — откроется чат с <a href="{HELPER_ACCOUNT_URL}">@vmeda_helper</a>, '
        "сообщение с суммой уже будет готово. Останется его отправить — и тебе пришлют реквизиты для перевода.\n\n"
        "Спасибо огромное за поддержку! 🙏😇"
    )

def get_rubles_donation_keyboard(amount: int):
    template = (
        f"Привет! Хочу перевести {amount}₽ в поддержку бота VMEDA_examen_bot 🙏 "
        "Подскажи, пожалуйста, реквизиты для перевода."
    )
    url = f"{HELPER_ACCOUNT_URL}?text={urllib.parse.quote(template)}"
    builder = InlineKeyboardBuilder()
    builder.row(InlineKeyboardButton(text="💸 Пожертвовать рубли", url=url))
    builder.row(InlineKeyboardButton(text="🔙 Назад", callback_data="donate_rubles_menu"))
    return builder.as_markup()

async def send_stars_invoice(chat_id: int, stars: int) -> None:
    await bot.send_invoice(
        chat_id=chat_id,
        title="Поддержка автора бота",
        description=f"Спасибо за поддержку VMEDA_examen_bot! Пожертвование: {stars} ⭐",
        payload=f"donate_stars_{stars}_{chat_id}_{int(time.time())}",
        provider_token="",
        currency="XTR",
        prices=[LabeledPrice(label="Поддержка автора", amount=stars)],
    )

# ==================== СКРЫТЫЕ БИЛЕТЫ (40-50) ====================
HIDDEN_TICKET_RANGE = (40, 50)
FORCE_VISIBLE_TICKETS = {"40A"}  # исключения из скрытого диапазона — показывать всегда

def _ticket_number_part(ticket_num: str) -> int:
    """Достаёт числовую часть номера билета (например, '20A' -> 20)."""
    digits = ""
    for ch in str(ticket_num):
        if ch.isdigit():
            digits += ch
        else:
            break
    return int(digits) if digits else 0

def is_ticket_visible(ticket_num: str) -> bool:
    if str(ticket_num) in FORCE_VISIBLE_TICKETS:
        return True
    n = _ticket_number_part(ticket_num)
    return not (HIDDEN_TICKET_RANGE[0] <= n <= HIDDEN_TICKET_RANGE[1])

def _ticket_sort_key(ticket_num: str):
    digits = ""
    letters = ""
    for ch in str(ticket_num):
        if ch.isdigit():
            digits += ch
        else:
            letters += ch
    return (int(digits) if digits else 0, letters)

VISIBLE_TICKETS = [t for t in TICKETS if is_ticket_visible(str(t.get("num")))]
VISIBLE_TICKET_NUMS = sorted(
    [str(t.get("num")) for t in VISIBLE_TICKETS],
    key=_ticket_sort_key
)

def _normalize_ticket_num(s: str) -> str:
    """Убирает пробелы и приводит букву А к единому виду (кириллица/латиница, регистр)."""
    return (s or "").strip().upper().replace(" ", "").replace("А", "A")

TICKET_LOOKUP = {_normalize_ticket_num(k): v for k, v in TICKETS_DICT.items()}

# ==================== ПОИСК ПО КЛЮЧЕВЫМ СЛОВАМ ====================
SEARCH_RESULTS_LIMIT = 25

def _extract_words(text: str) -> list:
    return re.findall(r"[a-zа-яё]+", (text or "").lower().replace("ё", "е"))

def _word_stem(word: str) -> str:
    """Грубый стеммер: отбрасывает окончание, чтобы находить разные словоформы
    ("плазмодий" / "плазмодия" / "плазмодии")."""
    n = len(word)
    if n <= 4:
        return word
    if n <= 6:
        return word[:-1]
    return word[:-2]

def search_questions_by_keyword(query: str, limit: int = SEARCH_RESULTS_LIMIT):
    query_stems = [_word_stem(w) for w in _extract_words(query) if len(w) >= 3]
    if not query_stems:
        return []
    matches = []
    for num in sorted(QUESTIONS.keys(), key=lambda x: int(x)):
        # Игнорируем короткие служебные слова ("и", "с", "у"), иначе они ложно
        # совпадают с любым стеммом запроса через startswith.
        title_words = [w for w in _extract_words(QUESTIONS[num].get("title", "")) if len(w) >= 3]
        if all(any(tw.startswith(qs) for tw in title_words) for qs in query_stems):
            matches.append(num)
            if len(matches) >= limit:
                break
    return matches

def search_operative_surgery(query: str, limit: int = 15):
    """Простой регистронезависимый поиск по подстроке — стеммированный IDF-индекс как у
    search_questions_by_keyword тут overkill для четырёх плоских списков. Ищет по названию/тексту
    темы (все 61 темы полнотекстовые — см. operative_surgery.json v2), названию инструмента,
    названию структуры проекции, названию практической станции."""
    q = query.strip().lower()
    if not q:
        return [], [], [], []
    topics = [
        t for t in OPERATIVE_SURGERY["topics"]
        if q in t["title"].lower() or any(q in s["text"].lower() for s in t["subtopics"])
    ][:limit]
    instruments = [
        (group["group"], item["name"])
        for group in OPERATIVE_SURGERY["instrument_groups"]
        for item in group["items"]
        if q in item["name"].lower()
    ][:limit]
    projections = [
        (group["group"], item)
        for group in OPERATIVE_SURGERY["projections"]
        for item in group["items"]
        if q in item["structure"].lower()
    ][:limit]
    stations = [
        (group["group"], name)
        for group in OPERATIVE_SURGERY["practical_stations"]
        for name in group["items"]
        if q in name.lower()
    ][:limit]
    return topics, instruments, projections, stations

OH_SEARCH_PENDING: set = set()  # user_id, ждущих следующее текстовое сообщение как поисковый
# запрос по разделу «Оперативная хирургия» (см. cb_oh_search_prompt в handlers/operative_surgery.py
# и handle_oh_search_query ниже) — тот же паттерн, что ADMIN_PENDING/ASSISTANT_PENDING, только без
# многошагового действия, просто факт "этот юзер сейчас ищет".

PHYS_SEARCH_PENDING: set = set()  # тот же паттерн, что OH_SEARCH_PENDING выше, для раздела
# «Нормальная физиология» (см. cb_phys_search_prompt в handlers/physiology.py и
# handle_phys_search_query ниже).

# AI RAG-lite (индекс/поиск/подмес материалов ВМедА) перенесён в ai/rag.py — см. ai_rag.configure()
# (вызывается один раз при старте, см. конец файла) и ai_rag.search_snippets_multi()/format_context()
# (используются в обработчиках AI-режима).

async def is_subscribed(user_id: int) -> bool:
    try:
        member = await bot.get_chat_member(CHANNEL_ID, user_id)
        return member.status in [ChatMemberStatus.MEMBER, ChatMemberStatus.ADMINISTRATOR, ChatMemberStatus.CREATOR]
    except Exception:
        return False

# is_admin / is_assistant_admin / is_admin_or_assistant перенесены в services/access.py —
# см. импорт и реэкспорт в начале файла.

CAPTION_LIMIT = 1024

async def safe_edit_text(message, text, **kwargs) -> None:
    """Как edit_text, но если сообщение больше не текстовое (например, стало фото),
    удаляет его и отправляет новое вместо падения с ошибкой."""
    try:
        await message.edit_text(text, **kwargs)
    except TelegramBadRequest:
        await message.delete()
        await message.answer(text, **kwargs)

async def send_answer(target, body: str, short_caption: str, question: dict, keyboard, edit: bool) -> None:
    """Показывает текст вопроса+ответа. Если у вопроса есть картинка-схема, она всегда
    приходит первым сообщением:
    - при коротком ответе (уместился в лимит подписи Telegram) — единое сообщение "фото + текст";
    - при длинном ответе — сначала фото (с коротким заголовком), затем отдельным сообщением
      полный текст ответа (объединить в одно сообщение технически невозможно: Telegram
      ограничивает подпись к фото 1024 символами).
    target — CallbackQuery.message при edit=True, либо обычное Message при edit=False.
    При edit=True старое сообщение удаляется, а не редактируется — иначе оно осталось бы
    на прежнем месте в чате, выше нового фото."""
    image_name = question.get("image")
    image_path = os.path.join(IMAGES_DIR, image_name) if image_name else None
    if image_path and not os.path.exists(image_path):
        logger.warning("Изображение не найдено: %s", image_path)
        image_path = None

    if not image_path:
        if edit:
            await safe_edit_text(target, body, parse_mode="HTML", reply_markup=keyboard)
        else:
            await target.answer(body, parse_mode="HTML", reply_markup=keyboard)
        return

    photo = FSInputFile(image_path)
    if edit:
        await target.delete()

    if len(body) <= CAPTION_LIMIT:
        await target.answer_photo(photo, caption=body, parse_mode="HTML", reply_markup=keyboard)
        return

    caption = short_caption if len(short_caption) <= CAPTION_LIMIT else short_caption[:CAPTION_LIMIT - 1] + "…"
    try:
        await target.answer_photo(photo, caption=caption, parse_mode="HTML")
    except Exception:
        logger.exception("Не удалось отправить изображение %s", image_path)
    await target.answer(body, parse_mode="HTML", reply_markup=keyboard)

# ==================== ВЫГРУЗКА РАЗДЕЛОВ В WORD-ФАЙЛ ====================
_HTML_TOKEN_RE = re.compile(r"(<b>|</b>|<i>|</i>|<u>|</u>)")

def strip_html_tags(text: str) -> str:
    """Контент хранится с Telegram-HTML-разметкой (<b>, <i>, <u>) — вне докс-раннов она не нужна."""
    return _HTML_TOKEN_RE.sub("", text)

def add_html_run(paragraph, text: str) -> None:
    """Разбирает Telegram-HTML-подмножество (<b>, <i>, <u>) в форматированные докс-раны."""
    bold = italic = underline = False
    for token in _HTML_TOKEN_RE.split(text):
        if token == "<b>":
            bold = True
        elif token == "</b>":
            bold = False
        elif token == "<i>":
            italic = True
        elif token == "</i>":
            italic = False
        elif token == "<u>":
            underline = True
        elif token == "</u>":
            underline = False
        elif token:
            run = paragraph.add_run(token)
            run.bold = bold
            run.italic = italic
            run.underline = underline

def add_html_paragraphs(doc, text: str) -> None:
    """text может содержать \n\n (разрыв абзаца) и \n (перенос строки внутри абзаца)."""
    for block in text.split("\n\n"):
        p = doc.add_paragraph()
        for i, line in enumerate(block.split("\n")):
            if i > 0:
                p.add_run().add_break()
            add_html_run(p, line)

def add_labeled_field(doc, label: str, value: str) -> None:
    p = doc.add_paragraph()
    p.add_run(f"{label}:").bold = True
    add_html_paragraphs(doc, value)

def add_topic_block(doc, topic: dict) -> None:
    """Общий для physics_tasks.json и chemistry_tasks.json блок темы: заголовок, вступление,
    формулы/алгоритм и разобранные задачи."""
    doc.add_heading(topic["title"], level=2)
    if topic.get("intro"):
        add_html_paragraphs(doc, topic["intro"])
    add_html_paragraphs(doc, topic["formulas"])
    for task in topic.get("tasks", []):
        doc.add_heading(f"{task['num']}. {task['title']}", level=3)
        add_labeled_field(doc, "Условие", task["condition"])
        add_labeled_field(doc, "Решение", task["solution"])

def docx_filename(title: str) -> str:
    return re.sub(r"[^\w\-]+", "_", title, flags=re.UNICODE).strip("_")[:60] + ".docx"

def build_docx_file(title: str, fill_fn) -> BufferedInputFile:
    doc = DocxDocument()
    doc.add_heading(title, level=0)
    fill_fn(doc)
    buf = io.BytesIO()
    doc.save(buf)
    return BufferedInputFile(buf.getvalue(), filename=docx_filename(title))

def build_biology_tickets_file() -> BufferedInputFile:
    def fill(doc):
        for ticket in TICKETS:
            doc.add_heading(ticket["title"], level=1)
            for q in ticket["questions"]:
                doc.add_heading(f"{q['num']}. {q['title']}", level=2)
                add_html_paragraphs(doc, q["answer"])
    return build_docx_file("Биология — все билеты (вопросы и ответы)", fill)

def build_physics_full_file() -> BufferedInputFile:
    def fill(doc):
        doc.add_heading("Часть 1. Тестовая часть — 186 вопросов (по алфавиту)", level=1)
        items = sorted(PHYSICS_QUESTIONS.values(), key=lambda v: v["title"])
        for item in items:
            doc.add_heading(item["title"], level=2)
            add_html_paragraphs(doc, item["answer"])
        doc.add_page_break()
        doc.add_heading("Часть 2. Шаблоны решения задач по всем темам", level=1)
        for topic in PHYSICS_TASKS.values():
            add_topic_block(doc, topic)
    return build_docx_file("Физика — тестовая часть и шаблоны решения задач", fill)

def build_physics_grade45_file() -> BufferedInputFile:
    def fill(doc):
        items = sorted(PHYSICS_GRADE45_QUESTIONS.values(), key=lambda v: v["title"])
        for i, item in enumerate(items):
            heading = doc.add_heading(item["title"], level=1)
            heading.paragraph_format.space_before = Pt(0 if i == 0 else 30)
            heading.paragraph_format.space_after = Pt(12)
            add_html_paragraphs(doc, item["answer"])
    return build_docx_file("Физика — (60 вопросов) на 4/5", fill)

_ALGORITHM_BLOCK_RE = re.compile(
    r"\n\n<b>Алгоритм решения:</b>.*?(?=\n\n<b>Единицы измерения:</b>)", re.DOTALL
)

def build_physics_tasks_cheatsheet_file() -> BufferedInputFile:
    """Краткая шпаргалка по всем типам задач по физике — только формулы и обозначения
    (пошаговый алгоритм решения из physics_tasks.json намеренно опущен для краткости)."""
    def fill(doc):
        for num in sorted(PHYSICS_TASKS.keys(), key=int):
            topic = PHYSICS_TASKS[num]
            heading = doc.add_heading(topic["title"], level=1)
            heading.paragraph_format.space_before = Pt(0 if num == "1" else 24)
            heading.paragraph_format.space_after = Pt(6)
            formulas_only = _ALGORITHM_BLOCK_RE.sub("", topic["formulas"])
            add_html_paragraphs(doc, formulas_only)
    return build_docx_file("Физика — шпаргалка по формулам к задачам", fill)

def build_physics_ticket_tasks_file() -> BufferedInputFile:
    def fill(doc):
        for num in sorted(PHYSICS_TEST_TICKETS.keys(), key=int):
            ticket = PHYSICS_TEST_TICKETS[num]
            tasks = ticket.get("tasks")
            if not tasks:
                continue
            doc.add_heading(f"{ticket['title']} — Часть 2. Задачи", level=1)
            for task in tasks:
                doc.add_heading(f"Задача {task['num']}. {task['title']}", level=2)
                add_labeled_field(doc, "Условие", task["condition"])
                add_labeled_field(doc, "Решение", task["solution"])
    return build_docx_file("Физика — билеты с задачами (ответы)", fill)

_LAB_FIELD_LABELS = {
    "positive_sol": "Положительный золь", "negative_sol": "Отрицательный золь",
    "procedure": "Методика", "theory": "Теория", "titrant": "Титрант",
    "indicator": "Индикатор", "buffer": "Буфер", "equations": "Уравнения реакций",
    "calculations": "Расчёты",
}

def build_chemistry_labs_file() -> BufferedInputFile:
    def fill(doc):
        for lab in CHEMISTRY_LABS["labs"]:
            doc.add_heading(f"Лабораторная работа {lab['number']}. {lab['theme']}", level=1)
            add_labeled_field(doc, "Условие", lab["condition"])
            for exp in lab.get("experiments", []):
                doc.add_heading(exp.get("name", ""), level=2)
                for key in ("description", "mechanism", "technique", "sorbent", "eluent", "procedure"):
                    value = exp.get(key)
                    if value:
                        add_html_paragraphs(doc, value)
            for key, label in _LAB_FIELD_LABELS.items():
                value = lab.get(key)
                if value:
                    add_labeled_field(doc, label, value)
    return build_docx_file("Химия — все лабораторные работы", fill)

def build_chemistry_tasks_file() -> BufferedInputFile:
    def fill(doc):
        for topic in CHEMISTRY_TASKS.values():
            add_topic_block(doc, topic)
    return build_docx_file("Химия — все задачи", fill)

# ==================== КЛАВИАТУРЫ ====================
def _anatomy_menu_label(user_id: int = None) -> str:
    sub_anatomy = user_id is not None and has_subscription_anatomy_access(user_id)
    if user_id is not None and is_admin(user_id):
        return "🔥🦴 Анатомия (админ)"
    elif anatomy_handlers.ANATOMY_MAINTENANCE_MODE:
        return "🦴 Анатомия (техобслуживание)"
    elif sub_anatomy:
        return "🔥🦴 Анатомия 💎"
    else:
        return "🔥🦴 Анатомия"


def _histology_menu_label(user_id: int = None) -> str:
    sub_histology = user_id is not None and has_subscription_histology_access(user_id)
    if HISTOLOGY_PUBLIC:
        return "🔬 Гистология"
    elif user_id is not None and is_admin(user_id):
        return "🔬 Гистология (админ)"
    elif is_section_promo_active("histology"):
        return "🔬 Гистология 🎉"
    elif user_id is not None and get_referral_count_this_month(user_id) >= REFERRAL_FULL_ACCESS_THRESHOLD:
        return "🔬 Гистология"
    elif sub_histology:
        return "🔬 Гистология 💎"
    elif user_id is not None and has_histology_temp_access(user_id):
        return "🔬 Гистология (пробный период)"
    else:
        return "🔬 Гистология (рефералы/подписка)"


def get_main_menu(user_id: int = None):
    builder = InlineKeyboardBuilder()
    builder.button(text="🤖 VMedA AI (бета)", callback_data="ai_menu")
    builder.button(text="1️⃣ Первый курс", callback_data="course_menu:1")
    builder.button(text="2️⃣ Второй курс", callback_data="course_menu:2")
    builder.button(text="👥 Пригласить друзей", callback_data="referral_info")
    builder.button(text="🏆 Рейтинг", callback_data="referral_leaderboard")
    rollcall_confirmed_count = len(stats["rollcall_confirmed"])
    builder.button(
        text=f"📋 Перекличка ({rollcall_confirmed_count}/{ROLLCALL_GROUP_COUNT})",
        callback_data="rollcall_menu"
    )
    battle_label = "⚔️ Битва рефералов 🔥" if is_battle_active() else "⚔️ Битва рефералов"
    builder.button(text=battle_label, callback_data="referral_battle")
    if user_id is not None and has_active_subscription(user_id):
        sub_label = "💎 Моя подписка"
    else:
        sub_label = "💎 Подписка без рефералов"
    builder.button(text=sub_label, callback_data="subscription_menu")
    builder.button(text="😇 Поддержать автора 💰", callback_data="support_menu")
    builder.adjust(1)
    return builder.as_markup()


# «1️⃣ Первый курс» / «2️⃣ Второй курс» — группировка предметов по году обучения на главном экране
# (см. запрос пользователя). Анатомия и Гистология входят в оба курса — их динамические
# (зависящие от подписки/рефералов/техобслуживания) подписи вынесены в _anatomy_menu_label()/
# _histology_menu_label() выше, чтобы не дублировать логику между этой клавиатурой и (в будущем)
# любым другим местом, где эти кнопки понадобятся. Сами callback_data кнопок (menu_biology,
# anatomy_root, histology_menu, oh:menu, phys:menu, ...) не меняются — гейтинг/маршрутизация
# работают ровно как раньше, меняется только то, с какого экрана до них можно добраться.
COURSE_SUBJECTS = {
    1: [
        ("⚛️ Физика", "menu_physics"),
        ("🧪 Химия", "menu_chemistry"),
        ("🧬 Биология", "menu_biology"),
        (None, "anatomy_root"),      # label resolved dynamically, see below
        (None, "histology_menu"),
    ],
    2: [
        (None, "anatomy_root"),
        (None, "histology_menu"),
        ("🧠 Нормальная физиология", "phys:menu"),
        ("🔪 Оперативная хирургия", "oh:menu"),
    ],
}


def get_course_menu_text(course: int) -> str:
    title = "1️⃣ ПЕРВЫЙ КУРС" if course == 1 else "2️⃣ ВТОРОЙ КУРС"
    return f"{title}\n{DIVIDER}\n\nВыбери предмет:"


def get_course_menu_keyboard(course: int, user_id: int = None):
    builder = InlineKeyboardBuilder()
    for label, callback_data in COURSE_SUBJECTS[course]:
        if label is None:
            label = _anatomy_menu_label(user_id) if callback_data == "anatomy_root" else _histology_menu_label(user_id)
        builder.button(text=label, callback_data=callback_data)
    builder.adjust(1)
    builder.row(InlineKeyboardButton(text="🔙 Назад в меню", callback_data="back_to_main"))
    return builder.as_markup()

def get_referral_back_keyboard():
    builder = InlineKeyboardBuilder()
    builder.row(InlineKeyboardButton(text="💎 Открыть доступ без рефералов", callback_data="subscription_menu"))
    builder.row(InlineKeyboardButton(text="🚀 Запустить Medical_vpn_bot", url=MEDICAL_VPN_URL))
    builder.row(InlineKeyboardButton(text="🔙 Назад в меню", callback_data="back_to_main"))
    return builder.as_markup()

def get_referral_full_access_keyboard(user_id: int = None):
    builder = InlineKeyboardBuilder()
    builder.button(text="⚔️ Битва рефералов", callback_data="referral_battle")
    sub_label = "💎 Моя подписка" if user_id is not None and has_active_subscription(user_id) else "💎 Подписка"
    builder.button(text=sub_label, callback_data="subscription_menu")
    builder.button(text="🔙 Назад в меню", callback_data="back_to_main")
    builder.adjust(1)
    return builder.as_markup()

# Главное меню Биологии, режим опроса (флэш-карточки) и клавиатуры билетов/списка вопросов
# перенесены в handlers/biology.py — см. dp.include_router(biology_handlers.router) и реэкспорт
# дальше по файлу.

def get_search_results_keyboard(nums: list):
    builder = InlineKeyboardBuilder()
    for num in nums:
        title = QUESTIONS[num].get("title", "")
        short_title = title if len(title) <= 60 else title[:57] + "…"
        builder.button(text=f"{num}. {short_title}", callback_data=f"q:{num}")
    builder.adjust(1)
    builder.row(InlineKeyboardButton(text="🔙 Назад", callback_data="menu_questions"))
    return builder.as_markup()

def get_question_answer_keyboard(q_num: str):
    builder = InlineKeyboardBuilder()
    n = int(q_num)
    nav = []
    if str(n - 1) in QUESTIONS:
        nav.append(InlineKeyboardButton(text="⬅️ Предыдущий вопрос", callback_data=f"q:{n - 1}"))
    if str(n + 1) in QUESTIONS:
        nav.append(InlineKeyboardButton(text="Следующий вопрос ➡️", callback_data=f"q:{n + 1}"))
    if nav:
        builder.row(*nav)
    builder.row(InlineKeyboardButton(text="🎲 Случайный вопрос", callback_data="question_random"))
    builder.row(InlineKeyboardButton(text="🔢 Ввести номер вручную", callback_data="question_by_number"))
    page = (n - 1) // 50 + 1
    builder.row(InlineKeyboardButton(text="🔙 Назад", callback_data=f"qpage:{page}"))
    return builder.as_markup()

def get_question_page_keyboard(page: int):
    builder = InlineKeyboardBuilder()
    start = (page - 1) * 50 + 1
    end = min(page * 50, 185)
    for i in range(start, end + 1, 5):
        row = [InlineKeyboardButton(text=f"🟢 {num}", callback_data=f"q:{num}") for num in range(i, min(i + 5, end + 1))]
        builder.row(*row)
    nav = []
    if page > 1:
        nav.append(InlineKeyboardButton(text="⬅️ Назад", callback_data=f"qpage:{page-1}"))
    if page < 4:
        nav.append(InlineKeyboardButton(text="Вперёд ➡️", callback_data=f"qpage:{page+1}"))
    if nav:
        builder.row(*nav)
    builder.row(InlineKeyboardButton(text="🔙 К списку страниц", callback_data="menu_questions"))
    return builder.as_markup()

def get_ticket_questions_keyboard(ticket_num: str):
    builder = InlineKeyboardBuilder()
    ticket = TICKETS_DICT.get(ticket_num, {})
    questions = ticket.get("questions", [])
    for q in questions:
        q_num = q.get("num")
        builder.button(text=f"🟢 Вопрос {q_num}", callback_data=f"ticket_q:{ticket_num}:{q_num}")
    builder.adjust(1)
    builder.row(InlineKeyboardButton(text="🔙 Назад к билетам", callback_data="menu_tickets"))
    return builder.as_markup()

# ==================== ФИЗИКА (клавиатуры) ====================
def get_physics_menu():
    builder = InlineKeyboardBuilder()
    builder.button(text="📝 Тестовая часть (186 вопросов)", callback_data="physics_test")
    builder.button(text="📘 Билеты", callback_data="physics_tickets")
    builder.button(text="🧮 Задачи", callback_data="physics_tasks")
    builder.button(text="❓ (60 вопросов) на 4/5", callback_data="physics_grade45")
    builder.button(text="⭐ Доп. вопросы от преподавателей", callback_data="physics_extra")
    builder.button(text="📄 186 вопросов + шаблоны задач (файл)", callback_data="download_physics_full")
    builder.button(text="📄 (60 вопросов) на 4/5 (файл)", callback_data="download_physics_grade45")
    builder.button(text="📄 Ответы на задачи билетов (файл)", callback_data="download_physics_ticket_tasks")
    builder.button(text="📄 Шпаргалка по формулам к задачам (файл)", callback_data="download_physics_tasks_cheatsheet")
    builder.adjust(1)
    builder.row(InlineKeyboardButton(text="🔙 Назад в меню", callback_data="back_to_main"))
    return builder.as_markup()

def get_physics_tickets_menu():
    builder = InlineKeyboardBuilder()
    builder.button(text="📝 Тестовые билеты", callback_data="physics_test_tickets")
    builder.button(text="📖 Билеты теоретической части", callback_data="physics_theory_tickets")
    builder.button(text="🧮 Билеты с задачами", callback_data="physics_task_tickets")
    builder.adjust(1)
    builder.row(InlineKeyboardButton(text="🔙 Назад", callback_data="menu_physics"))
    return builder.as_markup()

def get_physics_task_tickets_keyboard():
    builder = InlineKeyboardBuilder()
    for num in sorted(PHYSICS_TASK_TICKETS.keys(), key=int):
        builder.button(text=f"🧮 {PHYSICS_TASK_TICKETS[num]['title']}", callback_data=f"phys_task_ticket:{num}")
    builder.adjust(2)
    builder.row(InlineKeyboardButton(text="🔙 Назад", callback_data="physics_tickets"))
    return builder.as_markup()

def get_physics_task_ticket_list_keyboard(num: str):
    builder = InlineKeyboardBuilder()
    for task in PHYSICS_TASK_TICKETS[num]["tasks"]:
        builder.button(text=f"📝 Задача {task['num']}", callback_data=f"phys_task_ticket_show:{num}:{task['num']}")
    builder.adjust(1)
    builder.row(InlineKeyboardButton(text="🔙 К списку билетов", callback_data="physics_task_tickets"))
    return builder.as_markup()

def get_physics_task_ticket_detail_keyboard(num: str, task_num: int):
    builder = InlineKeyboardBuilder()
    tasks = PHYSICS_TASK_TICKETS[num]["tasks"]
    nums = [t["num"] for t in tasks]
    idx = nums.index(task_num)
    nav = []
    if idx > 0:
        nav.append(InlineKeyboardButton(text="⬅️ Предыдущая", callback_data=f"phys_task_ticket_show:{num}:{nums[idx-1]}"))
    if idx < len(nums) - 1:
        nav.append(InlineKeyboardButton(text="Следующая ➡️", callback_data=f"phys_task_ticket_show:{num}:{nums[idx+1]}"))
    if nav:
        builder.row(*nav)
    builder.row(InlineKeyboardButton(text="🔙 К списку задач", callback_data=f"phys_task_ticket:{num}"))
    return builder.as_markup()

def get_physics_test_tickets_keyboard():
    builder = InlineKeyboardBuilder()
    for num in sorted(PHYSICS_TEST_TICKETS.keys(), key=int):
        builder.button(text=f"📄 {PHYSICS_TEST_TICKETS[num]['title']}", callback_data=f"phys_test_ticket:{num}")
    builder.adjust(2)
    builder.row(InlineKeyboardButton(text="🔙 Назад", callback_data="physics_tickets"))
    return builder.as_markup()

def get_physics_test_ticket_detail_keyboard(num: str):
    builder = InlineKeyboardBuilder()
    ticket = PHYSICS_TEST_TICKETS.get(num, {})
    if ticket.get("tasks"):
        builder.row(InlineKeyboardButton(text="🧮 Часть 2. Задачи", callback_data=f"phys_test_ticket_tasks:{num}"))
    builder.row(InlineKeyboardButton(text="🔙 К списку билетов", callback_data="physics_test_tickets"))
    return builder.as_markup()

def get_physics_test_ticket_task_list_keyboard(num: str):
    builder = InlineKeyboardBuilder()
    for task in PHYSICS_TEST_TICKETS[num]["tasks"]:
        builder.button(text=f"📝 Задача {task['num']}", callback_data=f"phys_test_ticket_task_show:{num}:{task['num']}")
    builder.adjust(1)
    builder.row(InlineKeyboardButton(text="🔙 К билету", callback_data=f"phys_test_ticket:{num}"))
    return builder.as_markup()

def get_physics_test_ticket_task_detail_keyboard(num: str, task_num: int):
    builder = InlineKeyboardBuilder()
    tasks = PHYSICS_TEST_TICKETS[num]["tasks"]
    nums = [t["num"] for t in tasks]
    idx = nums.index(task_num)
    nav = []
    if idx > 0:
        nav.append(InlineKeyboardButton(text="⬅️ Предыдущая", callback_data=f"phys_test_ticket_task_show:{num}:{nums[idx-1]}"))
    if idx < len(nums) - 1:
        nav.append(InlineKeyboardButton(text="Следующая ➡️", callback_data=f"phys_test_ticket_task_show:{num}:{nums[idx+1]}"))
    if nav:
        builder.row(*nav)
    builder.row(InlineKeyboardButton(text="🔙 К списку задач", callback_data=f"phys_test_ticket_tasks:{num}"))
    return builder.as_markup()

def get_physics_theory_tickets_keyboard():
    builder = InlineKeyboardBuilder()
    for num in sorted(PHYSICS_THEORY_TICKETS.keys(), key=int):
        builder.button(text=f"📖 {PHYSICS_THEORY_TICKETS[num]['title']}", callback_data=f"phys_theory_ticket:{num}")
    builder.adjust(2)
    builder.row(InlineKeyboardButton(text="🔙 Назад", callback_data="physics_tickets"))
    return builder.as_markup()

def get_physics_theory_ticket_detail_keyboard(num: str):
    builder = InlineKeyboardBuilder()
    ticket = PHYSICS_THEORY_TICKETS[num]
    for i, q in enumerate(ticket["questions"]):
        label = q["title"] if len(q["title"]) <= 60 else q["title"][:57] + "..."
        builder.button(text=f"{i + 1}. {label}", callback_data=f"phys_theory_q:{num}:{i}")
    builder.adjust(1)
    builder.row(InlineKeyboardButton(text="🔙 К списку билетов", callback_data="physics_theory_tickets"))
    return builder.as_markup()

def get_physics_theory_question_keyboard(num: str, idx: int):
    builder = InlineKeyboardBuilder()
    total = len(PHYSICS_THEORY_TICKETS[num]["questions"])
    nav = []
    if idx > 0:
        nav.append(InlineKeyboardButton(text="⬅️ Предыдущий вопрос", callback_data=f"phys_theory_q:{num}:{idx - 1}"))
    if idx < total - 1:
        nav.append(InlineKeyboardButton(text="Следующий вопрос ➡️", callback_data=f"phys_theory_q:{num}:{idx + 1}"))
    if nav:
        builder.row(*nav)
    builder.row(InlineKeyboardButton(text="🔙 К билету", callback_data=f"phys_theory_ticket:{num}"))
    return builder.as_markup()

def get_physics_tasks_topics_keyboard():
    builder = InlineKeyboardBuilder()
    for num, topic in sorted(PHYSICS_TASKS.items(), key=lambda x: int(x[0])):
        builder.button(text=f"📂 {topic['title']}", callback_data=f"phystask_topic:{num}")
    builder.adjust(1)
    builder.row(InlineKeyboardButton(text="🔙 Назад", callback_data="menu_physics"))
    return builder.as_markup()

def get_physics_task_topic_keyboard(topic_num: str):
    builder = InlineKeyboardBuilder()
    builder.button(text="📐 Формулы и алгоритм", callback_data=f"phystask_formulas:{topic_num}")
    builder.button(text="📋 Список задач", callback_data=f"phystask_list:{topic_num}")
    builder.adjust(1)
    builder.row(InlineKeyboardButton(text="🔙 К темам", callback_data="physics_tasks"))
    return builder.as_markup()

def get_physics_formulas_keyboard(topic_num: str):
    builder = InlineKeyboardBuilder()
    builder.row(InlineKeyboardButton(text="🔙 Назад", callback_data=f"phystask_topic:{topic_num}"))
    return builder.as_markup()

def get_physics_task_list_keyboard(topic_num: str):
    builder = InlineKeyboardBuilder()
    topic = PHYSICS_TASKS[topic_num]
    for task in topic["tasks"]:
        builder.button(text=f"📝 Задача {task['num']}", callback_data=f"phystask_show:{topic_num}:{task['num']}")
    builder.adjust(1)
    builder.row(InlineKeyboardButton(text="🔙 Назад", callback_data=f"phystask_topic:{topic_num}"))
    return builder.as_markup()

def get_physics_task_detail_keyboard(topic_num: str, task_num: int):
    builder = InlineKeyboardBuilder()
    tasks = PHYSICS_TASKS[topic_num]["tasks"]
    nums = [t["num"] for t in tasks]
    idx = nums.index(task_num)
    nav = []
    if idx > 0:
        nav.append(InlineKeyboardButton(text="⬅️ Предыдущая", callback_data=f"phystask_show:{topic_num}:{nums[idx-1]}"))
    if idx < len(nums) - 1:
        nav.append(InlineKeyboardButton(text="Следующая ➡️", callback_data=f"phystask_show:{topic_num}:{nums[idx+1]}"))
    if nav:
        builder.row(*nav)
    builder.row(InlineKeyboardButton(text="🔙 К списку задач", callback_data=f"phystask_list:{topic_num}"))
    return builder.as_markup()

def get_physics_test_pages():
    builder = InlineKeyboardBuilder()
    builder.button(text="📄 Страница 1 (1-50)", callback_data="physics_page:1")
    builder.button(text="📄 Страница 2 (51-100)", callback_data="physics_page:2")
    builder.button(text="📄 Страница 3 (101-150)", callback_data="physics_page:3")
    builder.button(text="📄 Страница 4 (151-186)", callback_data="physics_page:4")
    builder.adjust(1)
    builder.row(InlineKeyboardButton(text="🔙 Назад", callback_data="menu_physics"))
    return builder.as_markup()

def get_physics_question_keyboard(page: int):
    builder = InlineKeyboardBuilder()
    start = (page - 1) * 50 + 1
    end = min(page * 50, 186)
    for i in range(start, end + 1, 5):
        row = [InlineKeyboardButton(text=f"🟢 {num}", callback_data=f"physics_q:{num}") for num in range(i, min(i + 5, end + 1))]
        builder.row(*row)
    nav = []
    if page > 1:
        nav.append(InlineKeyboardButton(text="⬅️ Назад", callback_data=f"physics_page:{page-1}"))
    if page < 4:
        nav.append(InlineKeyboardButton(text="Вперёд ➡️", callback_data=f"physics_page:{page+1}"))
    if nav:
        builder.row(*nav)
    builder.row(InlineKeyboardButton(text="🔙 К страницам", callback_data="physics_test"))
    return builder.as_markup()

def get_physics_answer_keyboard(q_num: str):
    builder = InlineKeyboardBuilder()
    n = int(q_num)
    nav = []
    if str(n - 1) in PHYSICS_QUESTIONS:
        nav.append(InlineKeyboardButton(text="⬅️ Предыдущий вопрос", callback_data=f"physics_q:{n - 1}"))
    if str(n + 1) in PHYSICS_QUESTIONS:
        nav.append(InlineKeyboardButton(text="Следующий вопрос ➡️", callback_data=f"physics_q:{n + 1}"))
    if nav:
        builder.row(*nav)
    page = (n - 1) // 50 + 1
    builder.row(InlineKeyboardButton(text="🔙 Назад", callback_data=f"physics_page:{page}"))
    return builder.as_markup()

def get_physics_grade45_keyboard():
    builder = InlineKeyboardBuilder()
    nums = sorted(PHYSICS_GRADE45_QUESTIONS.keys(), key=int)
    for i in range(0, len(nums), 5):
        row = [InlineKeyboardButton(text=f"🟢 {n}", callback_data=f"physics45_q:{n}") for n in nums[i:i + 5]]
        builder.row(*row)
    builder.row(InlineKeyboardButton(text="🔙 Назад", callback_data="menu_physics"))
    return builder.as_markup()

def get_physics_grade45_answer_keyboard(q_num: str):
    builder = InlineKeyboardBuilder()
    nums = sorted(PHYSICS_GRADE45_QUESTIONS.keys(), key=int)
    idx = nums.index(q_num)
    nav = []
    if idx > 0:
        nav.append(InlineKeyboardButton(text="⬅️ Предыдущий вопрос", callback_data=f"physics45_q:{nums[idx - 1]}"))
    if idx < len(nums) - 1:
        nav.append(InlineKeyboardButton(text="Следующий вопрос ➡️", callback_data=f"physics45_q:{nums[idx + 1]}"))
    if nav:
        builder.row(*nav)
    builder.row(InlineKeyboardButton(text="🔙 К списку вопросов", callback_data="physics_grade45"))
    return builder.as_markup()

def get_physics_extra_keyboard():
    builder = InlineKeyboardBuilder()
    for num in sorted(PHYSICS_EXTRA_QUESTIONS.keys(), key=int):
        builder.button(text=f"⭐ {PHYSICS_EXTRA_QUESTIONS[num]['title']}", callback_data=f"physics_extra_q:{num}")
    builder.adjust(1)
    builder.row(InlineKeyboardButton(text="🔙 Назад", callback_data="menu_physics"))
    return builder.as_markup()

def get_physics_extra_answer_keyboard(q_num: str):
    builder = InlineKeyboardBuilder()
    nums = sorted(PHYSICS_EXTRA_QUESTIONS.keys(), key=int)
    idx = nums.index(q_num)
    nav = []
    if idx > 0:
        nav.append(InlineKeyboardButton(text="⬅️ Предыдущий вопрос", callback_data=f"physics_extra_q:{nums[idx - 1]}"))
    if idx < len(nums) - 1:
        nav.append(InlineKeyboardButton(text="Следующий вопрос ➡️", callback_data=f"physics_extra_q:{nums[idx + 1]}"))
    if nav:
        builder.row(*nav)
    builder.row(InlineKeyboardButton(text="🔙 К списку вопросов", callback_data="physics_extra"))
    return builder.as_markup()

# ==================== ХИМИЯ (клавиатуры) ====================
def get_chemistry_menu():
    builder = InlineKeyboardBuilder()
    builder.button(text="📚 Теория", callback_data="chemistry_theory")
    builder.button(text="📝 Задачи", callback_data="chemistry_tasks")
    builder.button(text="🧪 Лабораторные работы", callback_data="chemistry_labs")
    builder.button(text="🎫 Билеты", callback_data="chemistry_tickets")
    builder.button(text="📄 Все лабораторные работы (файл)", callback_data="download_chemistry_labs")
    builder.button(text="📄 Все задачи (файл)", callback_data="download_chemistry_tasks")
    builder.adjust(1)
    builder.row(InlineKeyboardButton(text="🔙 Назад в меню", callback_data="back_to_main"))
    return builder.as_markup()

# chemistry_tickets_access_ok перенесена в services/access.py — см. импорт и реэкспорт
# в начале файла, рядом с исходным местом секции "РЕФЕРАЛЬНАЯ СИСТЕМА".

def get_chemistry_tickets_locked_text() -> str:
    cheapest = cheapest_gated3_tier()
    return (
        f"🎫 <b>Билеты по химии</b>\n{DIVIDER}\n\n"
        f"Раздел закрыт дополнительным условием: нужно {REFERRAL_FULL_ACCESS_THRESHOLD} "
        f"реферала в этом месяце или подписка от 89₽ (например, «{cheapest['emoji']} {cheapest['title']}» за "
        f"{cheapest['price_rub']}₽ / {cheapest['price_stars']}⭐) — обычного доступа к Химии для "
        "билетов недостаточно.\n\n"
        "Пригласи друзей или оформи подписку, чтобы открыть раздел."
    )

def get_chemistry_tickets_locked_keyboard():
    builder = InlineKeyboardBuilder()
    builder.button(text="👥 Пригласить друзей", callback_data="referral_info")
    builder.button(text="💎 Оформить подписку", callback_data="subscription_menu")
    builder.button(text="🔙 Назад", callback_data="menu_chemistry")
    builder.adjust(1)
    return builder.as_markup()

def get_chemistry_tickets_menu():
    builder = InlineKeyboardBuilder()
    builder.button(text="📖 Билеты теории", callback_data="chem_theory_tickets")
    builder.button(text="🧮 Билеты практики", callback_data="chem_practice_tickets")
    builder.adjust(1)
    builder.row(InlineKeyboardButton(text="🔙 Назад", callback_data="menu_chemistry"))
    return builder.as_markup()

def get_chemistry_theory_tickets_keyboard():
    builder = InlineKeyboardBuilder()
    for num in sorted(CHEMISTRY_THEORY_TICKETS.keys(), key=int):
        builder.button(text=f"📖 {CHEMISTRY_THEORY_TICKETS[num]['title']}", callback_data=f"chem_theory_ticket:{num}")
    builder.adjust(2)
    builder.row(InlineKeyboardButton(text="🔙 Назад", callback_data="chemistry_tickets"))
    return builder.as_markup()

def get_chemistry_theory_ticket_detail_keyboard(num: str):
    builder = InlineKeyboardBuilder()
    ticket = CHEMISTRY_THEORY_TICKETS[num]
    for i, q in enumerate(ticket["questions"]):
        label = q["title"] if len(q["title"]) <= 60 else q["title"][:57] + "..."
        builder.button(text=f"{i + 1}. {label}", callback_data=f"chem_theory_q:{num}:{i}")
    builder.adjust(1)
    builder.row(InlineKeyboardButton(text="🔙 К списку билетов", callback_data="chem_theory_tickets"))
    return builder.as_markup()

def get_chemistry_theory_question_keyboard(num: str, idx: int):
    builder = InlineKeyboardBuilder()
    total = len(CHEMISTRY_THEORY_TICKETS[num]["questions"])
    nav = []
    if idx > 0:
        nav.append(InlineKeyboardButton(text="⬅️ Предыдущий вопрос", callback_data=f"chem_theory_q:{num}:{idx - 1}"))
    if idx < total - 1:
        nav.append(InlineKeyboardButton(text="Следующий вопрос ➡️", callback_data=f"chem_theory_q:{num}:{idx + 1}"))
    if nav:
        builder.row(*nav)
    builder.row(InlineKeyboardButton(text="🔙 К списку билетов", callback_data="chem_theory_tickets"))
    return builder.as_markup()

def get_chemistry_practice_tickets_keyboard():
    builder = InlineKeyboardBuilder()
    for num in sorted(CHEMISTRY_PRACTICE_TICKETS.keys(), key=int):
        builder.button(text=f"🧮 {CHEMISTRY_PRACTICE_TICKETS[num]['title']}", callback_data=f"chem_practice_ticket:{num}")
    builder.adjust(2)
    builder.row(InlineKeyboardButton(text="🔙 Назад", callback_data="chemistry_tickets"))
    return builder.as_markup()

def get_chemistry_practice_ticket_keyboard():
    builder = InlineKeyboardBuilder()
    builder.row(InlineKeyboardButton(text="🔙 К списку билетов", callback_data="chem_practice_tickets"))
    return builder.as_markup()

def get_chemistry_theory_list():
    builder = InlineKeyboardBuilder()
    for num in range(1, 17):
        builder.button(text=f"📖 Тема {num}", callback_data=f"chem_theory:{num}")
    builder.adjust(2)
    builder.row(InlineKeyboardButton(text="🔙 Назад", callback_data="menu_chemistry"))
    return builder.as_markup()

def get_theory_navigation(current_num: int):
    builder = InlineKeyboardBuilder()
    if current_num > 1:
        builder.button(text="⬅️ Предыдущая", callback_data=f"chem_theory:{current_num-1}")
    if current_num < 16:
        builder.button(text="Следующая ➡️", callback_data=f"chem_theory:{current_num+1}")
    builder.adjust(2)
    builder.row(InlineKeyboardButton(text="🔙 К списку тем", callback_data="chemistry_theory_list"))
    return builder.as_markup()

def get_labs_keyboard():
    builder = InlineKeyboardBuilder()
    for lab in CHEMISTRY_LABS["labs"]:
        builder.button(text=f"🧪 Лаба {lab['number']}", callback_data=f"lab:{lab['number']}")
    builder.adjust(2)
    builder.row(InlineKeyboardButton(text="🔙 Назад", callback_data="menu_chemistry"))
    return builder.as_markup()

def get_chemistry_tasks_topics_keyboard():
    builder = InlineKeyboardBuilder()
    for num, topic in sorted(CHEMISTRY_TASKS.items(), key=lambda x: int(x[0])):
        builder.button(text=f"📂 {topic['title']}", callback_data=f"chemtask_topic:{num}")
    builder.adjust(1)
    builder.row(InlineKeyboardButton(text="🔙 Назад", callback_data="menu_chemistry"))
    return builder.as_markup()

def get_chemistry_task_topic_keyboard(topic_num: str):
    builder = InlineKeyboardBuilder()
    builder.button(text="📐 Формулы и алгоритм", callback_data=f"chemtask_formulas:{topic_num}")
    builder.button(text="📋 Список задач", callback_data=f"chemtask_list:{topic_num}")
    builder.adjust(1)
    builder.row(InlineKeyboardButton(text="🔙 К темам", callback_data="chemistry_tasks"))
    return builder.as_markup()

def get_chemistry_formulas_keyboard(topic_num: str):
    builder = InlineKeyboardBuilder()
    builder.row(InlineKeyboardButton(text="🔙 Назад", callback_data=f"chemtask_topic:{topic_num}"))
    return builder.as_markup()

def get_chemistry_task_list_keyboard(topic_num: str):
    builder = InlineKeyboardBuilder()
    topic = CHEMISTRY_TASKS[topic_num]
    for task in topic["tasks"]:
        builder.button(text=f"📝 Задача {task['num']}", callback_data=f"chemtask_show:{topic_num}:{task['num']}")
    builder.adjust(1)
    builder.row(InlineKeyboardButton(text="🔙 Назад", callback_data=f"chemtask_topic:{topic_num}"))
    return builder.as_markup()

def get_chemistry_task_detail_keyboard(topic_num: str, task_num: int):
    builder = InlineKeyboardBuilder()
    tasks = CHEMISTRY_TASKS[topic_num]["tasks"]
    nums = [t["num"] for t in tasks]
    idx = nums.index(task_num)
    nav = []
    if idx > 0:
        nav.append(InlineKeyboardButton(text="⬅️ Предыдущая", callback_data=f"chemtask_show:{topic_num}:{nums[idx-1]}"))
    if idx < len(nums) - 1:
        nav.append(InlineKeyboardButton(text="Следующая ➡️", callback_data=f"chemtask_show:{topic_num}:{nums[idx+1]}"))
    if nav:
        builder.row(*nav)
    builder.row(InlineKeyboardButton(text="🔙 К списку задач", callback_data=f"chemtask_list:{topic_num}"))
    return builder.as_markup()

# ==================== ГЛУБОКИЕ ССЫЛКИ (t.me/BOT?start=...) ====================
SECTION_DEEPLINKS = {
    "physics_tasks": (
        f"🧮 <b>Задачи по физике</b>\n{DIVIDER}\n\nВыбери тему:",
        lambda uid: get_physics_tasks_topics_keyboard(),
    ),
    "support_menu": (
        get_support_text(),
        lambda uid: get_support_keyboard(uid),
    ),
}

# ==================== ОБРАБОТЧИКИ ====================
@dp.message(CommandStart())
async def cmd_start(message: Message):
    user_id = message.from_user.id
    is_new_user = user_id not in stats["total_users"]
    stats["total_users"].add(user_id)
    stats["start_count"] += 1
    save_stats()

    payload = message.text.split(maxsplit=1)
    deep_link_key = None
    if len(payload) > 1:
        if payload[1].startswith("ref_"):
            referrer_id_str = payload[1][len("ref_"):]
            if referrer_id_str.isdigit():
                await register_referral(int(referrer_id_str), user_id)
        else:
            deep_link_key = payload[1]

    if not await is_subscribed(user_id):
        builder = InlineKeyboardBuilder()
        builder.button(text="📢 Открыть канал Vmeda_examen", url="https://t.me/Vmeda_examen")
        await message.answer(
            "👋 <b>Добро пожаловать!</b>\n\n"
            "Этот бот поможет подготовиться к экзаменам ВМедА:\n"
            "🧬 биология · ⚛️ физика · 🧪 химия\n\n"
            f"{DIVIDER}\n"
            "🔒 Чтобы пользоваться ботом, подпишись на канал:\n"
            "👉 https://t.me/Vmeda_examen\n\n"
            "После подписки нажми /start ещё раз.",
            parse_mode="HTML",
            reply_markup=builder.as_markup()
        )
        return

    if deep_link_key in SECTION_DEEPLINKS:
        text, keyboard_func = SECTION_DEEPLINKS[deep_link_key]
        await message.answer(text, parse_mode="HTML", reply_markup=keyboard_func(user_id))
        return

    if deep_link_key and deep_link_key.startswith("q_"):
        q_num = deep_link_key[len("q_"):]
        if q_num in QUESTIONS:
            stats["question_opened"][q_num] = stats["question_opened"].get(q_num, 0) + 1
            save_stats()
            q = QUESTIONS[q_num]
            header = f"❓ <b>Вопрос {q_num}</b>"
            body = f"{header}\n{DIVIDER}\n\n<b>{q['title']}</b>\n\n{q['answer']}"
            short_caption = f"{header}\n{DIVIDER}\n\n<b>{q['title']}</b>"
            await send_answer(message, body, short_caption, q, get_question_answer_keyboard(q_num), edit=False)
            return

    greeting = "🎉 <b>С возвращением!</b>" if not is_new_user else "👋 <b>Привет!</b>"
    await message.answer(
        f"{greeting}\n\nВыбери предмет для подготовки:",
        parse_mode="HTML",
        reply_markup=get_main_menu(user_id)
    )

@dp.message(Command("stats"))
async def cmd_stats(message: Message):
    if not is_admin(message.from_user.id):
        return
    text = (
        "📊 <b>Статистика бота</b>\n"
        f"{DIVIDER}\n"
        f"👥 Уникальных пользователей: <b>{len(stats['total_users'])}</b>\n"
        f"▶️ Запусков бота: <b>{stats['start_count']}</b>\n"
        f"❓ Вопросов просмотрено: <b>{sum(stats['question_opened'].values())}</b>\n"
        f"🎲 Случайных билетов открыто: <b>{stats['random_ticket_used']}</b>\n"
        f"🎲 Случайных вопросов открыто: <b>{stats['random_question_used']}</b>\n"
        f"📢 Рассылок отправлено: <b>{stats.get('broadcast_count', 0)}</b>"
    )
    await message.answer(text, parse_mode="HTML")

@dp.message(Command("broadcast"))
async def cmd_broadcast(message: Message):
    if not is_admin(message.from_user.id):
        return

    text = message.html_text.split(maxsplit=1)
    if len(text) < 2 or not text[1].strip():
        await message.answer(
            "✏️ <b>Публичное сообщение от администрации</b>\n\n"
            "Использование:\n<code>/broadcast Текст сообщения</code>",
            parse_mode="HTML"
        )
        return

    announcement = text[1]
    body = (
        "📢 <b>Сообщение от администрации</b>\n"
        f"{DIVIDER}\n\n"
        f"{announcement}"
    )

    recipients = list(stats["total_users"])
    status = await message.answer(f"⏳ Рассылка запущена для {len(recipients)} пользователей...")

    sent, failed = 0, 0
    for user_id in recipients:
        try:
            await bot.send_message(user_id, body, parse_mode="HTML")
            sent += 1
        except Exception:
            failed += 1
        await asyncio.sleep(0.05)

    stats["broadcast_count"] = stats.get("broadcast_count", 0) + 1
    save_stats()

    await safe_edit_text(
        status,
        "✅ <b>Рассылка завершена</b>\n"
        f"{DIVIDER}\n"
        f"Доставлено: <b>{sent}</b>\n"
        f"Не доставлено: <b>{failed}</b>",
        parse_mode="HTML"
    )

# ==================== АДМИН-ПАНЕЛЬ ====================
ADMIN_PENDING: dict = {}  # admin_id -> {"action": ...}
ADMIN_CHANNEL_POST_PREVIEW: dict = {}  # admin_id -> {"text": ..., "buttons": [(label, url), ...]}
from handlers import admin as admin_handlers  # noqa: E402 — mid-file by design, see handlers/admin.py
dp.include_router(admin_handlers.router)

ADMIN_USERLIST_PAGE_SIZE = admin_handlers.ADMIN_USERLIST_PAGE_SIZE
parse_channel_post_buttons = admin_handlers.parse_channel_post_buttons
build_channel_post_builder = admin_handlers.build_channel_post_builder
build_channel_post_keyboard = admin_handlers.build_channel_post_keyboard
get_admin_menu = admin_handlers.get_admin_menu
get_admin_battle_keyboard = admin_handlers.get_admin_battle_keyboard
get_admin_battle_text = admin_handlers.get_admin_battle_text
get_admin_announcements_keyboard = admin_handlers.get_admin_announcements_keyboard
get_admin_back_keyboard = admin_handlers.get_admin_back_keyboard
resolve_user_by_username = admin_handlers.resolve_user_by_username
format_admin_target_label = admin_handlers.format_admin_target_label
format_user_line = admin_handlers.format_user_line
get_admin_userlist_page = admin_handlers.get_admin_userlist_page
cb_admin_panel = admin_handlers.cb_admin_panel
cb_admin_battle_menu = admin_handlers.cb_admin_battle_menu
cb_admin_announcements_menu = admin_handlers.cb_admin_announcements_menu
cb_admin_battle_last_results = admin_handlers.cb_admin_battle_last_results
cb_admin_battle_start_confirm = admin_handlers.cb_admin_battle_start_confirm
cb_admin_battle_start_go = admin_handlers.cb_admin_battle_start_go
cb_admin_histology_promo_confirm = admin_handlers.cb_admin_histology_promo_confirm
cb_admin_histology_promo_go = admin_handlers.cb_admin_histology_promo_go
cb_admin_global_promo_confirm = admin_handlers.cb_admin_global_promo_confirm
cb_admin_global_promo_go = admin_handlers.cb_admin_global_promo_go
cb_admin_global_promo_12h_confirm = admin_handlers.cb_admin_global_promo_12h_confirm
cb_admin_global_promo_12h_go = admin_handlers.cb_admin_global_promo_12h_go
cb_admin_restore_restrictions_confirm = admin_handlers.cb_admin_restore_restrictions_confirm
cb_admin_restore_restrictions_go = admin_handlers.cb_admin_restore_restrictions_go
cb_admin_battle_end_confirm = admin_handlers.cb_admin_battle_end_confirm
cb_admin_battle_end_go = admin_handlers.cb_admin_battle_end_go
cb_admin_battle_remind_confirm = admin_handlers.cb_admin_battle_remind_confirm
cb_admin_battle_remind_go = admin_handlers.cb_admin_battle_remind_go
cb_admin_restore_access_confirm = admin_handlers.cb_admin_restore_access_confirm
cb_admin_restore_access_go = admin_handlers.cb_admin_restore_access_go
cb_admin_referral_reminder_confirm = admin_handlers.cb_admin_referral_reminder_confirm
cb_admin_referral_reminder_go = admin_handlers.cb_admin_referral_reminder_go
cb_admin_discount_promo_confirm = admin_handlers.cb_admin_discount_promo_confirm
cb_admin_discount_promo_go = admin_handlers.cb_admin_discount_promo_go
cb_admin_stats = admin_handlers.cb_admin_stats
get_admin_stats_keyboard = admin_handlers.get_admin_stats_keyboard
cb_admin_ai_breaker_reset = admin_handlers.cb_admin_ai_breaker_reset
get_ai_cache_queue_text = admin_handlers.get_ai_cache_queue_text
get_ai_cache_queue_keyboard = admin_handlers.get_ai_cache_queue_keyboard
cb_admin_ai_cache_queue = admin_handlers.cb_admin_ai_cache_queue
cb_admin_ai_cache_approve = admin_handlers.cb_admin_ai_cache_approve
cb_admin_ai_cache_reject = admin_handlers.cb_admin_ai_cache_reject
cb_admin_export_stats = admin_handlers.cb_admin_export_stats
cb_admin_userlist = admin_handlers.cb_admin_userlist
cb_admin_grant_prompt = admin_handlers.cb_admin_grant_prompt
cb_admin_revoke_prompt = admin_handlers.cb_admin_revoke_prompt
cb_admin_grant_anatomy_demo_prompt = admin_handlers.cb_admin_grant_anatomy_demo_prompt
cb_admin_revoke_anatomy_demo_prompt = admin_handlers.cb_admin_revoke_anatomy_demo_prompt
cb_admin_grant_assistant_prompt = admin_handlers.cb_admin_grant_assistant_prompt
cb_admin_revoke_assistant_prompt = admin_handlers.cb_admin_revoke_assistant_prompt
cb_admin_grant_payment_admin_prompt = admin_handlers.cb_admin_grant_payment_admin_prompt
cb_admin_revoke_payment_admin_prompt = admin_handlers.cb_admin_revoke_payment_admin_prompt
cb_admin_dm_prompt = admin_handlers.cb_admin_dm_prompt
cb_admin_donation_prompt = admin_handlers.cb_admin_donation_prompt
cb_admin_subscription_prompt = admin_handlers.cb_admin_subscription_prompt
cb_admin_announce_support_confirm = admin_handlers.cb_admin_announce_support_confirm
cb_admin_announce_support_go = admin_handlers.cb_admin_announce_support_go
cb_admin_announce_subscription_confirm = admin_handlers.cb_admin_announce_subscription_confirm
cb_admin_announce_subscription_go = admin_handlers.cb_admin_announce_subscription_go
cb_admin_announce_anatomy_confirm = admin_handlers.cb_admin_announce_anatomy_confirm
cb_admin_announce_anatomy_go = admin_handlers.cb_admin_announce_anatomy_go
cb_admin_announce_anatomy_exam_confirm = admin_handlers.cb_admin_announce_anatomy_exam_confirm
cb_admin_announce_anatomy_exam_go = admin_handlers.cb_admin_announce_anatomy_exam_go
cb_admin_announce_anatomy_latin_confirm = admin_handlers.cb_admin_announce_anatomy_latin_confirm
cb_admin_announce_anatomy_latin_go = admin_handlers.cb_admin_announce_anatomy_latin_go
cb_admin_announce_ai_confirm = admin_handlers.cb_admin_announce_ai_confirm
cb_admin_announce_ai_go = admin_handlers.cb_admin_announce_ai_go
cb_admin_channel_post_prompt = admin_handlers.cb_admin_channel_post_prompt
cb_admin_channel_post_go = admin_handlers.cb_admin_channel_post_go
cb_admin_channel_post_cancel = admin_handlers.cb_admin_channel_post_cancel
get_assistant_admin_menu_text = admin_handlers.get_assistant_admin_menu_text
get_assistant_admin_menu_keyboard = admin_handlers.get_assistant_admin_menu_keyboard
get_assistant_back_keyboard = admin_handlers.get_assistant_back_keyboard
get_assistant_stats_text = admin_handlers.get_assistant_stats_text
cb_assistant_panel = admin_handlers.cb_assistant_panel
cb_assistant_stats = admin_handlers.cb_assistant_stats
cb_assistant_dm_prompt = admin_handlers.cb_assistant_dm_prompt
cb_assistant_dm_approve = admin_handlers.cb_assistant_dm_approve
cb_assistant_dm_reject = admin_handlers.cb_assistant_dm_reject
get_payment_admin_menu_text = admin_handlers.get_payment_admin_menu_text
get_payment_admin_menu_keyboard = admin_handlers.get_payment_admin_menu_keyboard
cb_payment_admin_panel = admin_handlers.cb_payment_admin_panel

@dp.message(Command("admin"))
async def cmd_admin(message: Message):
    user_id = message.from_user.id
    if is_admin(user_id):
        await message.answer(
            f"🛠 <b>Админ-панель</b>\n{DIVIDER}\n\nВыбери действие:",
            parse_mode="HTML",
            reply_markup=get_admin_menu()
        )
        return
    if is_assistant_admin(user_id):
        await message.answer(
            get_assistant_admin_menu_text(),
            parse_mode="HTML",
            reply_markup=get_assistant_admin_menu_keyboard()
        )
        return
    if is_payment_admin(user_id):
        await message.answer(
            get_payment_admin_menu_text(),
            parse_mode="HTML",
            reply_markup=get_payment_admin_menu_keyboard()
        )

# Все callback_query-хендлеры и клавиатуры/тексты-хелперы админ-панели (cb_admin_panel и далее,
# 46 хендлеров) перенесены в handlers/admin.py — см. импорт и dp.include_router(admin_handlers.router)
# в начале этого блока. ADMIN_PENDING/ADMIN_CHANNEL_POST_PREVIEW остаются здесь (handle_admin_pending_action
# ниже — @dp.message(F.text) напрямую на dp, не через Router, см. docstring handlers/admin.py) —
# обращения к ним из перенесённых хендлеров идут через tb.ADMIN_PENDING/tb.ADMIN_CHANNEL_POST_PREVIEW.

@dp.message(F.text)
async def handle_admin_pending_action(message: Message):
    admin_id = message.from_user.id
    if not is_admin(admin_id) or admin_id not in ADMIN_PENDING:
        raise SkipHandler
    if message.text.startswith("/"):
        raise SkipHandler

    pending = ADMIN_PENDING[admin_id]
    action = pending["action"]

    if action in (
        "grant", "revoke", "grant_anatomy_demo", "revoke_anatomy_demo",
        "grant_assistant_admin", "revoke_assistant_admin",
        "grant_payment_admin", "revoke_payment_admin",
        "dm_username", "record_donation_username", "record_subscription_username",
    ):
        raw_input = message.text.strip()
        username, target_id = resolve_user_by_username(raw_input)
        if not target_id:
            identifier = raw_input.lstrip("@")
            if identifier.isdigit():
                await message.answer(
                    f"⚠️ Пользователь с ID <code>{identifier}</code> не найден — он ещё не писал боту.\n"
                    "Попробуй ещё раз или вернись в /admin.",
                    parse_mode="HTML"
                )
            else:
                await message.answer(
                    f"⚠️ Пользователь @{identifier} не найден — он ещё не писал боту, либо сменил username. "
                    "Можно также ввести его числовой ID.\n"
                    "Попробуй ещё раз или вернись в /admin.",
                    parse_mode="HTML"
                )
            return

        label = format_admin_target_label(username, target_id)

        if action == "grant":
            if target_id not in stats["manual_access_granted"]:
                stats["manual_access_granted"].append(target_id)
                save_stats()
            del ADMIN_PENDING[admin_id]
            await message.answer(f"✅ Доступ выдан {label}.", parse_mode="HTML")
            try:
                await bot.send_message(
                    target_id,
                    "🎉 Администратор открыл тебе полный доступ к боту без необходимости приглашать друзей!",
                    parse_mode="HTML"
                )
            except Exception:
                logger.exception("Не удалось уведомить пользователя %s о выдаче доступа", target_id)

        elif action == "revoke":
            if target_id in stats["manual_access_granted"]:
                stats["manual_access_granted"].remove(target_id)
                save_stats()
            del ADMIN_PENDING[admin_id]
            await message.answer(
                f"✅ Ручной доступ для {label} отозван.\n"
                "Если у пользователя уже есть свои рефералы, доступ всё равно останется открытым.",
                parse_mode="HTML"
            )

        elif action == "grant_anatomy_demo":
            if target_id not in stats["manual_anatomy_demo_granted"]:
                stats["manual_anatomy_demo_granted"].append(target_id)
                save_stats()
            del ADMIN_PENDING[admin_id]
            await message.answer(f"✅ Демо-доступ к Анатомии выдан {label}.", parse_mode="HTML")
            try:
                await bot.send_message(
                    target_id,
                    "🦴 Администратор открыл тебе демо-доступ к разделу «Анатомия»!",
                    parse_mode="HTML"
                )
            except Exception:
                logger.exception("Не удалось уведомить пользователя %s о выдаче демо-доступа к анатомии", target_id)

        elif action == "revoke_anatomy_demo":
            if target_id in stats["manual_anatomy_demo_granted"]:
                stats["manual_anatomy_demo_granted"].remove(target_id)
                save_stats()
            del ADMIN_PENDING[admin_id]
            await message.answer(f"✅ Демо-доступ к Анатомии для {label} отозван.", parse_mode="HTML")

        elif action == "grant_assistant_admin":
            if target_id not in stats["assistant_admins"]:
                stats["assistant_admins"].append(target_id)
                save_stats()
            del ADMIN_PENDING[admin_id]
            await message.answer(f"✅ {label} назначен(а) помощником администратора.", parse_mode="HTML")
            try:
                await bot.send_message(
                    target_id,
                    "🧑‍💼 Тебя назначили помощником администратора!\n\n"
                    "Открой /admin — там доступна статистика бота и возможность написать "
                    "пользователю (сообщение уйдёт только после подтверждения главным админом). "
                    "Также у тебя теперь есть доступ ко всем разделам бота.",
                    parse_mode="HTML"
                )
            except Exception:
                logger.exception("Не удалось уведомить пользователя %s о назначении помощником", target_id)

        elif action == "revoke_assistant_admin":
            if target_id in stats["assistant_admins"]:
                stats["assistant_admins"].remove(target_id)
                save_stats()
            del ADMIN_PENDING[admin_id]
            await message.answer(f"✅ {label} больше не помощник администратора.", parse_mode="HTML")

        elif action == "grant_payment_admin":
            if target_id not in stats["payment_admins"]:
                stats["payment_admins"].append(target_id)
                save_stats()
            del ADMIN_PENDING[admin_id]
            await message.answer(f"✅ {label} назначен(а) админом платежей.", parse_mode="HTML")
            try:
                await bot.send_message(
                    target_id,
                    "💳 Тебя назначили админом платежей!\n\n"
                    "Открой /admin — там доступно подтверждение рублёвых заявок на оплату (те же "
                    "one-tap кнопки, что приходят обычным админам) и рассылка анонсов из подраздела "
                    "«Анонсы».",
                    parse_mode="HTML"
                )
            except Exception:
                logger.exception("Не удалось уведомить пользователя %s о назначении админом платежей", target_id)

        elif action == "revoke_payment_admin":
            if target_id in stats["payment_admins"]:
                stats["payment_admins"].remove(target_id)
                save_stats()
            del ADMIN_PENDING[admin_id]
            await message.answer(f"✅ {label} больше не админ платежей.", parse_mode="HTML")

        elif action == "dm_username":
            ADMIN_PENDING[admin_id] = {"action": "dm_message", "target_id": target_id, "target_label": label}
            await message.answer(f"✅ Нашёл {label}. Теперь отправь текст сообщения для него.", parse_mode="HTML")

        elif action == "record_donation_username":
            ADMIN_PENDING[admin_id] = {"action": "record_donation_amount", "target_id": target_id, "target_label": label}
            await message.answer(f"✅ Нашёл {label}. Теперь пришли сумму в рублях (целое число).", parse_mode="HTML")

        elif action == "record_subscription_username":
            ADMIN_PENDING[admin_id] = {"action": "record_subscription_tier", "target_id": target_id, "target_label": label}
            tier_lines = "\n".join(
                f"{t} — {cfg['title']} ({cfg['price_rub']}₽)" for t, cfg in ACTIVE_SUBSCRIPTION_TIERS.items()
            )
            await message.answer(
                f"✅ Нашёл {label}. Выбери тариф кнопкой ниже или пришли номер:\n\n{tier_lines}",
                parse_mode="HTML",
                reply_markup=get_admin_tier_reply_keyboard()
            )
        return

    if action == "record_donation_amount":
        target_id = pending["target_id"]
        target_label = pending["target_label"]
        raw = message.text.strip().replace(" ", "")
        if not raw.isdigit() or int(raw) <= 0:
            await message.answer("⚠️ Введи, пожалуйста, положительное целое число рублей.")
            return
        amount = int(raw)
        del ADMIN_PENDING[admin_id]
        uid_str = str(target_id)
        stats["donor_rubles"][uid_str] = stats["donor_rubles"].get(uid_str, 0) + amount
        save_stats()
        await message.answer(f"✅ Записано пожертвование {amount}₽ от {target_label}.", parse_mode="HTML")
        return

    if action == "record_subscription_tier":
        target_id = pending["target_id"]
        target_label = pending["target_label"]
        raw = message.text.strip()
        if raw in ("❌ Отмена", "Отмена"):
            del ADMIN_PENDING[admin_id]
            await message.answer("Отменено.", reply_markup=ReplyKeyboardRemove())
            return
        tier_match = re.match(r"\d+", raw)
        tier_id = int(tier_match.group()) if tier_match else None
        if tier_id not in ACTIVE_SUBSCRIPTION_TIERS:
            tier_lines = "\n".join(
                f"{t} — {cfg['title']}" for t, cfg in ACTIVE_SUBSCRIPTION_TIERS.items()
            )
            await message.answer(f"⚠️ Введи номер тарифа из списка:\n\n{tier_lines}", reply_markup=get_admin_tier_reply_keyboard())
            return
        cfg = ACTIVE_SUBSCRIPTION_TIERS[tier_id]
        if cfg.get("subject_choice_required"):
            ADMIN_PENDING[admin_id] = {
                "action": "record_subscription_subject",
                "target_id": target_id, "target_label": target_label, "tier_id": tier_id,
            }
            await message.answer(
                f"✅ Тариф «{cfg['title']}». Какой предмет выбрать?",
                reply_markup=get_admin_subject_reply_keyboard()
            )
            return
        del ADMIN_PENDING[admin_id]
        await message.answer(
            f"✅ Подписка «{cfg['title']}» выдана {target_label}.",
            parse_mode="HTML",
            reply_markup=ReplyKeyboardRemove()
        )
        await grant_subscription_and_notify_buyer(target_id, tier_id, "rubles_manual", cfg["price_rub"])
        return

    if action == "record_subscription_subject":
        target_id = pending["target_id"]
        target_label = pending["target_label"]
        tier_id = pending["tier_id"]
        raw = message.text.strip()
        if raw in ("❌ Отмена", "Отмена"):
            del ADMIN_PENDING[admin_id]
            await message.answer("Отменено.", reply_markup=ReplyKeyboardRemove())
            return
        subject = ADMIN_SUBJECT_LABELS_RU.get(raw)
        if not subject:
            await message.answer(
                "⚠️ Выбери предмет кнопкой ниже.",
                reply_markup=get_admin_subject_reply_keyboard()
            )
            return
        cfg = SUBSCRIPTION_TIERS[tier_id]
        del ADMIN_PENDING[admin_id]
        await message.answer(
            f"✅ Подписка «{cfg['title']}» ({SUBJECT_TITLES[subject]}) выдана {target_label}.",
            parse_mode="HTML",
            reply_markup=ReplyKeyboardRemove()
        )
        await grant_subscription_and_notify_buyer(target_id, tier_id, "rubles_manual", cfg["price_rub"], subject)
        return

    if action == "dm_message":
        target_id = pending["target_id"]
        target_label = pending["target_label"]
        del ADMIN_PENDING[admin_id]
        try:
            await bot.send_message(
                target_id,
                f"✉️ <b>Личное сообщение от администрации</b>\n{DIVIDER}\n\n{message.html_text}",
                parse_mode="HTML"
            )
            await message.answer(f"✅ Сообщение отправлено {target_label}.", parse_mode="HTML")
        except Exception:
            logger.exception("Не удалось отправить личное сообщение пользователю %s", target_id)
            await message.answer(f"⚠️ Не удалось отправить сообщение {target_label} — возможно, он заблокировал бота.", parse_mode="HTML")
        return

    if action == "channel_post_text":
        ADMIN_PENDING[admin_id] = {"action": "channel_post_buttons", "text": message.html_text}
        await message.answer(
            "🔘 <b>Кнопки под постом</b>\n\n"
            "Пришли по одной кнопке на строке в формате:\n"
            "<code>Текст кнопки | https://ссылка</code>\n\n"
            "Можно несколько строк — будет несколько кнопок друг под другом.\n"
            "Если кнопки не нужны — пришли «-».",
            parse_mode="HTML"
        )
        return

    if action == "channel_post_buttons":
        raw = message.text or ""
        if raw.strip() in ("-", "нет", "пропустить", "skip"):
            buttons = []
        else:
            buttons = parse_channel_post_buttons(raw)
            if buttons is None:
                await message.answer(
                    "⚠️ Не понял формат. Каждая строка: <code>Текст кнопки | https://ссылка</code>. "
                    "Либо пришли «-», если кнопки не нужны.",
                    parse_mode="HTML"
                )
                return
        post_text = pending["text"]
        del ADMIN_PENDING[admin_id]
        ADMIN_CHANNEL_POST_PREVIEW[admin_id] = {"text": post_text, "buttons": buttons}
        builder = build_channel_post_builder(buttons)
        builder.row(InlineKeyboardButton(text="✅ Опубликовать в канал", callback_data="admin_channel_post_go"))
        builder.row(InlineKeyboardButton(text="❌ Отмена", callback_data="admin_channel_post_cancel"))
        await message.answer(
            f"👀 <b>Предпросмотр поста для {CHANNEL_ID}:</b>\n{DIVIDER}\n\n{post_text}",
            parse_mode="HTML",
            reply_markup=builder.as_markup()
        )
        return

# ==================== ПОМОЩНИК АДМИНИСТРАТОРА ====================
# Отдельная, сильно урезанная версия админ-панели для user_id из stats["assistant_admins"]
# (назначаются/снимаются только полным админом — см. admin_grant_assistant_prompt/
# admin_revoke_assistant_prompt выше). Помощник НЕ получает доступ к обычной админ-панели
# (get_admin_menu/cb_admin_panel/handle_admin_pending_action остаются is_admin-only) — только
# к своему собственному меню из двух пунктов: урезанная статистика и сообщение пользователю,
# которое не уходит напрямую, а ставится на подтверждение всем ADMIN_IDS (одно-тап confirm/
# reject, тот же паттерн гонки через pop(), что и у admin_confirm_sub/rollcall_confirm).
# Доступ ко ВСЕМ разделам контента у помощника уже есть через is_admin_or_assistant() —
# это отдельный, самостоятельный механизм, определённый в самом начале файла.
ASSISTANT_PENDING: dict = {}  # assistant_id -> {"action": ..., ...}
ASSISTANT_DM_REQUESTS: dict = {}  # request_id (str) -> {assistant_id, assistant_label, target_id, target_label, text_html}
_assistant_dm_request_seq = 0

def _next_assistant_dm_request_id() -> str:
    global _assistant_dm_request_seq
    _assistant_dm_request_seq += 1
    return str(_assistant_dm_request_seq)

# Клавиатуры/тексты панели помощника и её callback_query-хендлеры (cb_assistant_panel,
# cb_assistant_stats, cb_assistant_dm_prompt) перенесены в handlers/admin.py — см. импорт и
# dp.include_router(admin_handlers.router) в начале блока «АДМИН-ПАНЕЛЬ» выше. ASSISTANT_PENDING/
# ASSISTANT_DM_REQUESTS остаются здесь (см. docstring handlers/admin.py) — обращения к ним из
# перенесённых хендлеров идут через tb.ASSISTANT_PENDING/tb.ASSISTANT_DM_REQUESTS.
async def notify_admins_of_assistant_dm_request(req_id: str) -> None:
    req = ASSISTANT_DM_REQUESTS[req_id]
    text = (
        f"🧑‍💼 <b>Помощник просит согласовать сообщение</b>\n{DIVIDER}\n\n"
        f"От: {req['assistant_label']}\n"
        f"Кому: {req['target_label']}\n\n"
        f"Текст:\n{req['text_html']}"
    )
    builder = InlineKeyboardBuilder()
    builder.button(text="✅ Отправить", callback_data=f"assistant_dm_approve:{req_id}")
    builder.button(text="❌ Отклонить", callback_data=f"assistant_dm_reject:{req_id}")
    builder.adjust(1)
    for admin_id in ADMIN_IDS:
        try:
            await bot.send_message(admin_id, text, parse_mode="HTML", reply_markup=builder.as_markup())
        except Exception:
            logger.exception("Не удалось уведомить админа %s о запросе помощника на сообщение", admin_id)

@dp.message(F.text)
async def handle_assistant_pending_action(message: Message):
    assistant_id = message.from_user.id
    if not is_assistant_admin(assistant_id) or assistant_id not in ASSISTANT_PENDING:
        raise SkipHandler
    if message.text.startswith("/"):
        raise SkipHandler

    pending = ASSISTANT_PENDING[assistant_id]
    action = pending["action"]

    if action == "dm_username":
        raw_input = message.text.strip()
        username, target_id = resolve_user_by_username(raw_input)
        if not target_id:
            identifier = raw_input.lstrip("@")
            if identifier.isdigit():
                await message.answer(f"⚠️ Пользователь с ID <code>{identifier}</code> не найден — он ещё не писал боту.", parse_mode="HTML")
            else:
                await message.answer(f"⚠️ Пользователь @{identifier} не найден — он ещё не писал боту, либо сменил username.", parse_mode="HTML")
            return
        label = format_admin_target_label(username, target_id)
        ASSISTANT_PENDING[assistant_id] = {"action": "dm_message", "target_id": target_id, "target_label": label}
        await message.answer(f"✅ Нашёл {label}. Теперь отправь текст сообщения для него.", parse_mode="HTML")
        return

    if action == "dm_message":
        target_id = pending["target_id"]
        target_label = pending["target_label"]
        del ASSISTANT_PENDING[assistant_id]
        assistant_label = format_admin_target_label(
            stats["user_username"].get(str(assistant_id)), assistant_id
        )
        req_id = _next_assistant_dm_request_id()
        ASSISTANT_DM_REQUESTS[req_id] = {
            "assistant_id": assistant_id,
            "assistant_label": assistant_label,
            "target_id": target_id,
            "target_label": target_label,
            "text_html": message.html_text,
        }
        await message.answer(
            f"✅ Запрос отправлен главному админу на согласование — сообщение для {target_label} "
            "уйдёт после подтверждения.",
            parse_mode="HTML"
        )
        await notify_admins_of_assistant_dm_request(req_id)
        return

# cb_assistant_dm_approve/cb_assistant_dm_reject перенесены в handlers/admin.py — см. импорт
# и dp.include_router(admin_handlers.router) в начале блока «АДМИН-ПАНЕЛЬ» выше.

# ==================== МЕНЮ ====================
# Главное меню Биологии и режим опроса (quiz_*) перенесены в handlers/biology.py вместе с
# билетами/вопросами — см. dp.include_router(biology_handlers.router) и реэкспорт дальше по файлу.

@dp.callback_query(F.data == "back_to_main")
async def cb_back_to_main(callback: CallbackQuery):
    await callback.answer()
    await safe_edit_text(
        callback.message,
        "🏠 <b>Главное меню</b>\n\nВыбери предмет для подготовки:",
        parse_mode="HTML",
        reply_markup=get_main_menu(callback.from_user.id)
    )

@dp.callback_query(F.data.startswith("course_menu:"))
async def cb_course_menu(callback: CallbackQuery):
    course = int(callback.data.split(":")[1])
    await callback.answer()
    await safe_edit_text(
        callback.message,
        get_course_menu_text(course),
        parse_mode="HTML",
        reply_markup=get_course_menu_keyboard(course, callback.from_user.id)
    )

@dp.callback_query(F.data == "referral_info")
async def cb_referral_info(callback: CallbackQuery):
    await callback.answer()
    user_id = callback.from_user.id
    keyboard = get_referral_full_access_keyboard(user_id) if has_free_access(user_id) else get_referral_back_keyboard()
    await safe_edit_text(
        callback.message,
        get_referral_status_text(user_id),
        parse_mode="HTML",
        reply_markup=keyboard
    )

@dp.callback_query(F.data == "referral_leaderboard")
async def cb_referral_leaderboard(callback: CallbackQuery):
    await callback.answer()
    await safe_edit_text(
        callback.message,
        get_referral_leaderboard_text(callback.from_user.id),
        parse_mode="HTML",
        reply_markup=get_referral_leaderboard_keyboard()
    )

@dp.callback_query(F.data == "referral_battle")
async def cb_referral_battle(callback: CallbackQuery):
    await callback.answer()
    await safe_edit_text(
        callback.message,
        get_battle_text(callback.from_user.id),
        parse_mode="HTML",
        reply_markup=get_battle_keyboard(),
        disable_web_page_preview=True,
    )

@dp.callback_query(F.data == "support_menu")
async def cb_support_menu(callback: CallbackQuery):
    await callback.answer()
    DONATION_PENDING.pop(callback.from_user.id, None)
    await safe_edit_text(callback.message, get_support_text(), parse_mode="HTML", reply_markup=get_support_keyboard(callback.from_user.id))

@dp.callback_query(F.data == "toggle_donor_visibility")
async def cb_toggle_donor_visibility(callback: CallbackQuery):
    uid_str = str(callback.from_user.id)
    hidden = not stats["donor_hide_name"].get(uid_str, False)
    stats["donor_hide_name"][uid_str] = hidden
    save_stats()
    await callback.answer("Теперь ты анонимен в рейтинге" if hidden else "Теперь твой ник виден в рейтинге")
    await safe_edit_text(callback.message, get_support_text(), parse_mode="HTML", reply_markup=get_support_keyboard(callback.from_user.id))

@dp.callback_query(F.data == "donors_leaderboard")
async def cb_donors_leaderboard(callback: CallbackQuery):
    await callback.answer()
    await safe_edit_text(
        callback.message,
        get_donors_leaderboard_text(),
        parse_mode="HTML",
        reply_markup=get_donors_leaderboard_keyboard()
    )

@dp.callback_query(F.data == "donate_stars_menu")
async def cb_donate_stars_menu(callback: CallbackQuery):
    await callback.answer()
    DONATION_PENDING.pop(callback.from_user.id, None)
    await safe_edit_text(callback.message, get_stars_menu_text(), parse_mode="HTML", reply_markup=get_stars_menu_keyboard())

@dp.callback_query(F.data.startswith("donate_stars_amount:"))
async def cb_donate_stars_amount(callback: CallbackQuery):
    await callback.answer()
    amount = int(callback.data.split(":")[1])
    await safe_edit_text(
        callback.message,
        get_visibility_choice_text(amount, " ⭐"),
        parse_mode="HTML",
        reply_markup=get_stars_visibility_keyboard(amount)
    )

@dp.callback_query(F.data.startswith("donate_stars_confirm:"))
async def cb_donate_stars_confirm(callback: CallbackQuery):
    await callback.answer()
    _, amount_s, visibility = callback.data.split(":")
    stats["donor_hide_name"][str(callback.from_user.id)] = (visibility == "anon")
    save_stats()
    await send_stars_invoice(callback.from_user.id, int(amount_s))

@dp.callback_query(F.data == "donate_stars_custom")
async def cb_donate_stars_custom(callback: CallbackQuery):
    await callback.answer()
    DONATION_PENDING[callback.from_user.id] = {"type": "stars"}
    builder = InlineKeyboardBuilder()
    builder.row(InlineKeyboardButton(text="🔙 Назад", callback_data="donate_stars_menu"))
    await safe_edit_text(
        callback.message,
        f"✏️ <b>Своё количество звёзд</b>\n{DIVIDER}\n\nВведи число от {STARS_MIN} до {STARS_MAX}:",
        parse_mode="HTML",
        reply_markup=builder.as_markup()
    )

@dp.callback_query(F.data == "donate_rubles_menu")
async def cb_donate_rubles_menu(callback: CallbackQuery):
    await callback.answer()
    DONATION_PENDING.pop(callback.from_user.id, None)
    await safe_edit_text(callback.message, get_rubles_menu_text(), parse_mode="HTML", reply_markup=get_rubles_menu_keyboard())

@dp.callback_query(F.data.startswith("donate_rubles_amount:"))
async def cb_donate_rubles_amount(callback: CallbackQuery):
    await callback.answer()
    amount = int(callback.data.split(":")[1])
    await safe_edit_text(
        callback.message,
        get_visibility_choice_text(amount, "₽"),
        parse_mode="HTML",
        reply_markup=get_rubles_visibility_keyboard(amount)
    )

@dp.callback_query(F.data.startswith("donate_rubles_confirm:"))
async def cb_donate_rubles_confirm(callback: CallbackQuery):
    await callback.answer()
    _, amount_s, visibility = callback.data.split(":")
    amount = int(amount_s)
    stats["donor_hide_name"][str(callback.from_user.id)] = (visibility == "anon")
    save_stats()
    await safe_edit_text(
        callback.message,
        get_rubles_donation_message_text(amount),
        parse_mode="HTML",
        reply_markup=get_rubles_donation_keyboard(amount),
        disable_web_page_preview=True,
    )

@dp.callback_query(F.data == "donate_rubles_custom")
async def cb_donate_rubles_custom(callback: CallbackQuery):
    await callback.answer()
    DONATION_PENDING[callback.from_user.id] = {"type": "rubles"}
    builder = InlineKeyboardBuilder()
    builder.row(InlineKeyboardButton(text="🔙 Назад", callback_data="donate_rubles_menu"))
    await safe_edit_text(
        callback.message,
        f"✏️ <b>Своя сумма в рублях</b>\n{DIVIDER}\n\nВведи сумму числом (от {RUBLES_MIN} до {RUBLES_MAX}):",
        parse_mode="HTML",
        reply_markup=builder.as_markup()
    )

# ==================== ПЛАТНАЯ ПОДПИСКА (UI и оплата) ====================
def format_subscription_expiry(expires) -> str:
    if expires is None:
        return "навсегда"
    # МСК, а не системный часовой пояс контейнера (UTC) — иначе дата истечения, вычисленная как
    # московская полночь (см. APP_TIMEZONE/services/access.py), могла бы отобразиться на день раньше.
    return f"до {datetime.fromtimestamp(expires, APP_TIMEZONE).strftime('%d.%m.%Y')}"

async def grant_subscription_and_notify_buyer(
    target_id: int, tier_id: int, method: str, price: int, subject: str | None = None
) -> None:
    """Общая точка для всех трёх путей выдачи подписки (Stars, ручное подтверждение рублей
    админом, быстрое подтверждение по кнопке) — выдаёт тариф и шлёт покупателю одно и то же
    сообщение об активации + апсейл на следующий тариф. method различает источник для
    статистики платежей: "stars"/"rubles" — реальная подтверждённая оплата, "rubles_manual" —
    ручная выдача админом (например, бесплатно другу) и в выручку не считается."""
    grant_subscription(target_id, tier_id, method, price, subject)
    cfg = SUBSCRIPTION_TIERS[tier_id]
    sub = get_subscription(target_id)
    scope_label = get_subscription_scope_label(sub)
    text = (
        f"🎉 <b>Подписка «{cfg['title']}» активирована!</b>\n\n"
        f"Доступ {scope_label} открыт — {format_subscription_expiry(sub['expires'])}.\n"
        "Правило про рефералов для тебя больше не действует. Спасибо за поддержку! 🙏😇"
    )
    text += get_tier_upsell_text(tier_id)
    keyboard = get_tier_upsell_keyboard(tier_id)
    try:
        await bot.send_message(target_id, text, parse_mode="HTML", reply_markup=keyboard)
    except Exception:
        logger.exception("Не удалось уведомить пользователя %s о выдаче подписки", target_id)

def get_admin_payment_confirm_text(cfg: dict, user, subject: str | None = None, price: int | None = None) -> str:
    price = price if price is not None else cfg["price_rub"]
    subject_line = f"\nПредмет: {SUBJECT_TITLES[subject]}" if subject else ""
    discount_line = " (со скидкой 10%)" if price != cfg["price_rub"] else ""
    return (
        f"💰 <b>Запрос на подтверждение оплаты</b>\n{DIVIDER}\n\n"
        f"Тариф: «{cfg['title']}» — {price}₽{discount_line}{subject_line}\n"
        f"От: {html.escape(user.full_name)} "
        f"({f'@{user.username} ' if user.username else ''}ID <code>{user.id}</code>)\n\n"
        "Нажми ниже, когда увидишь перевод в чате с @vmeda_helper — подписка выдастся сразу."
    )

def get_admin_payment_confirm_keyboard(tier_id: int, target_id: int, subject: str | None = None, price: int | None = None):
    cfg = SUBSCRIPTION_TIERS[tier_id]
    price = price if price is not None else cfg["price_rub"]
    builder = InlineKeyboardBuilder()
    builder.button(
        text="✅ Подтвердить оплату",
        callback_data=f"admin_confirm_sub:{tier_id}:{target_id}:{subject or '-'}:{price}"
    )
    builder.button(
        text="❌ Отклонить",
        callback_data=f"admin_reject_sub:{tier_id}:{target_id}:{subject or '-'}"
    )
    builder.adjust(1)
    return builder.as_markup()

async def notify_admins_of_payment_request(
    tier_id: int, target_id: int, user, subject: str | None = None, price: int | None = None
) -> None:
    cfg = SUBSCRIPTION_TIERS[tier_id]
    text = get_admin_payment_confirm_text(cfg, user, subject, price)
    keyboard = get_admin_payment_confirm_keyboard(tier_id, target_id, subject, price)
    # Админы платежей (отдельная роль, stats["payment_admins"]) тоже должны увидеть one-tap
    # кнопку подтверждения — иначе назначение роли ничего не даёт, см. cb_admin_confirm_sub.
    for recipient_id in ADMIN_IDS | set(stats["payment_admins"]):
        try:
            await bot.send_message(recipient_id, text, parse_mode="HTML", reply_markup=keyboard)
        except Exception:
            logger.exception("Не удалось уведомить админа %s о запросе оплаты", recipient_id)

def get_my_subscription_status_block(user_id: int) -> str:
    sub = get_subscription(user_id)
    if not sub or not has_active_subscription(user_id):
        return ""
    cfg = SUBSCRIPTION_TIERS.get(sub["tier"], {})
    scope_label = get_subscription_scope_label(sub)
    return (
        f"✅ У тебя активна подписка «{cfg.get('title', '')}»\n"
        f"Доступ {scope_label} — {format_subscription_expiry(sub['expires'])}.\n\n"
    )

def _next_upsell_tier_id(tier_id: int) -> int | None:
    cfg = SUBSCRIPTION_TIERS[tier_id]
    candidates = sorted(
        (t for t, c in ACTIVE_SUBSCRIPTION_TIERS.items() if c["price_rub"] > cfg["price_rub"]),
        key=lambda t: ACTIVE_SUBSCRIPTION_TIERS[t]["price_rub"]
    )
    return candidates[0] if candidates else None

def get_tier_upsell_text(tier_id: int) -> str:
    cfg = SUBSCRIPTION_TIERS[tier_id]
    nxt_id = _next_upsell_tier_id(tier_id)
    if nxt_id is None:
        return ""
    nxt = SUBSCRIPTION_TIERS[nxt_id]
    diff_rub = nxt["price_rub"] - cfg["price_rub"]
    diff_stars = nxt["price_stars"] - cfg["price_stars"]
    return (
        f"\n\n💡 <b>Выгоднее:</b> тариф «{nxt['emoji']} {nxt['title']}» — всего на "
        f"<b>{diff_rub}₽ / {diff_stars}⭐</b> дороже (<b>{nxt['price_rub']}₽ / {nxt['price_stars']}⭐</b> "
        f"вместо {cfg['price_rub']}₽), а даёт: {nxt['benefits'][0].lower()}."
    )

def get_tier_upsell_keyboard(tier_id: int):
    nxt_id = _next_upsell_tier_id(tier_id)
    if nxt_id is None:
        return None
    nxt = SUBSCRIPTION_TIERS[nxt_id]
    builder = InlineKeyboardBuilder()
    builder.button(text=f"⬆️ Перейти на «{nxt['short']}» за {nxt['price_rub']}₽", callback_data=f"sub_tier:{nxt_id}")
    return builder.as_markup()

def get_subscription_menu_text(user_id: int) -> str:
    lines = [f"💎 <b>Подписка без рефералов</b>\n{DIVIDER}\n"]
    lines.append(
        "⚠️ Разработка и содержание бота требуют серьёзных затрат — поэтому в дополнение "
        "к бесплатному доступу за рефералов мы вынуждены были добавить платные подписки. "
        "Так бот сможет и дальше жить, обновляться и получать новые разделы.\n"
    )
    lines.append(
        "🔬 Раздел <b>Гистологии</b> уже полностью готов и проработан: все микрофотографии "
        "и протоколы-описания взяты именно с препаратов академии, а содержание сверено "
        "с преподавателями.\n"
    )
    status = get_my_subscription_status_block(user_id)
    if status:
        lines.append(status)
    lines.append(
        "Не хочешь ждать или звать друзей? Открой доступ сразу оплатой — без рефералов "
        "и ограничений. Выбери вариант:\n\n"
        "⏳ Зачёркнутая цена — во сколько подписка будет обходиться с сентября, успей купить сейчас!\n"
    )
    best21, best26 = SUBSCRIPTION_TIERS[21], SUBSCRIPTION_TIERS[26]
    lines.append(
        "🏆 <b>ТОП-2 предложения:</b>\n"
        f"👉 «{best21['emoji']} {best21['title']}» — {best21['price_rub']}₽, "
        f"или «{best26['emoji']} {best26['title']}» — {best26['price_rub']}₽ — закрывает всё сразу! 🔥\n"
    )
    for tier_id, cfg in sorted_active_tiers():
        if cfg.get("badge"):
            lines.append(f"<b>{cfg['badge']}</b>")
        lines.append(f"{cfg['emoji']} <b>{cfg['title']}</b> — {get_tier_price_line(cfg)}")
        lines.append(f"• {cfg['benefits'][0]}")
        lines.append("")
    lines.append(
        "👆 Полное описание каждого тарифа — при выборе ниже.\n\n"
        "После оплаты правило про рефералов для тебя больше не действует — доступ "
        "открывается сразу и держится всё оплаченное время."
    )
    return "\n".join(lines)

def get_subscription_menu_keyboard():
    builder = InlineKeyboardBuilder()
    for tier_id, cfg in sorted_active_tiers():
        badge = f"{cfg['badge']} — " if cfg.get("badge") else ""
        builder.button(
            text=f"{badge}{cfg['emoji']} {cfg['short']} — {cfg['price_rub']}₽/{cfg['price_stars']}⭐",
            callback_data=f"sub_tier:{tier_id}"
        )
    builder.adjust(1)
    builder.row(InlineKeyboardButton(text="🔙 Назад в меню", callback_data="back_to_main"))
    return builder.as_markup()

# ==================== ВЫБОР КУРСА (витрина подписки) ====================
# "subscription_menu" (главная точка входа — кнопка "💎 Подписка" в главном меню, тизеры, все
# "Назад" в экранах покупки) теперь сначала показывает выбор курса, а не сразу плоский список
# всех активных тарифов — 9 карточек одним экраном плохо читаются, когда линейка выросла с 7 до
# 9 тарифов и стала курс-специфичной (см. спецификацию новой тарифной системы). Плоский список
# (get_subscription_menu_text/_keyboard выше, поведение не менялось) остаётся доступен через
# кнопку "📦 Все тарифы" на любом экране курса. Выбор курса нигде не сохраняется — состояние
# целиком живёт в callback_data, ничего не пишется в stats.
FIRST_YEAR_TIER_IDS = (21, 23, 25, 26)  # Месяц / До зачёта по химии / Весь 1 курс / До конца 2 курса
SECOND_YEAR_AUTUMN_TIER_IDS = (20, 22, 24, 26)  # Пересдача 7 дней / Все пересдачи / Зимняя сессия / До конца 2 курса
SECOND_YEAR_WINTER_TIER_IDS = (21, 24, 26, 27)  # Месяц / Зимняя сессия / До конца 2 курса / 2 года MAX
_COURSE_TITLES = {"year1": "1 курс", "year2": "2 курс"}

def _second_year_tier_ids() -> tuple[int, ...]:
    """До 1 ноября 2026 — сезон пересдач (август-октябрь), после — обычная витрина 2 курса
    (см. NOV_1_2026_CUTOFF — тот же рубеж, на котором истекает тариф 22 "Все пересдачи")."""
    return SECOND_YEAR_AUTUMN_TIER_IDS if time.time() < NOV_1_2026_CUTOFF else SECOND_YEAR_WINTER_TIER_IDS

def _course_tier_ids(course: str) -> tuple[int, ...]:
    if course == "year1":
        return FIRST_YEAR_TIER_IDS
    if course == "year2":
        return _second_year_tier_ids()
    return ()

def get_subscription_course_picker_text(user_id: int) -> str:
    lines = [f"💎 <b>Подписка без рефералов</b>\n{DIVIDER}\n"]
    status = get_my_subscription_status_block(user_id)
    if status:
        lines.append(status)
    lines.append("🎓 Выбери свой курс — покажем подходящие тарифы:")
    return "\n".join(lines)

def get_subscription_course_picker_keyboard():
    builder = InlineKeyboardBuilder()
    builder.button(text="1️⃣ Первый курс", callback_data="subscription_course:year1")
    builder.button(text="2️⃣ Второй курс", callback_data="subscription_course:year2")
    builder.button(text="📦 Все тарифы", callback_data="subscription_all_tiers")
    builder.adjust(1)
    builder.row(InlineKeyboardButton(text="🔙 Назад в меню", callback_data="back_to_main"))
    return builder.as_markup()

def get_subscription_course_tiers_text(user_id: int, course: str) -> str:
    lines = [f"💎 <b>Подписка — {_COURSE_TITLES.get(course, course)}</b>\n{DIVIDER}\n"]
    status = get_my_subscription_status_block(user_id)
    if status:
        lines.append(status)
    for tier_id in _course_tier_ids(course):
        cfg = SUBSCRIPTION_TIERS[tier_id]
        if cfg.get("badge"):
            lines.append(f"<b>{cfg['badge']}</b>")
        lines.append(f"{cfg['emoji']} <b>{cfg['title']}</b> — {get_tier_price_line(cfg)}")
        for b in cfg["benefits"]:
            lines.append(f"• {b}")
        lines.append("")
    lines.append("📦 Не подошёл ни один вариант? Смотри «Все тарифы» ниже.")
    return "\n".join(lines)

def get_subscription_course_tiers_keyboard(course: str):
    builder = InlineKeyboardBuilder()
    for tier_id in _course_tier_ids(course):
        cfg = SUBSCRIPTION_TIERS[tier_id]
        badge = f"{cfg['badge']} — " if cfg.get("badge") else ""
        builder.button(
            text=f"{badge}{cfg['emoji']} {cfg['short']} — {cfg['price_rub']}₽/{cfg['price_stars']}⭐",
            callback_data=f"sub_tier:{tier_id}"
        )
    builder.button(text="📦 Все тарифы", callback_data="subscription_all_tiers")
    builder.adjust(1)
    builder.row(InlineKeyboardButton(text="🔙 Назад", callback_data="subscription_menu"))
    return builder.as_markup()

def get_sub_tier_text(tier_id: int) -> str:
    cfg = SUBSCRIPTION_TIERS[tier_id]
    lines = [f"{cfg['emoji']} <b>{cfg['title']}</b>"]
    if cfg.get("badge"):
        lines.append(f"<b>{cfg['badge']}</b>")
    lines.append(f"{DIVIDER}\n")
    if cfg.get("joke"):
        lines.append(f"<i>{cfg['joke']}</i>\n")
    for b in cfg["benefits"]:
        lines.append(f"• {b}")
    lines.append(f"\nЦена: {get_tier_price_line(cfg)}")
    lines.append("⏳ Зачёркнутая цена — во сколько подписка будет обходиться с сентября, успей купить сейчас!")
    lines.append(get_tier_upsell_text(tier_id))
    if cfg.get("subject_choice_required"):
        lines.append("\nСначала выбери предмет, потом способ оплаты:")
    else:
        lines.append("\nВыбери способ оплаты:")
    return "\n".join(lines)

def get_sub_tier_keyboard(tier_id: int):
    cfg = SUBSCRIPTION_TIERS[tier_id]
    builder = InlineKeyboardBuilder()
    if cfg.get("subject_choice_required"):
        builder.button(text="🧬 Биология", callback_data=f"sub_subject:{tier_id}:biology")
        builder.button(text="⚛️ Физика", callback_data=f"sub_subject:{tier_id}:physics")
        builder.button(text="🧪 Химия", callback_data=f"sub_subject:{tier_id}:chemistry")
        builder.adjust(1)
    else:
        builder.button(text=f"⭐ Оплатить {cfg['price_stars']} звёзд", callback_data=f"buy_sub_stars:{tier_id}")
        builder.button(text=f"💵 Оплатить {cfg['price_rub']}₽", callback_data=f"buy_sub_rubles:{tier_id}")
        builder.adjust(1)
    nxt_kb = get_tier_upsell_keyboard(tier_id)
    if nxt_kb:
        builder.row(nxt_kb.inline_keyboard[0][0])
    builder.row(InlineKeyboardButton(text="🔙 Назад", callback_data="subscription_menu"))
    return builder.as_markup()

def get_sub_subject_keyboard(tier_id: int, subject: str):
    cfg = SUBSCRIPTION_TIERS[tier_id]
    builder = InlineKeyboardBuilder()
    builder.button(
        text=f"⭐ Оплатить {cfg['price_stars']} звёзд",
        callback_data=f"buy_sub_stars_subj:{tier_id}:{subject}"
    )
    builder.button(
        text=f"💵 Оплатить {cfg['price_rub']}₽",
        callback_data=f"buy_sub_rubles_subj:{tier_id}:{subject}"
    )
    builder.adjust(1)
    builder.row(InlineKeyboardButton(text="🔙 Назад", callback_data=f"sub_tier:{tier_id}"))
    return builder.as_markup()

def get_sub_rubles_message_text(tier_id: int, subject: str | None = None, price: int | None = None) -> str:
    cfg = SUBSCRIPTION_TIERS[tier_id]
    price = price if price is not None else cfg["price_rub"]
    subject_line = f" ({SUBJECT_TITLES[subject]})" if subject else ""
    return (
        f"💵 <b>Оплата подписки «{cfg['title']}»{subject_line} — {price}₽</b>\n{DIVIDER}\n\n"
        f'Нажми на кнопку ниже — откроется чат с <a href="{HELPER_ACCOUNT_URL}">@vmeda_helper</a>, '
        "сообщение с тарифом уже будет готово. Отправь его и переведи по присланным реквизитам — "
        "как только оплата подтвердится, подписка будет включена вручную.\n\n"
        "Спасибо, что поддерживаешь бота! 🙏"
    )

def get_sub_rubles_keyboard(tier_id: int, subject: str | None = None, price: int | None = None):
    cfg = SUBSCRIPTION_TIERS[tier_id]
    price = price if price is not None else cfg["price_rub"]
    subject_line = f" ({SUBJECT_TITLES[subject]})" if subject else ""
    template = (
        f"Привет! Хочу оформить подписку «{cfg['title']}»{subject_line} за {price}₽ в боте "
        "VMEDA_examen_bot. Подскажи, пожалуйста, реквизиты для перевода."
    )
    url = f"{HELPER_ACCOUNT_URL}?text={urllib.parse.quote(template)}"
    back_data = f"sub_subject:{tier_id}:{subject}" if subject else f"sub_tier:{tier_id}"
    builder = InlineKeyboardBuilder()
    builder.row(InlineKeyboardButton(text="💸 Написать @vmeda_helper", url=url))
    builder.row(InlineKeyboardButton(text="🔙 Назад", callback_data=back_data))
    return builder.as_markup()

async def send_subscription_stars_invoice(
    chat_id: int, tier_id: int, subject: str | None = None, discount: bool = False
) -> None:
    cfg = SUBSCRIPTION_TIERS[tier_id]
    price_stars = discount_price(cfg["price_stars"]) if discount else cfg["price_stars"]
    subject_line = f" ({SUBJECT_TITLES[subject]})" if subject else ""
    discount_line = " — скидка 10%" if discount else ""
    await bot.send_invoice(
        chat_id=chat_id,
        title=f"Подписка: {cfg['title']}{subject_line}{discount_line}",
        description=f"VMEDA_examen_bot — подписка «{cfg['title']}»{subject_line}{discount_line}. Доступ откроется сразу после оплаты.",
        payload=f"sub_stars_{tier_id}_{subject or '-'}_{chat_id}_{int(time.time())}",
        provider_token="",
        currency="XTR",
        prices=[LabeledPrice(label=cfg["title"], amount=price_stars)],
    )

def get_discount_offer_text(tier_id: int) -> str:
    cfg = SUBSCRIPTION_TIERS[tier_id]
    lines = [
        f"🔥 <b>Скидка {int(DISCOUNT_RATE * 100)}% — специально для тебя!</b>\n{DIVIDER}\n",
        f"{cfg['emoji']} <b>{cfg['title']}</b>\n",
        f"Цена со скидкой: <b>{discount_price(cfg['price_rub'])}₽</b> <s>{cfg['price_rub']}₽</s> / "
        f"<b>{discount_price(cfg['price_stars'])}⭐</b> <s>{cfg['price_stars']}⭐</s>\n",
    ]
    lines += [f"• {b}" for b in cfg["benefits"]]
    lines.append("\n⏳ Скидка действует ограниченное время — не тяни с оплатой!\n\nВыбери способ оплаты:")
    return "\n".join(lines)

def get_discount_offer_keyboard(tier_id: int):
    cfg = SUBSCRIPTION_TIERS[tier_id]
    builder = InlineKeyboardBuilder()
    builder.button(
        text=f"⭐ Оплатить {discount_price(cfg['price_stars'])} звёзд",
        callback_data=f"buy_sub_stars_discount:{tier_id}"
    )
    builder.button(
        text=f"💵 Оплатить {discount_price(cfg['price_rub'])}₽",
        callback_data=f"buy_sub_rubles_discount:{tier_id}"
    )
    builder.adjust(1)
    builder.row(InlineKeyboardButton(text="🔙 Назад", callback_data="subscription_menu"))
    return builder.as_markup()

def get_subscription_teaser_keyboard():
    builder = InlineKeyboardBuilder()
    builder.row(InlineKeyboardButton(text="💎 Открыть доступ без рефералов", callback_data="subscription_menu"))
    builder.row(InlineKeyboardButton(text="🚀 Запустить Medical_vpn_bot", url=MEDICAL_VPN_URL))
    return builder.as_markup()

@dp.callback_query(F.data == "subscription_menu")
async def cb_subscription_menu(callback: CallbackQuery):
    await callback.answer()
    await safe_edit_text(
        callback.message,
        get_subscription_course_picker_text(callback.from_user.id),
        parse_mode="HTML",
        reply_markup=get_subscription_course_picker_keyboard(),
        disable_web_page_preview=True,
    )

@dp.callback_query(F.data == "subscription_all_tiers")
async def cb_subscription_all_tiers(callback: CallbackQuery):
    await callback.answer()
    await safe_edit_text(
        callback.message,
        get_subscription_menu_text(callback.from_user.id),
        parse_mode="HTML",
        reply_markup=get_subscription_menu_keyboard(),
        disable_web_page_preview=True,
    )

@dp.callback_query(F.data.startswith("subscription_course:"))
async def cb_subscription_course(callback: CallbackQuery):
    course = callback.data.split(":", 1)[1]
    if course not in _COURSE_TITLES:
        await callback.answer("Курс не найден", show_alert=True)
        return
    await callback.answer()
    await safe_edit_text(
        callback.message,
        get_subscription_course_tiers_text(callback.from_user.id, course),
        parse_mode="HTML",
        reply_markup=get_subscription_course_tiers_keyboard(course),
        disable_web_page_preview=True,
    )

def get_tier_retired_text(tier_id: int) -> str:
    cfg = SUBSCRIPTION_TIERS[tier_id]
    return (
        f"{cfg['emoji']} <b>{cfg['title']}</b>\n{DIVIDER}\n\n"
        "🚫 Этот тариф больше не продаётся — линейка подписок обновилась.\n\n"
        "Посмотри актуальные варианты ниже 👇"
    )

def get_tier_retired_keyboard():
    builder = InlineKeyboardBuilder()
    builder.row(InlineKeyboardButton(text="💎 Посмотреть актуальные тарифы", callback_data="subscription_menu"))
    return builder.as_markup()

@dp.callback_query(F.data.startswith("sub_tier:"))
async def cb_sub_tier(callback: CallbackQuery):
    tier_id = int(callback.data.split(":")[1])
    if tier_id not in SUBSCRIPTION_TIERS:
        await callback.answer("Тариф не найден", show_alert=True)
        return
    await callback.answer()
    if tier_id not in ACTIVE_SUBSCRIPTION_TIERS:
        await safe_edit_text(
            callback.message,
            get_tier_retired_text(tier_id),
            parse_mode="HTML",
            reply_markup=get_tier_retired_keyboard()
        )
        return
    await safe_edit_text(
        callback.message,
        get_sub_tier_text(tier_id),
        parse_mode="HTML",
        reply_markup=get_sub_tier_keyboard(tier_id)
    )

@dp.callback_query(F.data.startswith("sub_subject:"))
async def cb_sub_subject(callback: CallbackQuery):
    _, tier_id_raw, subject = callback.data.split(":")
    tier_id = int(tier_id_raw)
    if tier_id not in SUBSCRIPTION_TIERS or subject not in SUBJECT_TITLES:
        await callback.answer("Тариф не найден", show_alert=True)
        return
    await callback.answer()
    if tier_id not in ACTIVE_SUBSCRIPTION_TIERS:
        await safe_edit_text(
            callback.message,
            get_tier_retired_text(tier_id),
            parse_mode="HTML",
            reply_markup=get_tier_retired_keyboard()
        )
        return
    cfg = SUBSCRIPTION_TIERS[tier_id]
    text = (
        f"{cfg['emoji']} <b>{cfg['title']}</b> — {SUBJECT_TITLES[subject]}\n{DIVIDER}\n\n"
        f"Цена: {get_tier_price_line(cfg)}\n"
        "⏳ Зачёркнутая цена — во сколько подписка будет обходиться с сентября, успей купить сейчас!\n\n"
        "Выбери способ оплаты:"
    )
    await safe_edit_text(callback.message, text, parse_mode="HTML", reply_markup=get_sub_subject_keyboard(tier_id, subject))

@dp.callback_query(F.data.startswith("buy_sub_stars:"))
async def cb_buy_sub_stars(callback: CallbackQuery):
    tier_id = int(callback.data.split(":")[1])
    if tier_id not in SUBSCRIPTION_TIERS:
        await callback.answer("Тариф не найден", show_alert=True)
        return
    if tier_id not in ACTIVE_SUBSCRIPTION_TIERS:
        await callback.answer("Этот тариф больше не продаётся", show_alert=True)
        return
    await callback.answer()
    await send_subscription_stars_invoice(callback.from_user.id, tier_id)

@dp.callback_query(F.data.startswith("buy_sub_stars_subj:"))
async def cb_buy_sub_stars_subj(callback: CallbackQuery):
    _, tier_id_raw, subject = callback.data.split(":")
    tier_id = int(tier_id_raw)
    if tier_id not in SUBSCRIPTION_TIERS or subject not in SUBJECT_TITLES:
        await callback.answer("Тариф не найден", show_alert=True)
        return
    if tier_id not in ACTIVE_SUBSCRIPTION_TIERS:
        await callback.answer("Этот тариф больше не продаётся", show_alert=True)
        return
    await callback.answer()
    await send_subscription_stars_invoice(callback.from_user.id, tier_id, subject)

@dp.callback_query(F.data.startswith("buy_sub_rubles:"))
async def cb_buy_sub_rubles(callback: CallbackQuery):
    tier_id = int(callback.data.split(":")[1])
    if tier_id not in SUBSCRIPTION_TIERS:
        await callback.answer("Тариф не найден", show_alert=True)
        return
    if tier_id not in ACTIVE_SUBSCRIPTION_TIERS:
        await callback.answer("Этот тариф больше не продаётся", show_alert=True)
        return
    await callback.answer()
    await safe_edit_text(
        callback.message,
        get_sub_rubles_message_text(tier_id),
        parse_mode="HTML",
        reply_markup=get_sub_rubles_keyboard(tier_id),
        disable_web_page_preview=True,
    )
    await notify_admins_of_payment_request(tier_id, callback.from_user.id, callback.from_user)

@dp.callback_query(F.data.startswith("buy_sub_rubles_subj:"))
async def cb_buy_sub_rubles_subj(callback: CallbackQuery):
    _, tier_id_raw, subject = callback.data.split(":")
    tier_id = int(tier_id_raw)
    if tier_id not in SUBSCRIPTION_TIERS or subject not in SUBJECT_TITLES:
        await callback.answer("Тариф не найден", show_alert=True)
        return
    if tier_id not in ACTIVE_SUBSCRIPTION_TIERS:
        await callback.answer("Этот тариф больше не продаётся", show_alert=True)
        return
    await callback.answer()
    await safe_edit_text(
        callback.message,
        get_sub_rubles_message_text(tier_id, subject),
        parse_mode="HTML",
        reply_markup=get_sub_rubles_keyboard(tier_id, subject),
        disable_web_page_preview=True,
    )
    await notify_admins_of_payment_request(tier_id, callback.from_user.id, callback.from_user, subject)

@dp.callback_query(F.data.startswith("sub_discount:"))
async def cb_sub_discount(callback: CallbackQuery):
    tier_id = int(callback.data.split(":")[1])
    cfg = SUBSCRIPTION_TIERS.get(tier_id)
    if not cfg or cfg.get("subject_choice_required"):
        await callback.answer("Тариф не найден", show_alert=True)
        return
    if tier_id not in ACTIVE_SUBSCRIPTION_TIERS:
        await callback.answer("Этот тариф больше не продаётся", show_alert=True)
        return
    await callback.answer()
    await safe_edit_text(
        callback.message,
        get_discount_offer_text(tier_id),
        parse_mode="HTML",
        reply_markup=get_discount_offer_keyboard(tier_id),
    )

@dp.callback_query(F.data.startswith("buy_sub_stars_discount:"))
async def cb_buy_sub_stars_discount(callback: CallbackQuery):
    tier_id = int(callback.data.split(":")[1])
    if tier_id not in SUBSCRIPTION_TIERS:
        await callback.answer("Тариф не найден", show_alert=True)
        return
    if tier_id not in ACTIVE_SUBSCRIPTION_TIERS:
        await callback.answer("Этот тариф больше не продаётся", show_alert=True)
        return
    await callback.answer()
    await send_subscription_stars_invoice(callback.from_user.id, tier_id, discount=True)

@dp.callback_query(F.data.startswith("buy_sub_rubles_discount:"))
async def cb_buy_sub_rubles_discount(callback: CallbackQuery):
    tier_id = int(callback.data.split(":")[1])
    if tier_id not in SUBSCRIPTION_TIERS:
        await callback.answer("Тариф не найден", show_alert=True)
        return
    if tier_id not in ACTIVE_SUBSCRIPTION_TIERS:
        await callback.answer("Этот тариф больше не продаётся", show_alert=True)
        return
    await callback.answer()
    price = discount_price(SUBSCRIPTION_TIERS[tier_id]["price_rub"])
    await safe_edit_text(
        callback.message,
        get_sub_rubles_message_text(tier_id, price=price),
        parse_mode="HTML",
        reply_markup=get_sub_rubles_keyboard(tier_id, price=price),
        disable_web_page_preview=True,
    )
    await notify_admins_of_payment_request(tier_id, callback.from_user.id, callback.from_user, price=price)

@dp.callback_query(F.data.startswith("admin_confirm_sub:"))
async def cb_admin_confirm_sub(callback: CallbackQuery):
    if not (is_admin(callback.from_user.id) or is_payment_admin(callback.from_user.id)):
        await callback.answer()
        return
    parts = callback.data.split(":")
    _, tier_id_raw, target_id_raw, subject_raw = parts[:4]
    tier_id = int(tier_id_raw)
    target_id = int(target_id_raw)
    subject = subject_raw if subject_raw != "-" else None
    if tier_id not in SUBSCRIPTION_TIERS:
        await callback.answer("Тариф не найден", show_alert=True)
        return
    price = int(parts[4]) if len(parts) > 4 else SUBSCRIPTION_TIERS[tier_id]["price_rub"]

    existing = get_subscription(target_id)
    already_confirmed = (
        existing and existing.get("tier") == tier_id and existing.get("method") == "rubles"
        and has_active_subscription(target_id)
        and time.time() - existing.get("purchased_at", 0) < 600
    )
    if already_confirmed:
        await callback.answer("Уже подтверждено (скорее всего, другим админом)", show_alert=True)
        await safe_edit_text(
            callback.message,
            f"✅ Уже подтверждено — подписка «{SUBSCRIPTION_TIERS[tier_id]['title']}» "
            f"выдана {format_admin_target_label(None, target_id)}.",
            parse_mode="HTML"
        )
        return

    await callback.answer("Подтверждено ✅", show_alert=True)
    await grant_subscription_and_notify_buyer(target_id, tier_id, "rubles", price, subject)
    await safe_edit_text(
        callback.message,
        f"✅ Подтверждено — подписка «{SUBSCRIPTION_TIERS[tier_id]['title']}» "
        f"выдана {format_admin_target_label(None, target_id)}.",
        parse_mode="HTML"
    )

@dp.callback_query(F.data.startswith("admin_reject_sub:"))
async def cb_admin_reject_sub(callback: CallbackQuery):
    """Позволяет закрыть запрос на подтверждение оплаты, если покупатель так и не перевёл
    деньги (например, проигнорировал) — просто убирает заявку, ничего не выдаёт и не трогает
    статистику."""
    if not (is_admin(callback.from_user.id) or is_payment_admin(callback.from_user.id)):
        await callback.answer()
        return
    _, tier_id_raw, target_id_raw, subject_raw = callback.data.split(":")
    tier_id = int(tier_id_raw)
    target_id = int(target_id_raw)
    if tier_id not in SUBSCRIPTION_TIERS:
        await callback.answer("Тариф не найден", show_alert=True)
        return
    await callback.answer("Заявка отклонена", show_alert=True)
    await safe_edit_text(
        callback.message,
        f"❌ Отклонено — заявка на «{SUBSCRIPTION_TIERS[tier_id]['title']}» от "
        f"{format_admin_target_label(None, target_id)} закрыта без выдачи подписки.",
        parse_mode="HTML"
    )

@dp.pre_checkout_query()
async def handle_pre_checkout(pre_checkout_query) -> None:
    """Проверяет payload/тариф/сумму НЕПОСРЕДСТВЕННО перед оплатой, а не только доверяет тому, что
    сам инвойс когда-то был создан правильно (send_subscription_stars_invoice формирует payload/
    цену из живого SUBSCRIPTION_TIERS в момент выставления счёта — но между выставлением и оплатой
    тариф теоретически мог быть снят с продажи админом/деплоем, а сумма — не совпасть с ожидаемой
    из-за бага при формировании инвойса; отклонить здесь дешевле и безопаснее, чем молча выдать
    несоответствующую payload'у подписку в handle_successful_payment)."""
    payload = pre_checkout_query.invoice_payload or ""
    if payload.startswith("sub_stars_"):
        parts = payload.split("_")
        try:
            tier_id = int(parts[2])
        except (IndexError, ValueError):
            await pre_checkout_query.answer(ok=False, error_message="Некорректные данные платежа — попробуй оформить подписку заново.")
            return
        cfg = SUBSCRIPTION_TIERS.get(tier_id)
        if cfg is None or cfg.get("retired"):
            await pre_checkout_query.answer(ok=False, error_message="Этот тариф больше не продаётся — выбери другой в разделе «Подписка».")
            return
        expected_prices = {cfg["price_stars"], discount_price(cfg["price_stars"])}
        if pre_checkout_query.total_amount not in expected_prices:
            logger.error(
                "pre_checkout_query: сумма %s не совпадает с ожидаемой ценой тарифа %s (%s), payload=%s",
                pre_checkout_query.total_amount, tier_id, expected_prices, payload,
            )
            await pre_checkout_query.answer(ok=False, error_message="Сумма платежа не совпадает с ценой тарифа — попробуй оформить подписку заново.")
            return
    await pre_checkout_query.answer(ok=True)

@dp.message(F.successful_payment)
async def handle_successful_payment(message: Message):
    payment = message.successful_payment
    stars = payment.total_amount
    payload = payment.invoice_payload or ""

    # Идемпотентность: telegram_payment_charge_id — уникальный ID именно ЭТОЙ транзакции (у Stars
    # provider_payment_charge_id пуст, внешнего провайдера нет). Без этой проверки повторная
    # доставка одного и того же successful_payment (сетевой ретрай Telegram, повторный webhook и
    # т.п.) выдала бы подписку/донат ещё раз — списание у пользователя при этом одно, а не выдать
    # оплаченное только затем не дать выдать это же ещё раз важнее, чем пропустить с виду похожий,
    # но на самом деле новый платёж, у которого будет свой charge_id.
    charge_id = payment.telegram_payment_charge_id
    if charge_id and charge_id in stats["processed_payment_charge_ids"]:
        logger.warning(
            "Повторная доставка successful_payment для telegram_payment_charge_id=%s (user %s) — "
            "игнорируем, чтобы не выдать подписку/донат повторно", charge_id, message.from_user.id,
        )
        return
    if charge_id:
        stats["processed_payment_charge_ids"][charge_id] = {
            "user_id": message.from_user.id, "stars": stars, "payload": payload, "at": time.time(),
        }
        save_stats()

    if payload.startswith("sub_stars_"):
        parts = payload.split("_")
        tier_id = int(parts[2])
        subject = parts[3] if len(parts) > 3 and parts[3] != "-" else None
        cfg = SUBSCRIPTION_TIERS[tier_id]
        await grant_subscription_and_notify_buyer(message.from_user.id, tier_id, "stars", stars, subject)
        user = message.from_user
        for admin_id in ADMIN_IDS:
            try:
                await bot.send_message(
                    admin_id,
                    f"💎 <b>Новая подписка звёздами!</b>\n\n«{cfg['title']}» ({stars} ⭐) — "
                    f"{html.escape(user.full_name)} (ID <code>{user.id}</code>)",
                    parse_mode="HTML"
                )
            except Exception:
                logger.exception("Не удалось уведомить админа %s о подписке", admin_id)
        return

    stats["donations_stars_total"] += stars
    stats["donations_stars_count"] += 1
    uid_str = str(message.from_user.id)
    stats["donor_stars"][uid_str] = stats["donor_stars"].get(uid_str, 0) + stars
    save_stats()
    await message.answer(
        f"🎉 <b>Спасибо огромное за поддержку — {stars} ⭐!</b>\n\n"
        "Это очень помогает развивать бота дальше 🙏😇",
        parse_mode="HTML"
    )
    user = message.from_user
    for admin_id in ADMIN_IDS:
        try:
            await bot.send_message(
                admin_id,
                f"⭐ <b>Новое пожертвование звёздами!</b>\n\n{stars} ⭐ от "
                f"{html.escape(user.full_name)} (ID <code>{user.id}</code>)",
                parse_mode="HTML"
            )
        except Exception:
            logger.exception("Не удалось уведомить админа %s о пожертвовании", admin_id)

@dp.message(F.text)
async def handle_donation_pending_amount(message: Message):
    user_id = message.from_user.id
    if user_id not in DONATION_PENDING:
        raise SkipHandler
    if message.text.startswith("/"):
        raise SkipHandler

    pending = DONATION_PENDING[user_id]
    raw = message.text.strip().replace(" ", "")
    if not raw.isdigit():
        await message.answer("⚠️ Введи, пожалуйста, просто число.")
        return
    amount = int(raw)

    if pending["type"] == "stars":
        if not (STARS_MIN <= amount <= STARS_MAX):
            await message.answer(f"⚠️ Введи число от {STARS_MIN} до {STARS_MAX}.")
            return
        del DONATION_PENDING[user_id]
        await message.answer(
            get_visibility_choice_text(amount, " ⭐"),
            parse_mode="HTML",
            reply_markup=get_stars_visibility_keyboard(amount)
        )
    else:
        if not (RUBLES_MIN <= amount <= RUBLES_MAX):
            await message.answer(f"⚠️ Введи число от {RUBLES_MIN} до {RUBLES_MAX}.")
            return
        del DONATION_PENDING[user_id]
        await message.answer(
            get_visibility_choice_text(amount, "₽"),
            parse_mode="HTML",
            reply_markup=get_rubles_visibility_keyboard(amount)
        )

# Главное меню Физики и скачивание файлов перенесены в handlers/physics.py — см.
# dp.include_router(physics_handlers.router) и реэкспорт дальше по файлу.

# Главное меню Химии и скачивание файлов перенесены в handlers/chemistry.py — см.
# dp.include_router(chemistry_handlers.router) и реэкспорт дальше по файлу.

# ==================== ХИМИЯ (роутер/реэкспорт) ====================
# Хендлеры (теория с навигацией, билеты, задачи, лабораторные работы — все с уникальными
# callback_data-фильтрами, безопасно для порядка dp) вынесены в handlers/chemistry.py (свой
# Router) — здесь только регистрация роутера и реэкспорт имён. Клавиатурные билдеры и
# chemistry_tickets_access_ok остаются здесь (используются и cb_menu_chemistry/
# download_chemistry_*, которые тоже остаются) — см. docstring handlers/chemistry.py.
from handlers import chemistry as chemistry_handlers  # noqa: E402 — mid-file by design, see above

dp.include_router(chemistry_handlers.router)

cb_chemistry_theory = chemistry_handlers.cb_chemistry_theory
cb_show_theory_topic = chemistry_handlers.cb_show_theory_topic
cb_theory_list = chemistry_handlers.cb_theory_list
cb_chemistry_tickets = chemistry_handlers.cb_chemistry_tickets
cb_chem_theory_tickets = chemistry_handlers.cb_chem_theory_tickets
cb_chem_theory_ticket = chemistry_handlers.cb_chem_theory_ticket
cb_chem_theory_question = chemistry_handlers.cb_chem_theory_question
cb_chem_practice_tickets = chemistry_handlers.cb_chem_practice_tickets
cb_chem_practice_ticket = chemistry_handlers.cb_chem_practice_ticket
cb_chemistry_tasks = chemistry_handlers.cb_chemistry_tasks
cb_chemtask_topic = chemistry_handlers.cb_chemtask_topic
cb_chemtask_formulas = chemistry_handlers.cb_chemtask_formulas
cb_chemtask_list = chemistry_handlers.cb_chemtask_list
cb_chemtask_show = chemistry_handlers.cb_chemtask_show
cb_chemistry_labs = chemistry_handlers.cb_chemistry_labs
cb_show_lab = chemistry_handlers.cb_show_lab
cb_lab_summary = chemistry_handlers.cb_lab_summary
cb_lab_experiments = chemistry_handlers.cb_lab_experiments
cb_lab_calculations = chemistry_handlers.cb_lab_calculations
cb_menu_chemistry = chemistry_handlers.cb_menu_chemistry
cb_download_chemistry_labs = chemistry_handlers.cb_download_chemistry_labs
cb_download_chemistry_tasks = chemistry_handlers.cb_download_chemistry_tasks

# ==================== БИОЛОГИЯ — БИЛЕТЫ / ВОПРОСЫ ====================
# Билеты и вопросы (уникальные callback_data-фильтры, безопасно для порядка dp) вынесены в
# handlers/biology.py (свой Router) — здесь только регистрация роутера и реэкспорт имён. Режим
# опроса (quiz_*) и handle_question_number (F.text.isdigit(), сразу под этим блоком) сознательно
# НЕ перенесены — см. docstring handlers/biology.py.
from handlers import biology as biology_handlers  # noqa: E402 — mid-file by design, see above

dp.include_router(biology_handlers.router)

cb_random_ticket = biology_handlers.cb_random_ticket
show_ticket = biology_handlers.show_ticket
cb_ticket = biology_handlers.cb_ticket
cb_ticket_question = biology_handlers.cb_ticket_question
cb_question_page = biology_handlers.cb_question_page
cb_show_question = biology_handlers.cb_show_question
cb_question_random = biology_handlers.cb_question_random
cb_question_by_number = biology_handlers.cb_question_by_number
cb_question_search = biology_handlers.cb_question_search
get_biology_menu = biology_handlers.get_biology_menu
QUIZ_SESSION_SIZE = biology_handlers.QUIZ_SESSION_SIZE
QUIZ_SESSIONS = biology_handlers.QUIZ_SESSIONS
get_quiz_question_keyboard = biology_handlers.get_quiz_question_keyboard
get_quiz_answer_keyboard = biology_handlers.get_quiz_answer_keyboard
get_quiz_summary_keyboard = biology_handlers.get_quiz_summary_keyboard
start_quiz_session = biology_handlers.start_quiz_session
render_quiz_question = biology_handlers.render_quiz_question
render_quiz_answer = biology_handlers.render_quiz_answer
render_quiz_summary = biology_handlers.render_quiz_summary
get_ticket_keyboard = biology_handlers.get_ticket_keyboard
get_questions_main_menu = biology_handlers.get_questions_main_menu
get_biology_tickets_locked_text = biology_handlers.get_biology_tickets_locked_text
get_biology_tickets_locked_keyboard = biology_handlers.get_biology_tickets_locked_keyboard
cb_menu_biology = biology_handlers.cb_menu_biology
cb_download_biology_tickets = biology_handlers.cb_download_biology_tickets
cb_quiz_start = biology_handlers.cb_quiz_start
cb_quiz_show_answer = biology_handlers.cb_quiz_show_answer
cb_quiz_answer = biology_handlers.cb_quiz_answer
cb_quiz_stop = biology_handlers.cb_quiz_stop
cb_menu_tickets = biology_handlers.cb_menu_tickets
cb_menu_questions = biology_handlers.cb_menu_questions

@dp.message(F.text.isdigit())
async def handle_question_number(message: Message):
    q_num = message.text.strip()
    if q_num in QUESTIONS:
        stats["question_opened"][q_num] = stats["question_opened"].get(q_num, 0) + 1
        save_stats()
        q = QUESTIONS[q_num]
        header = f"❓ <b>Вопрос {q_num}</b>"
        body = f"{header}\n{DIVIDER}\n\n<b>{q['title']}</b>\n\n{q['answer']}"
        short_caption = f"{header}\n{DIVIDER}\n\n<b>{q['title']}</b>"
        await send_answer(message, body, short_caption, q, get_question_answer_keyboard(q_num), edit=False)
    else:
        await message.answer("⚠️ Вопрос с таким номером не найден.")

# ==================== VMEDA AI (MVP) ====================
# Phase 1 (MVP): пользователь присылает фото/текст задания -> модель отвечает напрямую,
# без Knowledge Base/верификации (это следующие фазы, см. согласованный roadmap).
# Работает поверх существующих паттернов бота: свой gate (не через has_free_access),
# свои ключи в stats.json (ai_usage, через .setdefault — обычная миграция), не трогает
# ни один существующий раздел/контент/тариф.
#
# Провайдеры/роутинг/RAG/промпты/подготовка фото вынесены в пакет ai/ (см. ai/service.py,
# ai/router.py, ai/rag.py, ai/prompts.py, ai/vision.py, ai/providers/*) — здесь остаётся
# только то, что завязано на stats/save_stats/is_admin/DIVIDER и сами @dp-хендлеры (перенос
# хендлеров в отдельный router — следующая фаза рефакторинга, см. CLAUDE.md).
AI_FREE_DAILY_LIMIT = 3
AI_SESSION_TIMEOUT_SECONDS = 20 * 60  # диалог считается закрытым после 20 минут без сообщений
AI_SESSIONS: dict = {}
# user_id -> {"task": TaskRepresentation|None, "messages": [...], "rag_context": str|None,
#             "bucket": str|None, "quick_answer": str|None, "last_active": ts, "processing": bool}
# "task" заполняется РОВНО один раз — при разборе самого первого сообщения сессии (фото или
# текста) через ai.vision_parser.parse_task (см. handle_ai_photo_input/handle_ai_text_input) —
# вместе с ним же в этот момент считаются rag_context (ai.rag.search_for_task) и bucket
# (ai.router.route_bucket); последующие ходы того же диалога (уточняющие вопросы, кнопка
# "Показать решение по шагам") переиспользуют то же rag_context/bucket и просто дописывают
# текст в "messages", не трогая vision/RAG повторно.

AI_USER_LOCKS: dict = {}  # user_id -> asyncio.Lock — ЖИВЁТ отдельно от AI_SESSIONS специально:
# session["processing"] (флаг ВНУТРИ словаря AI_SESSIONS[user_id]) не защищает от гонки, когда
# start_ai_session() целиком ЗАМЕНЯЕТ этот словарь новым (processing сбрасывается в False), пока
# СТАРЫЙ словарь всё ещё "processing": True — пользователь может тапнуть "AI" ещё раз посреди
# обработки предыдущего запроса, начав новую сессию и запустив второй, платный по квоте запрос,
# ДО того как первый вообще успеет списать квоту. Блокировка в отдельном, никогда не заменяемом
# словаре закрывает именно этот зазор, не трогая существующую (и по-прежнему валидную для
# дублей внутри одной и той же сессии) проверку session["processing"] — обе проверки нужны:
# lock.locked() ловит гонку между сессиями, session["processing"] — как и раньше, дубль внутри
# одной сессии (например, двойной тап по одной и той же кнопке).

def _get_ai_user_lock(user_id: int) -> asyncio.Lock:
    lock = AI_USER_LOCKS.get(user_id)
    if lock is None:
        lock = asyncio.Lock()
        AI_USER_LOCKS[user_id] = lock
    return lock

# ---- Бот-широкий предохранитель по одновременным запросам (независим от AI_USER_LOCKS выше,
# который ограничивает только ОДНОГО пользователя одним запросом за раз) — всплеск трафика (много
# разных пользователей, у каждого своя дневная квота) мог бы всё равно запустить неограниченное
# число дорогих запросов ОДНОВРЕМЕННО, раз лимит был только по-пользовательский ----
MAX_AI_CONCURRENT_REQUESTS = int(os.environ.get("AI_MAX_CONCURRENT_REQUESTS", "10"))

class _AIConcurrencyGate:
    """Небольшая неблокирующая замена asyncio.Semaphore под нашу семантику "нет слота — сразу
    отказ", а не "подождать, пока освободится" (обычное поведение Semaphore.acquire()/async with):
    очередь запросов молча ждущих своей очереди — то же самое накопление нагрузки, от которого
    вообще затевался этот предохранитель, просто отложенное на потом и невидимое пользователю (он
    просто долго не получает ответ вместо понятного "попробуй через минуту").

    try_acquire() — синхронный метод БЕЗ await внутри, поэтому проверка условия и инкремент
    счётчика происходят одним неразрывным шагом: между ними физически не может вклиниться другая
    корутина на том же event loop (в отличие от прежней схемы "check ai_concurrency_slot_available()
    ... await ... count += 1", где между проверкой и инкрементом успевали пройти чужие await-точки
    — несколько запросов могли одновременно увидеть свободный слот и все пройти проверку до того,
    как хоть один из них успевал застолбить его за собой, и потолок оказывался не жёстким)."""
    def __init__(self, limit: int):
        self._limit = limit
        self._count = 0

    def try_acquire(self) -> bool:
        if self._count >= self._limit:
            return False
        self._count += 1
        return True

    def release(self) -> None:
        self._count -= 1

AI_CONCURRENCY_GATE = _AIConcurrencyGate(MAX_AI_CONCURRENT_REQUESTS)

def get_ai_usage_today(user_id: int) -> int:
    entry = stats["ai_usage"].get(str(user_id))
    if not entry or entry.get("date") != local_today().isoformat():
        return 0
    return entry.get("count", 0)

def has_unlimited_ai(user_id: int) -> bool:
    return is_admin(user_id)

# ---- Платный AI-план подписки (независим от бесплатного дневного лимита выше) ----
# Новые тарифы (subscription_version 2) несут собственные ai_limit_type/ai_limit — "period"
# (общий пул запросов на весь срок действия тарифа, без сброса) либо "monthly" (сбрасывается
# каждый календарный месяц). Старые платные подписки (subscription_version отсутствует, т.е.
# читается как 1 — см. .get("subscription_version", 1)) получают фиксированный ежемесячный
# бонус LEGACY_PAID_AI_MONTHLY_BONUS, не завязанный ни на одно сохранённое в их записи поле —
# трогать саму запись подписки нельзя (см. CLAUDE.md/спецификацию новой тарифной линейки).
# Контентные права (anatomy/histology/scope/...) этот слой не читает и не меняет.

def _sub_ai_plan(user_id: int) -> tuple[str | None, int | None]:
    """(limit_type, limit) платного AI-плана подписки, либо (None, None) — тогда действует
    только обычный бесплатный дневной лимит (AI_FREE_DAILY_LIMIT)."""
    sub = get_subscription(user_id)
    if not sub or not has_active_subscription(user_id):
        return None, None
    if sub.get("subscription_version", 1) >= 2:
        cfg = SUBSCRIPTION_TIERS.get(sub.get("tier"), {})
        limit_type, limit = cfg.get("ai_limit_type"), cfg.get("ai_limit")
        if limit_type and limit:
            return limit_type, limit
        return None, None
    return "monthly", LEGACY_PAID_AI_MONTHLY_BONUS

def _current_ai_month_key() -> str:
    return local_today().strftime("%Y-%m")

def _get_sub_ai_used(sub: dict, limit_type: str) -> int:
    if limit_type == "monthly":
        entry = sub.get("ai_used_monthly")
        if not entry or entry.get("month") != _current_ai_month_key():
            return 0
        return entry.get("count", 0)
    return sub.get("ai_used_period", 0)

def _increment_sub_ai_usage(user_id: int, limit_type: str) -> None:
    sub = get_subscription(user_id)
    if sub is None:
        return
    if limit_type == "monthly":
        month = _current_ai_month_key()
        entry = sub.get("ai_used_monthly")
        if not entry or entry.get("month") != month:
            entry = {"month": month, "count": 0}
        entry["count"] += 1
        sub["ai_used_monthly"] = entry
    else:
        sub["ai_used_period"] = sub.get("ai_used_period", 0) + 1
    save_stats()

def sub_ai_requests_left(user_id: int) -> int | None:
    """None — нет отдельного платного AI-плана (см. _sub_ai_plan)."""
    limit_type, limit = _sub_ai_plan(user_id)
    if limit_type is None:
        return None
    sub = get_subscription(user_id)
    return max(0, limit - _get_sub_ai_used(sub, limit_type))

def increment_ai_usage(user_id: int) -> None:
    limit_type, _ = _sub_ai_plan(user_id)
    if limit_type is not None:
        _increment_sub_ai_usage(user_id, limit_type)
        return
    today = local_today().isoformat()
    entry = stats["ai_usage"].get(str(user_id))
    if not entry or entry.get("date") != today:
        entry = {"date": today, "count": 0}
    entry["count"] += 1
    stats["ai_usage"][str(user_id)] = entry
    save_stats()

def ai_requests_left(user_id: int) -> int:
    sub_left = sub_ai_requests_left(user_id)
    if sub_left is not None:
        return sub_left
    return max(0, AI_FREE_DAILY_LIMIT - get_ai_usage_today(user_id))

def ai_quota_ok(user_id: int) -> bool:
    return has_unlimited_ai(user_id) or ai_requests_left(user_id) > 0

def get_ai_quota_label(user_id: int) -> str:
    if has_unlimited_ai(user_id):
        return "♾ безлимит (админ)"
    limit_type, limit = _sub_ai_plan(user_id)
    if limit_type is not None:
        period_word = "в месяц" if limit_type == "monthly" else "на срок подписки"
        return f"{ai_requests_left(user_id)}/{limit} ({period_word})"
    return f"{ai_requests_left(user_id)}/{AI_FREE_DAILY_LIMIT}"

_AI_PROVIDER_PRICES = {
    "openai": (ai_openai.PRICE_INPUT_PER_1M, ai_openai.PRICE_OUTPUT_PER_1M),
    "grok": (ai_xai.PRICE_INPUT_PER_1M, ai_xai.PRICE_OUTPUT_PER_1M),
    "gemini": (ai_gemini.PRICE_INPUT_PER_1M, ai_gemini.PRICE_OUTPUT_PER_1M),
    # Эмбеддинги биллятся только по входным токенам — выходная цена 0, но запись всё равно идёт
    # через record_ai_cost с provider="openai-embeddings", чтобы расход RAG был виден отдельной
    # строкой в статистике, а не терялся молча (см. CLAUDE.md/архитектурный разбор, пункт 5).
    "openai-embeddings": (ai_rag.EMBEDDING_PRICE_PER_1M, 0.0),
}

def record_ai_cost(usage: dict) -> None:
    """Копит агрегированную стоимость AI-запросов — не пишет по записи на каждый запрос
    (раздуло бы stats.json), только бегущие суммы. Нужно, чтобы реально видеть эффект
    любых будущих оптимизаций (короткие ответы, кэш и т.д.), а не гадать на глаз.
    usage["provider"] ("openai"/"grok"/"gemini", по умолчанию "openai") выбирает прайс — у
    каждого провайдера своя цена за токен, общие totals остаются суммой по всем провайдерам, а
    by_provider хранит разбивку, чтобы в статистике было видно расход каждого провайдера
    отдельно от OpenAI."""
    input_tokens = usage.get("input_tokens", 0)
    output_tokens = usage.get("output_tokens", 0)
    provider = usage.get("provider", "openai")
    price_in, price_out = _AI_PROVIDER_PRICES.get(provider, _AI_PROVIDER_PRICES["openai"])
    cost = input_tokens * price_in / 1_000_000 + output_tokens * price_out / 1_000_000
    totals = stats["ai_cost_totals"]
    totals["requests"] += 1
    totals["input_tokens"] += input_tokens
    totals["output_tokens"] += output_tokens
    totals["cost_usd"] += cost
    by_provider = totals.setdefault("by_provider", {})
    p = by_provider.setdefault(provider, {"requests": 0, "input_tokens": 0, "output_tokens": 0, "cost_usd": 0.0})
    p["requests"] += 1
    p["input_tokens"] += input_tokens
    p["output_tokens"] += output_tokens
    p["cost_usd"] += cost
    _update_ai_cost_windows(cost)
    save_stats()
    if ai_circuit_breaker_tripped() and not stats["ai_cost_windows"].get("breaker_alerted"):
        try:
            asyncio.get_running_loop().create_task(_maybe_alert_admins_of_breaker_trip())
        except RuntimeError:
            pass  # вызвано вне event loop (например, из синхронного теста) — алерт просто не уйдёт

def get_ai_cost_stats_block() -> str:
    totals = stats["ai_cost_totals"]
    requests = totals["requests"]
    avg = totals["cost_usd"] / requests if requests else 0.0
    block = (
        f"\n🤖 <b>VMedA AI</b>\n"
        f"Запросов: <b>{requests}</b> (вход {totals['input_tokens']:,} / выход {totals['output_tokens']:,} ток.)\n"
        f"Расход: <b>${totals['cost_usd']:.4f}</b>, в среднем ${avg:.5f}/запрос"
    )
    by_provider = totals.get("by_provider", {})
    for provider, label in (("grok", "Grok"), ("gemini", "Gemini"), ("openai-embeddings", "RAG-эмбеддинги")):
        p = by_provider.get(provider)
        if p and p["requests"]:
            block += f"\n  из них {label}: {p['requests']} запр., ${p['cost_usd']:.4f}"
    return block.replace(",", " ")

# ---- Ценовой автовыключатель (circuit breaker) — на случай, если из-за какого-то бага (не
# concurrency-гонки выше, а, например, зацикленного клиента или скомпрометированного ключа) реально
# улетают деньги: MAX_AI_CONCURRENT_REQUESTS ограничивает только ОДНОВРЕМЕННЫЕ запросы, но не
# суммарный расход за час/сутки. Окна те же, что и у ai_used_monthly (см. CLAUDE.md) — ключ-строка
# периода + бегущая сумма, сравниваемая с порогом при каждой записи стоимости. ----
AI_COST_HOUR_LIMIT_USD = float(os.environ.get("AI_COST_HOUR_LIMIT_USD", "5.0"))
AI_COST_DAY_LIMIT_USD = float(os.environ.get("AI_COST_DAY_LIMIT_USD", "30.0"))

def _current_hour_key() -> str:
    return local_now().strftime("%Y-%m-%d-%H")

def _current_day_key() -> str:
    return local_today().isoformat()

def _update_ai_cost_windows(cost: float) -> None:
    """Вызывается из record_ai_cost() на КАЖДУЮ записанную стоимость (включая эмбеддинги) —
    копит расход в скользящих часовом/суточном окнах и взводит breaker_tripped, если порог
    превышен. Не сбрасывает breaker_tripped автоматически при смене окна — снятие блокировки нарочно
    осталось ручным (админ должен увидеть алерт и осознанно решить, что дело не в баге, а не
    полагаться на то, что через час всё само откроется и никто не заметит, что лимит вообще
    срабатывал)."""
    windows = stats["ai_cost_windows"]
    hour_key = _current_hour_key()
    if windows.get("hour_key") != hour_key:
        windows["hour_key"] = hour_key
        windows["hour_cost_usd"] = 0.0
    windows["hour_cost_usd"] += cost
    day_key = _current_day_key()
    if windows.get("day_key") != day_key:
        windows["day_key"] = day_key
        windows["day_cost_usd"] = 0.0
    windows["day_cost_usd"] += cost
    if windows["hour_cost_usd"] >= AI_COST_HOUR_LIMIT_USD or windows["day_cost_usd"] >= AI_COST_DAY_LIMIT_USD:
        windows["breaker_tripped"] = True

def ai_circuit_breaker_tripped() -> bool:
    return stats["ai_cost_windows"].get("breaker_tripped", False)

def reset_ai_circuit_breaker() -> None:
    """Ручной сброс из админ-панели — единственный способ разблокировать AI после срабатывания."""
    windows = stats["ai_cost_windows"]
    windows["breaker_tripped"] = False
    windows["breaker_alerted"] = False
    save_stats()

async def _maybe_alert_admins_of_breaker_trip() -> None:
    """Шлёт админам алерт РОВНО один раз за срабатывание (breaker_alerted — не per-заблокированный
    запрос, иначе при живом трафике админов завалило бы одинаковыми сообщениями)."""
    windows = stats["ai_cost_windows"]
    if windows.get("breaker_alerted"):
        return
    windows["breaker_alerted"] = True
    save_stats()
    text = (
        "🚨 <b>AI-автовыключатель сработал</b>\n\n"
        f"Расход за текущий час: ${windows['hour_cost_usd']:.2f} (лимит ${AI_COST_HOUR_LIMIT_USD:.2f})\n"
        f"Расход за сутки: ${windows['day_cost_usd']:.2f} (лимит ${AI_COST_DAY_LIMIT_USD:.2f})\n\n"
        "AI-режим временно отключён для всех пользователей. Проверь расход в статистике и сбрось "
        "выключатель из админ-панели, если это ожидаемый трафик, а не баг."
    )
    for admin_id in ADMIN_IDS:
        try:
            await bot.send_message(admin_id, text, parse_mode="HTML")
        except Exception:
            logger.exception("Не удалось уведомить админа %s о срабатывании AI-автовыключателя", admin_id)

def is_ai_session_active(user_id: int) -> bool:
    session = AI_SESSIONS.get(user_id)
    return bool(session) and time.time() - session["last_active"] < AI_SESSION_TIMEOUT_SECONDS

def start_ai_session(user_id: int) -> None:
    AI_SESSIONS[user_id] = {
        "task": None, "messages": [], "rag_context": None, "bucket": None,
        "quick_answer": None, "last_active": time.time(), "processing": False,
    }

def end_ai_session(user_id: int) -> None:
    AI_SESSIONS.pop(user_id, None)
    # AI_USER_LOCKS иначе растёт вечно — один Lock на каждого КОГДА-ЛИБО пользовавшегося AI
    # user_id, и ничего никогда его не убирало. Убираем только НЕзаблокированный lock — если он
    # прямо сейчас держится реальным запросом (async with lock: где-то в стеке), .locked() вернёт
    # True и мы его не трогаем: pop() из словаря не разрывает уже существующую ссылку на объект
    # у держащей его корутины, но create нового Lock под тем же user_id, пока старый ещё жив,
    # завёл бы ДВА разных объекта на одного пользователя — тогда второй параллельный запрос мог
    # бы получить "свежий" незалоченный lock и проскочить мимо защиты, которую AI_USER_LOCKS и
    # существует. Проверка и pop синхронны (без await между ними), так что гонки здесь нет.
    lock = AI_USER_LOCKS.get(user_id)
    if lock is not None and not lock.locked():
        AI_USER_LOCKS.pop(user_id, None)

def record_ai_attempts_cost(attempts_log: list) -> None:
    """Учитывает стоимость КАЖДОЙ попытки провайдера из attempts_log (см. ai.router.try_providers)
    — не только финально успешной. "refused" (контент-фильтр) реально тратит токены и должен
    попадать в себестоимость; "failed" (сетевая/API-ошибка) всегда несёт нулевой usage, так что
    его учёт — безопасный no-op. Вызывать вместо record_ai_cost(usage) на одну финальную попытку —
    иначе стоимость отказавших/сорвавшихся попыток при переключении на резервного провайдера
    нигде не фиксировалась (см. CLAUDE.md/архитектурный разбор AI-режима, пункт 9)."""
    for attempt in attempts_log:
        if attempt["status"] in ("success", "refused"):
            usage = dict(attempt["usage"])
            usage["provider"] = attempt["provider"]
            record_ai_cost(usage)

# ---- Кэш точных совпадений с модерацией (см. CLAUDE.md/архитектурный разбор AI-режима, пункт 6)
# ----
# Ключ — TaskRepresentation.fingerprint() (см. ai/task.py): нормализованный текст вопроса + данные
# условия, порядок слов не важен, разные исходные числа — разный fingerprint. Свежесгенерированный
# ответ НИКОГДА не раздаётся другим пользователям автоматически ("Модерация перед раздачей" —
# явный выбор архитектуры, а не auto-trust-with-first-answer: одна неверная генерация не должна
# разойтись по всем студентам, задавшим тот же вопрос, прежде чем админ её увидит) — см.
# get_cached_ai_answer/submit_ai_answer_for_moderation ниже и админ-очередь модерации в
# handlers/admin.py.
def get_cached_ai_answer(fingerprint: str) -> str | None:
    entry = stats["ai_answer_cache"].get(fingerprint)
    if not entry or entry.get("status") != "approved":
        return None
    entry["hits"] = entry.get("hits", 0) + 1
    save_stats()
    return entry["answer"]

# ---- Pre-parse кэш по СЫРОМУ тексту вопроса (только для текстовых сообщений — для фото
# fingerprint в принципе недостижим без распознавания) ----
# get_cached_ai_answer выше проверяется уже ПОСЛЕ vision-парсера — а значит "cache hit ничего не
# стоит" было верно только для RAG/эмбеддингов и solver'а, но не для самого parse_task(): даже при
# полном попадании в кэш этот вызов уже произошёл и уже оплачен. stats["ai_raw_text_aliases"]
# хранит {fingerprint СЫРОГО текста -> fingerprint РАЗОБРАННОГО задания} — если конкретно ЭТОТ
# сырой текст уже когда-то встречался и его разобранная версия попала в одобренный кэш, повторный
# буквальный ввод того же текста отдаётся вообще без единого обращения к модели (см.
# get_raw_text_precache_answer, вызывается из handle_ai_text_input ДО ai_vision_parser.parse_task).
# Не претендует на то же покрытие, что и обычный fingerprint (парсер может переформулировать текст
# и вытащить values/units, каких у необработанного текста ещё нет — значит, разные ФОРМУЛИРОВКИ
# одного и того же вопроса тут не совпадут, только буквально идентичный повторный ввод), но это
# именно то, что реально накапливается на популярных вопросах из общей базы билетов/тестов, когда
# разные студенты копируют один и тот же текст вопроса.
def get_raw_text_precache_answer(text: str) -> tuple[str, str] | None:
    """(answer, user_turn_content) при полном попадании — без единого обращения к модели, либо
    None (обычный путь: парсинг как раньше)."""
    raw_fingerprint = TaskRepresentation(raw_text=text).fingerprint()
    parsed_fingerprint = stats["ai_raw_text_aliases"].get(raw_fingerprint)
    if not parsed_fingerprint:
        return None
    cached_answer = get_cached_ai_answer(parsed_fingerprint)
    if cached_answer is None:
        return None
    text_part = TaskRepresentation(raw_text=text).to_prompt_text()
    text_part = (text_part + ai_prompts.QUICK_SUFFIX) if text_part else ai_prompts.QUICK_SUFFIX.strip()
    return cached_answer, text_part

def record_raw_text_alias(text: str, task) -> None:
    """Запоминает, во что разобрался ДАННЫЙ сырой текст — вызывается на каждом первом сообщении
    сессии (независимо от того, оказался ли РАЗОБРАННЫЙ fingerprint кэш-хитом или свежей
    генерацией), чтобы alias был готов к моменту, когда исходный кандидат станет approved.
    Перезаписывает существующий alias тем же сырым текстом — парсер не гарантированно
    детерминирован (OCR/переформулировка), так что один и тот же сырой текст теоретически может
    время от времени давать разные fingerprint'ы; в этом случае alias "плавает" на последнюю
    версию, а не залипает на первую — не идеально, но не идёт вразрез с корректностью (промах
    pre-cache просто скатывается на обычный путь с парсингом)."""
    raw_fingerprint = TaskRepresentation(raw_text=text).fingerprint()
    parsed_fingerprint = task.fingerprint()
    if stats["ai_raw_text_aliases"].get(raw_fingerprint) != parsed_fingerprint:
        stats["ai_raw_text_aliases"][raw_fingerprint] = parsed_fingerprint
        save_stats()

def submit_ai_answer_for_moderation(
    task, answer: str, confidence_action: str = ai_confidence.SERVE, confidence_reasons: list = None,
) -> None:
    """Не перезаписывает уже ОДОБРЕННУЮ запись тем же fingerprint'ом — она остаётся источником
    истины, пока админ не отклонит её отдельным действием. "pending"/"rejected" запись
    обновляется новым кандидатом — отклонённый вопрос не блокируется навсегда, у него будет шанс
    на новый (возможно, более удачный) сгенерированный ответ при следующем обращении.

    confidence_action/confidence_reasons — вывод ai.confidence.decide() (см. get_first_message_ai_answer
    ниже): не влияет на то, попадёт ли запись в очередь, только на порядок (см.
    get_next_pending_ai_cache_entry) и на то, что увидит админ на экране модерации."""
    fingerprint = task.fingerprint()
    existing = stats["ai_answer_cache"].get(fingerprint)
    if existing and existing.get("status") == "approved":
        return
    stats["ai_answer_cache"][fingerprint] = {
        "question_preview": task.question_text()[:300],
        "answer": answer,
        "subject": task.subject,
        "status": "pending",
        "created_at": time.time(),
        "hits": existing.get("hits", 0) if existing else 0,
        "confidence_action": confidence_action,
        "confidence_reasons": confidence_reasons or [],
    }
    save_stats()

def get_pending_ai_cache_count() -> int:
    return sum(1 for e in stats["ai_answer_cache"].values() if e.get("status") == "pending")

_AI_CACHE_PRIORITY = {ai_confidence.ESCALATE: 0, ai_confidence.VERIFY: 1, ai_confidence.SERVE: 2}

def get_next_pending_ai_cache_entry():
    """(fingerprint, entry) следующей записи на модерацию, либо (None, None), если очередь пуста.
    Не строго FIFO — сначала записи с более тревожным confidence_action (см. ai/confidence.py:
    ESCALATE, затем VERIFY, затем SERVE — сегодня у quick-ответа нет более сильного провайдера,
    на который можно было бы реально переключиться при низкой уверенности, поэтому "эскалация"
    здесь означает приоритет в очереди модерации, а не скрытый повторный запрос к модели), внутри
    одного уровня приоритета — старейшая первая."""
    pending = [(fp, e) for fp, e in stats["ai_answer_cache"].items() if e.get("status") == "pending"]
    if not pending:
        return None, None
    pending.sort(key=lambda item: (
        _AI_CACHE_PRIORITY.get(item[1].get("confidence_action", ai_confidence.SERVE), 2),
        item[1].get("created_at", 0),
    ))
    return pending[0]

def moderate_ai_cache_entry(fingerprint: str, approve: bool) -> bool:
    entry = stats["ai_answer_cache"].get(fingerprint)
    if not entry:
        return False
    entry["status"] = "approved" if approve else "rejected"
    save_stats()
    return True

AI_LOW_CONFIDENCE_NOTE = (
    "\n\n⚠️ Этот ответ не прошёл автоматическую проверку на согласованность с заданием — "
    "обязательно сверь его с материалами курса, прежде чем полагаться на него."
)

async def ensure_rag_context(session: dict) -> str:
    """Лениво считает session["rag_context"], если он ещё не считался (None — сентинел "не
    считали", в отличие от "" — "считали, снипетов не нашлось"). Первый ход сессии, попавший в
    кэш точных совпадений (см. get_first_message_ai_answer), НЕ платит за RAG вообще — но если
    пользователь потом жмёт "Показать решение по шагам", этому второму, уже небесплатному
    запросу контекст всё-таки нужен, иначе подробный разбор проиграет в качестве без видимой
    причины. Безопасно вызывать повторно — второй вызов на той же сессии не платит снова."""
    if session.get("rag_context") is not None:
        return session["rag_context"]
    if session.get("task") is None:
        return ""
    snippets, rag_usage = await ai_rag.search_for_task(session["task"])
    if rag_usage.get("input_tokens"):
        record_ai_cost({**rag_usage, "provider": "openai-embeddings"})
    session["rag_context"] = ai_rag.format_context(snippets)
    return session["rag_context"]

async def get_first_message_ai_answer(user_id: int, session: dict, task) -> tuple:
    """Первый ход AI-сессии: сначала проверяем кэш точных совпадений (только ОДОБРЕННЫЕ записи) —
    при попадании ответ отдаётся бесплатно, без обращения к модели, без RAG/эмбеддингов и без
    списания квоты/учёта стоимости (см. пункт 4 архитектурного разбора: раньше RAG считался ДО
    проверки кэша, и даже кэш-хит платил за эмбеддинги). При промахе — сначала считаем RAG-контекст
    (ensure_rag_context, тут уже реально нужен модели), затем обычный запрос к solve_ai_request
    (списывает квоту, учитывает стоимость всех попыток), затем ответ прогоняется через
    детерминированный валидатор (см. ai/validator.py), для calculation-заданий — ещё и через
    независимый математический verifier (ai/math_verifier.py, пересчитывает результат по
    распознанной формуле и сверяет с ответом), для mcq-заданий — через сверку с эталонной базой
    (ai/mcq_verifier.py, ai/reference_bank.py: 1040 вопросов теста кафедры анатомии с объективно
    известным правильным вариантом), и через confidence-роутер (ai/confidence.py) — при низкой
    уверенности пользователю честно показывается предупреждение (AI_LOW_CONFIDENCE_NOTE), а запись
    в очереди модерации получает более высокий приоритет на проверку. session["quick_answer"]
    всегда хранит ИСХОДНЫЙ ответ БЕЗ предупреждения — это canonical-якорь для "Показать решение по
    шагам" (ai.prompts.explain_followup_text), предупреждение не должно путать модель на следующем
    ходу.

    Возвращает (answer, user_turn) — user_turn в обоих случаях в том же формате, что и обычный
    ход solve(), включая суффикс краткого ответа (ai_prompts.QUICK_SUFFIX), чтобы история сессии
    не отличалась по форме от того, что получилось бы при реальном обращении к модели."""
    session.setdefault("task", task)  # ensure_rag_context reads session["task"] — обычные хендлеры
    # уже выставляют его сами до вызова этой функции, но полагаться на внешнюю дисциплину незачем
    cached_answer = get_cached_ai_answer(task.fingerprint())
    if cached_answer is not None:
        session["quick_answer"] = cached_answer
        text_part = task.to_prompt_text()
        text_part = (text_part + ai_prompts.QUICK_SUFFIX) if text_part else ai_prompts.QUICK_SUFFIX.strip()
        return cached_answer, {"role": "user", "content": text_part}

    rag_context = await ensure_rag_context(session)
    answer, user_turn, usage, attempts_log = await solve_ai_request(
        task=task, history=session["messages"], quick=True,
        bucket=session["bucket"], rag_context=rag_context,
    )
    increment_ai_usage(user_id)
    record_ai_attempts_cost(attempts_log)

    validation = ai_validator.validate_answer(task, answer)
    math_verification = ai_math_verifier.verify_calculation(task, answer)
    mcq_verification = ai_mcq_verifier.verify_mcq(task, answer)
    decision = ai_confidence.decide(
        task, validation, rag_grounded=bool(session.get("rag_context")),
        math_verification=math_verification, mcq_verification=mcq_verification,
    )
    submit_ai_answer_for_moderation(task, answer, decision.action, decision.reasons)
    session["quick_answer"] = answer

    display_answer = answer + AI_LOW_CONFIDENCE_NOTE if decision.action != ai_confidence.SERVE else answer
    return display_answer, user_turn

def get_ai_menu_text(user_id: int) -> str:
    availability = "" if ai_provider_available() else "\n\n🔧 Идут финальные настройки — совсем скоро запустим."
    return (
        f"🤖 <b>VMedA AI</b>\n{DIVIDER}\n\n"
        "AI-помощник, который разбирает задание по фото или тексту и сразу выдаёт решение: "
        "чёткий ответ и объяснение по шагам. Работает по биологии, физике, химии, анатомии и оперативной хирургии — "
        "тесты, билеты, контрольные, летучки. Просто присылаешь фото — получаешь разбор.\n\n"
        f"Бесплатных запросов сегодня: <b>{get_ai_quota_label(user_id)}</b>"
        f"{availability}"
    )

def get_ai_menu_keyboard():
    builder = InlineKeyboardBuilder()
    builder.button(text="📷 Отправить фото или текст задания", callback_data="ai_solve_start")
    builder.button(text="🔙 Назад в меню", callback_data="back_to_main")
    builder.adjust(1)
    return builder.as_markup()

def get_ai_waiting_keyboard():
    builder = InlineKeyboardBuilder()
    builder.button(text="🛑 Отмена", callback_data="ai_solve_cancel")
    builder.adjust(1)
    return builder.as_markup()

def get_ai_result_chunks(answer: str, user_id: int, session_active: bool, offer_explanation: bool = False) -> list:
    """Делит длинный ответ на несколько сообщений вместо физической обрезки (см.
    ai.service.split_answer_into_chunks) — заголовок только в первом куске, футер (напоминание
    сверяться с курсом + остаток квоты + подсказка про диалог) только в последнем. HTML-разметка
    (ai.service.format_answer_html) применяется к КАЖДОМУ куску сырого текста уже ПОСЛЕ
    разбиения — иначе резать уже готовый HTML рисковало бы разорвать тег пополам, а Telegram
    целиком отклоняет сообщение с несбалансированной разметкой."""
    if offer_explanation:
        continuation = "\n\n🧠 Это краткий ответ — нажми кнопку ниже, если нужно решение по шагам."
    elif session_active:
        continuation = "\n\n💬 Можешь сразу уточнить вопрос по этой же теме — я помню контекст диалога."
    else:
        continuation = ""
    footer = (
        f"\n\n💡 Сверяй важные ответы с материалами курса.\n"
        f"Осталось бесплатных запросов сегодня: {get_ai_quota_label(user_id)}"
        f"{continuation}"
    )
    raw_chunks = ai_service.split_answer_into_chunks(answer)
    last = len(raw_chunks) - 1
    texts = []
    for i, raw_chunk in enumerate(raw_chunks):
        text = ai_service.format_answer_html(raw_chunk)
        if i == 0:
            text = f"🤖 <b>Ответ AI</b>\n{DIVIDER}\n\n{text}"
        if i == last:
            text = f"{text}{footer}"
        texts.append(text)
    return texts

def get_ai_result_keyboard(session_active: bool, offer_explanation: bool = False):
    builder = InlineKeyboardBuilder()
    if offer_explanation and session_active:
        builder.button(text="🧠 Показать решение по шагам", callback_data="ai_show_explanation")
    if session_active:
        builder.button(text="🔚 Закончить диалог", callback_data="ai_session_end")
    else:
        builder.button(text="🔙 Назад в меню", callback_data="ai_menu")
    builder.adjust(1)
    return builder.as_markup()

async def send_ai_result(thinking, answer: str, user_id: int, session_active: bool, offer_explanation: bool = False) -> None:
    """Общая точка отправки ответа AI для всех трёх хендлеров (фото/текст/"показать решение") —
    редактирует "thinking"-заглушку первым куском, остальные куски (если ответ длинный) уходят
    отдельными сообщениями; клавиатура с действиями всегда только на последнем сообщении."""
    chunks = get_ai_result_chunks(answer, user_id, session_active, offer_explanation)
    keyboard = get_ai_result_keyboard(session_active, offer_explanation)
    await safe_edit_text(thinking, chunks[0], parse_mode="HTML", reply_markup=keyboard if len(chunks) == 1 else None)
    for chunk in chunks[1:-1]:
        await thinking.answer(chunk, parse_mode="HTML")
    if len(chunks) > 1:
        await thinking.answer(chunks[-1], parse_mode="HTML", reply_markup=keyboard)

@dp.callback_query(F.data == "ai_menu")
async def cb_ai_menu(callback: CallbackQuery):
    await callback.answer()
    end_ai_session(callback.from_user.id)
    await safe_edit_text(
        callback.message,
        get_ai_menu_text(callback.from_user.id),
        parse_mode="HTML",
        reply_markup=get_ai_menu_keyboard()
    )

@dp.callback_query(F.data == "ai_solve_start")
async def cb_ai_solve_start(callback: CallbackQuery):
    user_id = callback.from_user.id
    if not ai_provider_available():
        await callback.answer("AI сейчас на техническом обслуживании, загляни позже.", show_alert=True)
        return
    if ai_circuit_breaker_tripped():
        await callback.answer(
            "AI временно отключён из-за высокой нагрузки — администраторы уже знают, попробуй позже.",
            show_alert=True,
        )
        return
    if not ai_quota_ok(user_id):
        await callback.answer("На сегодня бесплатные AI-запросы закончились, попробуй завтра.", show_alert=True)
        return
    await callback.answer()
    start_ai_session(user_id)
    await safe_edit_text(
        callback.message,
        f"📷 <b>Жду задание</b>\n{DIVIDER}\n\nПришли фото задания или напиши его текстом одним сообщением. "
        "Дальше можно будет уточнять вопросы по теме — контекст диалога сохранится.",
        parse_mode="HTML",
        reply_markup=get_ai_waiting_keyboard()
    )

@dp.callback_query(F.data == "ai_solve_cancel")
async def cb_ai_solve_cancel(callback: CallbackQuery):
    end_ai_session(callback.from_user.id)
    await callback.answer()
    await safe_edit_text(
        callback.message,
        get_ai_menu_text(callback.from_user.id),
        parse_mode="HTML",
        reply_markup=get_ai_menu_keyboard()
    )

@dp.callback_query(F.data == "ai_session_end")
async def cb_ai_session_end(callback: CallbackQuery):
    end_ai_session(callback.from_user.id)
    await callback.answer("Диалог закончен")
    await safe_edit_text(
        callback.message,
        get_ai_menu_text(callback.from_user.id),
        parse_mode="HTML",
        reply_markup=get_ai_menu_keyboard()
    )

@dp.callback_query(F.data == "ai_show_explanation")
async def cb_ai_show_explanation(callback: CallbackQuery):
    """Второй, отдельный запрос — только по явному нажатию, только если после короткого
    ответа ещё остался дневной лимит. Большинство тестовых вопросов на этом и заканчиваются,
    не оплачивая длинный пошаговый разбор, который часто и не нужен."""
    user_id = callback.from_user.id
    if not is_ai_session_active(user_id):
        await callback.answer("Диалог уже закрыт, начни заново через меню AI.", show_alert=True)
        return
    session = AI_SESSIONS[user_id]
    if session["processing"]:
        await callback.answer()
        return
    lock = _get_ai_user_lock(user_id)
    if lock.locked():
        # другой запрос этого пользователя уже в работе — возможно, из другой (уже заменившей эту)
        # AI-сессии, session["processing"] выше такую гонку не ловит (см. AI_USER_LOCKS)
        await callback.answer()
        return
    if ai_circuit_breaker_tripped():
        await callback.answer(
            "AI временно отключён из-за высокой нагрузки — администраторы уже знают, попробуй позже.",
            show_alert=True,
        )
        return
    if not ai_quota_ok(user_id):
        end_ai_session(user_id)
        await callback.answer("На сегодня бесплатные AI-запросы закончились, попробуй завтра.", show_alert=True)
        return
    if not AI_CONCURRENCY_GATE.try_acquire():
        # проверяем последним, ПОСЛЕ всех остальных причин отказа — слот резервируется только
        # если запрос реально пойдёт в работу, иначе release() в finally было бы не с чем парить
        await callback.answer("Сейчас слишком много запросов к AI одновременно — попробуй через минуту.", show_alert=True)
        return
    await callback.answer()
    async with lock:
        session["processing"] = True
        thinking = await callback.message.answer("🤖 Готовлю решение по шагам...")
        task_type = session["task"].type if session.get("task") is not None else None
        followup_text = ai_prompts.explain_followup_text(session.get("quick_answer") or "", task_type)
        try:
            rag_context = await ensure_rag_context(session)
            answer, user_turn, usage, attempts_log = await solve_ai_request(
                text=followup_text, history=session["messages"], quick=False,
                bucket=session.get("bucket"), rag_context=rag_context,
            )
            increment_ai_usage(user_id)
            record_ai_attempts_cost(attempts_log)
            session["messages"].append(user_turn)
            session["messages"].append({"role": "assistant", "content": answer})
            session["last_active"] = time.time()
            session_active = ai_quota_ok(user_id)
            if not session_active:
                end_ai_session(user_id)
            await send_ai_result(thinking, answer, user_id, session_active)
        except AIRefusalError as exc:
            logger.warning("AI отказался дать подробный разбор пользователю %s", user_id)
            record_ai_attempts_cost(getattr(exc, "ai_attempts_log", []))
            await safe_edit_text(
                thinking,
                "⚠️ AI отказался отвечать на этот конкретный вопрос — похоже, сработал фильтр "
                "содержимого провайдера (так бывает на некоторых медицинских формулировках). "
                "Эта попытка не списана с дневного лимита — попробуй переформулировать вопрос."
            )
        except Exception as exc:
            logger.exception("Ошибка при получении подробного решения для пользователя %s", user_id)
            record_ai_attempts_cost(getattr(exc, "ai_attempts_log", []))
            await safe_edit_text(thinking, "⚠️ Не удалось получить решение. Попробуй ещё раз позже.")
        finally:
            AI_CONCURRENCY_GATE.release()
            if user_id in AI_SESSIONS:
                AI_SESSIONS[user_id]["processing"] = False

@dp.message(F.photo)
async def handle_ai_photo_input(message: Message):
    """Единственный обработчик входящих фото в боте — фото вне AI-режима бот
    никогда раньше не принимал, поэтому конфликтов с другими хендлерами нет."""
    user_id = message.from_user.id
    if not is_ai_session_active(user_id):
        return
    session = AI_SESSIONS[user_id]
    if session["processing"]:
        return  # предыдущее сообщение этого диалога ещё обрабатывается — вероятный случайный дубль
    lock = _get_ai_user_lock(user_id)
    if lock.locked():
        # запрос этого пользователя уже в работе (возможно, из другой, уже заменившей эту, сессии)
        return
    if ai_circuit_breaker_tripped():
        await message.answer("AI временно отключён из-за высокой нагрузки — администраторы уже знают, попробуй позже.")
        return
    if not ai_quota_ok(user_id):
        end_ai_session(user_id)
        await message.answer("На сегодня бесплатные AI-запросы закончились, попробуй завтра.")
        return
    if not AI_CONCURRENCY_GATE.try_acquire():
        # проверяем последним, ПОСЛЕ всех остальных причин отказа — слот резервируется только
        # если запрос реально пойдёт в работу, иначе release() в finally было бы не с чем парить
        await message.answer("Сейчас слишком много запросов к AI одновременно — попробуй через минуту.")
        return
    is_first = session["task"] is None  # самое первое сообщение сессии — сперва только краткий ответ
    async with lock:
        session["processing"] = True
        thinking = await message.answer("🤖 Разбираю задание, подожди немного...")
        try:
            photo = message.photo[-1]
            tg_file = await bot.get_file(photo.file_id)
            buf = await bot.download_file(tg_file.file_path)
            task_repr, parse_usage = await ai_vision_parser.parse_task(image_bytes=resize_image_for_ai(buf.read()))
            if parse_usage.get("input_tokens") or parse_usage.get("output_tokens"):
                record_ai_cost(parse_usage)
            if is_first:
                session["task"] = task_repr
                session["bucket"] = ai_router.route_bucket(task_repr)
                answer, user_turn = await get_first_message_ai_answer(user_id, session, task_repr)
            else:
                rag_context = await ensure_rag_context(session)
                answer, user_turn, usage, attempts_log = await solve_ai_request(
                    text=task_repr.to_prompt_text(), history=session["messages"], quick=False,
                    bucket=session.get("bucket"), rag_context=rag_context,
                )
                increment_ai_usage(user_id)
                record_ai_attempts_cost(attempts_log)
            session["messages"].append(user_turn)
            session["messages"].append({"role": "assistant", "content": answer})
            session["last_active"] = time.time()
            session_active = ai_quota_ok(user_id)
            if not session_active:
                end_ai_session(user_id)
            await send_ai_result(thinking, answer, user_id, session_active, offer_explanation=is_first)
        except AIRefusalError as exc:
            logger.warning("AI отказался разобрать фото от пользователя %s", user_id)
            record_ai_attempts_cost(getattr(exc, "ai_attempts_log", []))
            await safe_edit_text(
                thinking,
                "⚠️ AI отказался отвечать на это фото — похоже, сработал фильтр содержимого "
                "провайдера (так бывает на некоторых медицинских формулировках). Эта попытка не "
                "списана с дневного лимита — попробуй прислать вопрос текстом или переформулировать."
            )
        except Exception as exc:
            logger.exception("Ошибка при обработке AI-фото от пользователя %s", user_id)
            record_ai_attempts_cost(getattr(exc, "ai_attempts_log", []))
            await safe_edit_text(thinking, "⚠️ Не удалось обработать фото. Попробуй ещё раз или пришли текстом.")
        finally:
            AI_CONCURRENCY_GATE.release()
            if user_id in AI_SESSIONS:
                AI_SESSIONS[user_id]["processing"] = False

@dp.message(F.text)
async def handle_ai_text_input(message: Message):
    user_id = message.from_user.id
    if not is_ai_session_active(user_id):
        raise SkipHandler  # не AI-режим — пусть сообщение обработают остальные текстовые хендлеры
    if message.text.startswith("/"):
        raise SkipHandler
    session = AI_SESSIONS[user_id]
    if session["processing"]:
        raise SkipHandler  # предыдущее сообщение этого диалога ещё обрабатывается — вероятный случайный дубль
    lock = _get_ai_user_lock(user_id)
    if lock.locked():
        # запрос этого пользователя уже в работе (возможно, из другой, уже заменившей эту, сессии)
        raise SkipHandler
    if ai_circuit_breaker_tripped():
        await message.answer("AI временно отключён из-за высокой нагрузки — администраторы уже знают, попробуй позже.")
        return
    if not ai_quota_ok(user_id):
        end_ai_session(user_id)
        await message.answer("На сегодня бесплатные AI-запросы закончились, попробуй завтра.")
        return
    if not AI_CONCURRENCY_GATE.try_acquire():
        # проверяем последним, ПОСЛЕ всех остальных причин отказа — слот резервируется только
        # если запрос реально пойдёт в работу, иначе release() в finally было бы не с чем парить
        await message.answer("Сейчас слишком много запросов к AI одновременно — попробуй через минуту.")
        return
    is_first = session["task"] is None  # самое первое сообщение сессии — сперва только краткий ответ
    async with lock:
        session["processing"] = True
        thinking = await message.answer("🤖 Разбираю задание, подожди немного...")
        try:
            if is_first:
                precache = get_raw_text_precache_answer(message.text)
                if precache is not None:
                    # буквально тот же сырой текст уже встречался и разобрался в одобренный кэш —
                    # ни vision-парсер, ни RAG, ни solver не трогаются вообще (см. комментарий у
                    # get_raw_text_precache_answer)
                    cached_answer, text_part = precache
                    session["task"] = TaskRepresentation(raw_text=message.text)
                    session["bucket"] = ai_router.route_bucket(session["task"])
                    session["quick_answer"] = cached_answer
                    answer, user_turn = cached_answer, {"role": "user", "content": text_part}
                else:
                    task_repr, parse_usage = await ai_vision_parser.parse_task(text=message.text)
                    if parse_usage.get("input_tokens") or parse_usage.get("output_tokens"):
                        record_ai_cost(parse_usage)
                    session["task"] = task_repr
                    session["bucket"] = ai_router.route_bucket(task_repr)
                    answer, user_turn = await get_first_message_ai_answer(user_id, session, task_repr)
                    record_raw_text_alias(message.text, task_repr)
            else:
                rag_context = await ensure_rag_context(session)
                answer, user_turn, usage, attempts_log = await solve_ai_request(
                    text=message.text, history=session["messages"], quick=False,
                    bucket=session.get("bucket"), rag_context=rag_context,
                )
                increment_ai_usage(user_id)
                record_ai_attempts_cost(attempts_log)
            session["messages"].append(user_turn)
            session["messages"].append({"role": "assistant", "content": answer})
            session["last_active"] = time.time()
            session_active = ai_quota_ok(user_id)
            if not session_active:
                end_ai_session(user_id)
            await send_ai_result(thinking, answer, user_id, session_active, offer_explanation=is_first)
        except AIRefusalError as exc:
            logger.warning("AI отказался ответить на текст от пользователя %s", user_id)
            record_ai_attempts_cost(getattr(exc, "ai_attempts_log", []))
            await safe_edit_text(
                thinking,
                "⚠️ AI отказался отвечать на этот конкретный вопрос — похоже, сработал фильтр "
                "содержимого провайдера (так бывает на некоторых медицинских формулировках). "
                "Эта попытка не списана с дневного лимита — попробуй переформулировать вопрос."
            )
        except Exception as exc:
            logger.exception("Ошибка при обработке AI-текста от пользователя %s", user_id)
            record_ai_attempts_cost(getattr(exc, "ai_attempts_log", []))
            await safe_edit_text(thinking, "⚠️ Не удалось получить ответ от AI. Попробуй ещё раз позже.")
        finally:
            AI_CONCURRENCY_GATE.release()
            if user_id in AI_SESSIONS:
                AI_SESSIONS[user_id]["processing"] = False

# ==================== СКРЫТАЯ ФУНКЦИЯ (ВРЕМЕННО) ====================
# Если написать боту номер билета текстом (например "20А"), в чат придут все
# вопросы и ответы этого билета подряд, без кнопок. Без команд и упоминаний в меню.
@dp.message(F.text)
async def handle_hidden_ticket_dump(message: Message):
    ticket = TICKET_LOOKUP.get(_normalize_ticket_num(message.text))
    if not ticket:
        raise SkipHandler  # не билет — пусть попробует обработать поиск по словам
    ticket_num = ticket.get("num", "?")
    questions = ticket.get("questions", [])
    await message.answer(f"📘 <b>Билет {ticket_num}</b> — все ответы\n{DIVIDER}", parse_mode="HTML")
    for q in questions:
        q_num = q.get("num")
        body = f"❓ <b>Вопрос {q_num}</b>\n{DIVIDER}\n\n<b>{q.get('title', '')}</b>\n\n{q.get('answer', '')}"
        image_name = q.get("image")
        image_path = os.path.join(IMAGES_DIR, image_name) if image_name else None
        if image_path and os.path.exists(image_path):
            try:
                await message.answer_photo(FSInputFile(image_path))
            except Exception:
                logger.exception("Не удалось отправить изображение %s", image_path)
        await message.answer(body, parse_mode="HTML")

# ==================== ПОИСК ПО ОПЕРАТИВНОЙ ХИРУРГИИ (обработчик) ====================
# Должен стоять РАНЬШЕ handle_keyword_search ниже — тот не проверяет никакого pending-состояния
# и обрабатывает ЛЮБОЙ текст как биологический поиск, так что если зарегистрировать этот хендлер
# позже (например, через dp.include_router() в самом конце файла, как handlers/histology.py),
# handle_keyword_search перехватывал бы поисковый запрос по ОХ первым и он сюда никогда бы не
# доходил. Поэтому — как ADMIN_PENDING/ASSISTANT_PENDING — живёт прямо в telegram_bot.py, до
# handle_keyword_search по порядку в файле.
@dp.message(F.text)
async def handle_oh_search_query(message: Message):
    user_id = message.from_user.id
    if user_id not in OH_SEARCH_PENDING:
        raise SkipHandler
    if message.text.startswith("/"):
        raise SkipHandler
    OH_SEARCH_PENDING.discard(user_id)
    query = message.text.strip()
    safe_query = html.escape(query)
    topics, instruments, projections, stations = search_operative_surgery(query)
    builder = InlineKeyboardBuilder()
    if not topics and not instruments and not projections and not stations:
        builder.row(InlineKeyboardButton(text="🔙 К разделу", callback_data="oh:menu"))
        await message.answer(
            f"🔍 По запросу «{safe_query}» в разделе «Оперативная хирургия» ничего не найдено.",
            reply_markup=builder.as_markup()
        )
        return
    lines = [f"🔍 <b>Оперативная хирургия — результаты поиска:</b> «{safe_query}»\n{DIVIDER}"]
    if topics:
        lines.append("\n📚 <b>Темы:</b>")
        for t in topics:
            lines.append(f"• {t['number']}. {t['title']}")
            builder.button(text=f"{t['number']}. {t['title'][:40]}", callback_data=f"oh:topic:{t['id']}")
    if instruments:
        lines.append("\n🛠 <b>Инструменты:</b>")
        for group_name, name in instruments:
            lines.append(f"• {name} ({group_name})")
        builder.button(text="🛠 Открыть инструменты", callback_data="oh:instruments")
    if projections:
        lines.append("\n📍 <b>Проекции:</b>")
        for group_name, item in projections:
            lines.append(f"• {item['structure']}")
        builder.button(text="📍 Открыть проекции", callback_data="oh:projections")
    if stations:
        lines.append("\n🎓 <b>Практические станции:</b>")
        for group_name, name in stations:
            lines.append(f"• {name}")
        builder.button(text="🎓 Открыть станции", callback_data="oh:stations")
    builder.adjust(1)
    builder.row(InlineKeyboardButton(text="🔙 К разделу", callback_data="oh:menu"))
    await message.answer("\n".join(lines), parse_mode="HTML", reply_markup=builder.as_markup())

# ==================== ПОИСК ПО НОРМАЛЬНОЙ ФИЗИОЛОГИИ (обработчик) ====================
# Та же причина, что у handle_oh_search_query выше — должен стоять РАНЬШЕ handle_keyword_search,
# не может жить внутри handlers/physiology.py (dp.include_router() регистрируется в самом конце
# файла и попал бы в цепочку диспетчеризации ПОСЛЕ handle_keyword_search).
@dp.message(F.text)
async def handle_phys_search_query(message: Message):
    user_id = message.from_user.id
    if user_id not in PHYS_SEARCH_PENDING:
        raise SkipHandler
    if message.text.startswith("/"):
        raise SkipHandler
    PHYS_SEARCH_PENDING.discard(user_id)
    query = message.text.strip()
    safe_query = html.escape(query)
    results = physiology_handlers.search_physiology(query)
    builder = InlineKeyboardBuilder()
    if not results:
        builder.row(InlineKeyboardButton(text="🔙 К разделу", callback_data="phys:menu"))
        await message.answer(
            f"🔍 По запросу «{safe_query}» в разделе «Нормальная физиология» ничего не найдено.",
            reply_markup=builder.as_markup()
        )
        return
    lines = [f"🔍 <b>Нормальная физиология — результаты поиска:</b> «{safe_query}»\n{DIVIDER}\n"]
    for t in results:
        lines.append(f"• {t['order']}. {t['title']}")
        builder.button(text=f"{t['order']}. {t['short_title'][:40]}", callback_data=f"phys:topic:{t['topic_id']}")
    builder.adjust(1)
    builder.row(InlineKeyboardButton(text="🔙 К разделу", callback_data="phys:menu"))
    await message.answer("\n".join(lines), parse_mode="HTML", reply_markup=builder.as_markup())

# ==================== ПОИСК ПО КЛЮЧЕВЫМ СЛОВАМ (обработчик) ====================
@dp.message(F.text)
async def handle_keyword_search(message: Message):
    query = (message.text or "").strip()
    if not query or query.startswith("/"):
        return
    results = search_questions_by_keyword(query)
    safe_query = html.escape(query)
    if not results:
        await message.answer(
            f"🔍 По запросу «{safe_query}» ничего не найдено.\n"
            "Попробуй другое слово или загляни в раздел «📝 Вопросы»."
        )
        return
    suffix = f" (показаны первые {SEARCH_RESULTS_LIMIT})" if len(results) >= SEARCH_RESULTS_LIMIT else ""
    text = (
        f"🔍 <b>Результаты поиска:</b> «{safe_query}»\n{DIVIDER}\n\n"
        f"Найдено вопросов: {len(results)}{suffix}\nВыбери нужный:"
    )
    await message.answer(text, parse_mode="HTML", reply_markup=get_search_results_keyboard(results))

# ==================== ФИЗИКА (роутер/реэкспорт) ====================
# Хендлеры (тестовая часть, билеты, задачи — все с уникальными callback_data-фильтрами, безопасно
# для порядка dp) вынесены в handlers/physics.py (свой Router) — здесь только регистрация роутера
# и реэкспорт имён. Клавиатурные билдеры остаются здесь (используются и cb_menu_physics/
# download_physics_*, которые тоже остаются) — см. docstring handlers/physics.py.
from handlers import physics as physics_handlers  # noqa: E402 — mid-file by design, see above

dp.include_router(physics_handlers.router)

cb_physics_tickets = physics_handlers.cb_physics_tickets
cb_physics_theory_tickets = physics_handlers.cb_physics_theory_tickets
cb_phys_theory_ticket = physics_handlers.cb_phys_theory_ticket
cb_phys_theory_question = physics_handlers.cb_phys_theory_question
cb_physics_test_tickets = physics_handlers.cb_physics_test_tickets
cb_phys_test_ticket = physics_handlers.cb_phys_test_ticket
cb_phys_test_ticket_tasks = physics_handlers.cb_phys_test_ticket_tasks
cb_phys_test_ticket_task_show = physics_handlers.cb_phys_test_ticket_task_show
cb_physics_task_tickets = physics_handlers.cb_physics_task_tickets
cb_phys_task_ticket = physics_handlers.cb_phys_task_ticket
cb_phys_task_ticket_show = physics_handlers.cb_phys_task_ticket_show
cb_physics_test = physics_handlers.cb_physics_test
cb_physics_page = physics_handlers.cb_physics_page
cb_physics_question = physics_handlers.cb_physics_question
cb_physics_grade45 = physics_handlers.cb_physics_grade45
cb_physics_grade45_question = physics_handlers.cb_physics_grade45_question
cb_physics_extra = physics_handlers.cb_physics_extra
cb_physics_extra_question = physics_handlers.cb_physics_extra_question
cb_physics_tasks = physics_handlers.cb_physics_tasks
cb_phystask_topic = physics_handlers.cb_phystask_topic
cb_phystask_formulas = physics_handlers.cb_phystask_formulas
cb_phystask_list = physics_handlers.cb_phystask_list
cb_phystask_show = physics_handlers.cb_phystask_show
cb_menu_physics = physics_handlers.cb_menu_physics
cb_download_physics_full = physics_handlers.cb_download_physics_full
cb_download_physics_grade45 = physics_handlers.cb_download_physics_grade45
cb_download_physics_ticket_tasks = physics_handlers.cb_download_physics_ticket_tasks
cb_download_physics_tasks_cheatsheet = physics_handlers.cb_download_physics_tasks_cheatsheet

# ==================== АНАТОМИЯ (В РАЗРАБОТКЕ, ПОКА ДОСТУПНО ТОЛЬКО АДМИНАМ) ====================
# Хендлеры и вся логика раздела вынесены в handlers/anatomy.py (свой Router) — здесь только
# регистрация роутера на dp и реэкспорт имён, на которые ссылается остальной код файла (главное
# меню, админ-панель) и тесты, обращающиеся к ним как tb.<имя>. Сам импорт стоит здесь (в самом
# конце файла, как и раньше стояла вся секция), чтобы handlers/anatomy.py при своём импорте уже
# видел все нужные ему имена (stats, save_stats, DIVIDER, ACTIVE_SUBSCRIPTION_TIERS и т.д.)
# определёнными в этом модуле.
from handlers import anatomy as anatomy_handlers  # noqa: E402 — deliberately late, see above

dp.include_router(anatomy_handlers.router)

ANATOMY_PUBLIC = anatomy_handlers.ANATOMY_PUBLIC
# ANATOMY_MAINTENANCE_MODE is intentionally NOT re-exported as a flat copy here — unlike every
# other name in this block, it's a scalar bool that both production code (get_main_menu, below)
# and tests toggle at runtime. A flat `X = anatomy_handlers.X` assignment only snapshots the
# value at import time; reassigning the copy (`tb.ANATOMY_MAINTENANCE_MODE = ...`) would silently
# stop affecting the real flag `cb_anatomy_root` reads in handlers/anatomy.py, and vice versa —
# always go through `anatomy_handlers.ANATOMY_MAINTENANCE_MODE` (both here and in tests) so
# there's exactly one source of truth.
ANATOMY_FLASH_SESSION_SIZE = anatomy_handlers.ANATOMY_FLASH_SESSION_SIZE
ANATOMY_MATCH_SESSION_SIZE = anatomy_handlers.ANATOMY_MATCH_SESSION_SIZE
ANATOMY_LATIN_SESSION_SIZE = anatomy_handlers.ANATOMY_LATIN_SESSION_SIZE
ANATOMY_LATIN_ALL_SESSION_SIZE = anatomy_handlers.ANATOMY_LATIN_ALL_SESSION_SIZE
ANATOMY_LATIN_LEADERBOARD_MSG_LIMIT = anatomy_handlers.ANATOMY_LATIN_LEADERBOARD_MSG_LIMIT
ANATOMY_FLASH_SESSIONS = anatomy_handlers.ANATOMY_FLASH_SESSIONS
ANATOMY_MATCH_SESSIONS = anatomy_handlers.ANATOMY_MATCH_SESSIONS
ANATOMY_LATIN_SESSIONS = anatomy_handlers.ANATOMY_LATIN_SESSIONS
anatomy_access_ok = anatomy_handlers.anatomy_access_ok
ANATOMY_FREE_SECTIONS = anatomy_handlers.ANATOMY_FREE_SECTIONS
anatomy_section_access_ok = anatomy_handlers.anatomy_section_access_ok
get_anatomy_dev_alert_text = anatomy_handlers.get_anatomy_dev_alert_text
get_anatomy_topic_data = anatomy_handlers.get_anatomy_topic_data
get_topic_section_key = anatomy_handlers.get_topic_section_key
get_anatomy_locked_text = anatomy_handlers.get_anatomy_locked_text
get_anatomy_locked_keyboard = anatomy_handlers.get_anatomy_locked_keyboard
get_anatomy_menu_keyboard = anatomy_handlers.get_anatomy_menu_keyboard
get_anatomy_section_keyboard = anatomy_handlers.get_anatomy_section_keyboard
get_anatomy_topic_keyboard = anatomy_handlers.get_anatomy_topic_keyboard
get_anatomy_bones_keyboard = anatomy_handlers.get_anatomy_bones_keyboard
get_bone_title = anatomy_handlers.get_bone_title
get_bone_material_list = anatomy_handlers.get_bone_material_list
get_bone_flashcards = anatomy_handlers.get_bone_flashcards
get_bone_pairs = anatomy_handlers.get_bone_pairs
get_bone_mnemonics = anatomy_handlers.get_bone_mnemonics
get_bone_latin_terms = anatomy_handlers.get_bone_latin_terms
ANATOMY_ATLAS_CREDITS = anatomy_handlers.ANATOMY_ATLAS_CREDITS
ANATOMY_ALBUM_PAGE_SIZE = anatomy_handlers.ANATOMY_ALBUM_PAGE_SIZE
anatomy_page_count = anatomy_handlers.anatomy_page_count
get_bone_images = anatomy_handlers.get_bone_images
get_anatomy_bone_hub_keyboard = anatomy_handlers.get_anatomy_bone_hub_keyboard
get_anatomy_bone_hub_text = anatomy_handlers.get_anatomy_bone_hub_text
ANATOMY_FILE_ID_CACHE_PATH = anatomy_handlers.ANATOMY_FILE_ID_CACHE_PATH
_load_anatomy_file_id_cache = anatomy_handlers._load_anatomy_file_id_cache
ANATOMY_FILE_ID_CACHE = anatomy_handlers.ANATOMY_FILE_ID_CACHE
_write_anatomy_file_id_cache = anatomy_handlers._write_anatomy_file_id_cache
save_anatomy_file_id_cache = anatomy_handlers.save_anatomy_file_id_cache
_anatomy_image_key = anatomy_handlers._anatomy_image_key
_anatomy_image_media = anatomy_handlers._anatomy_image_media
_cache_anatomy_file_id = anatomy_handlers._cache_anatomy_file_id
build_input_media_photo = anatomy_handlers.build_input_media_photo
send_anatomy_album = anatomy_handlers.send_anatomy_album
get_topic_atlas_images = anatomy_handlers.get_topic_atlas_images
get_topic_latin_terms = anatomy_handlers.get_topic_latin_terms
get_all_latin_terms = anatomy_handlers.get_all_latin_terms
start_anatomy_latin_session = anatomy_handlers.start_anatomy_latin_session
get_anatomy_latin_keyboard = anatomy_handlers.get_anatomy_latin_keyboard
pick_anatomy_latin_distractors = anatomy_handlers.pick_anatomy_latin_distractors
render_anatomy_latin_question = anatomy_handlers.render_anatomy_latin_question
record_anatomy_latin_score = anatomy_handlers.record_anatomy_latin_score
get_anatomy_latin_leaderboard_text = anatomy_handlers.get_anatomy_latin_leaderboard_text
get_anatomy_latin_leaderboard_keyboard = anatomy_handlers.get_anatomy_latin_leaderboard_keyboard
render_anatomy_latin_summary = anatomy_handlers.render_anatomy_latin_summary
get_bone_material_keyboard = anatomy_handlers.get_bone_material_keyboard
get_bone_material_text = anatomy_handlers.get_bone_material_text
get_anatomy_material_keyboard = anatomy_handlers.get_anatomy_material_keyboard
get_anatomy_material_text = anatomy_handlers.get_anatomy_material_text
get_anatomy_material_list_keyboard = anatomy_handlers.get_anatomy_material_list_keyboard
start_anatomy_flash_session = anatomy_handlers.start_anatomy_flash_session
get_anatomy_flash_question_keyboard = anatomy_handlers.get_anatomy_flash_question_keyboard
get_anatomy_flash_answer_keyboard = anatomy_handlers.get_anatomy_flash_answer_keyboard
get_anatomy_flash_summary_keyboard = anatomy_handlers.get_anatomy_flash_summary_keyboard
render_anatomy_flash_question = anatomy_handlers.render_anatomy_flash_question
render_anatomy_flash_answer = anatomy_handlers.render_anatomy_flash_answer
render_anatomy_flash_summary = anatomy_handlers.render_anatomy_flash_summary
get_anatomy_all_pairs = anatomy_handlers.get_anatomy_all_pairs
start_anatomy_match_session = anatomy_handlers.start_anatomy_match_session
get_anatomy_match_keyboard = anatomy_handlers.get_anatomy_match_keyboard
render_anatomy_match_question = anatomy_handlers.render_anatomy_match_question
render_anatomy_match_summary = anatomy_handlers.render_anatomy_match_summary
get_anatomy_mnemonics_keyboard = anatomy_handlers.get_anatomy_mnemonics_keyboard
get_anatomy_mnemonic_text = anatomy_handlers.get_anatomy_mnemonic_text
get_bone_mnemonics_keyboard = anatomy_handlers.get_bone_mnemonics_keyboard
get_bone_mnemonic_text = anatomy_handlers.get_bone_mnemonic_text
get_anatomy_root_keyboard = anatomy_handlers.get_anatomy_root_keyboard
get_anatomy_maintenance_text = anatomy_handlers.get_anatomy_maintenance_text
get_anatomy_maintenance_keyboard = anatomy_handlers.get_anatomy_maintenance_keyboard
cb_anatomy_root = anatomy_handlers.cb_anatomy_root
cb_anatomy_menu = anatomy_handlers.cb_anatomy_menu
get_anatomy_video_text = anatomy_handlers.get_anatomy_video_text
cb_anatomy_section = anatomy_handlers.cb_anatomy_section
cb_anatomy_section_video = anatomy_handlers.cb_anatomy_section_video
cb_anatomy_topic = anatomy_handlers.cb_anatomy_topic
cb_anatomy_topic_video = anatomy_handlers.cb_anatomy_topic_video
cb_anatomy_bones = anatomy_handlers.cb_anatomy_bones
cb_anatomy_bone_hub = anatomy_handlers.cb_anatomy_bone_hub
cb_anatomy_bone_material = anatomy_handlers.cb_anatomy_bone_material
cb_anatomy_bone_slides = anatomy_handlers.cb_anatomy_bone_slides
cb_anatomy_bone_atlas = anatomy_handlers.cb_anatomy_bone_atlas
cb_anatomy_atlas = anatomy_handlers.cb_anatomy_atlas
cb_anatomy_bone_flash_start = anatomy_handlers.cb_anatomy_bone_flash_start
cb_anatomy_bone_match_start = anatomy_handlers.cb_anatomy_bone_match_start
cb_anatomy_bone_latin_start = anatomy_handlers.cb_anatomy_bone_latin_start
cb_anatomy_bone_mnemonics = anatomy_handlers.cb_anatomy_bone_mnemonics
cb_anatomy_material_list = anatomy_handlers.cb_anatomy_material_list
cb_anatomy_material = anatomy_handlers.cb_anatomy_material
cb_anatomy_flash_start = anatomy_handlers.cb_anatomy_flash_start
cb_anatomy_flash_show_answer = anatomy_handlers.cb_anatomy_flash_show_answer
cb_anatomy_flash_answer = anatomy_handlers.cb_anatomy_flash_answer
cb_anatomy_flash_stop = anatomy_handlers.cb_anatomy_flash_stop
cb_anatomy_match_start = anatomy_handlers.cb_anatomy_match_start
cb_anatomy_match_answer = anatomy_handlers.cb_anatomy_match_answer
cb_anatomy_match_stop = anatomy_handlers.cb_anatomy_match_stop
cb_anatomy_latin_all_start = anatomy_handlers.cb_anatomy_latin_all_start
cb_anatomy_latin_leaderboard = anatomy_handlers.cb_anatomy_latin_leaderboard
cb_anatomy_latin_start = anatomy_handlers.cb_anatomy_latin_start
cb_anatomy_latin_answer = anatomy_handlers.cb_anatomy_latin_answer
cb_anatomy_latin_stop = anatomy_handlers.cb_anatomy_latin_stop
cb_anatomy_mnemonics = anatomy_handlers.cb_anatomy_mnemonics
cb_anatomy_picture = anatomy_handlers.cb_anatomy_picture
get_anatomy_exam_menu_keyboard = anatomy_handlers.get_anatomy_exam_menu_keyboard
cb_anatomy_exam_menu = anatomy_handlers.cb_anatomy_exam_menu
get_anatomy_exam_practice_section = anatomy_handlers.get_anatomy_exam_practice_section
get_anatomy_exam_practice_section_keyboard = anatomy_handlers.get_anatomy_exam_practice_section_keyboard
cb_anatomy_exam_practice = anatomy_handlers.cb_anatomy_exam_practice
get_anatomy_exam_practice_question_list_keyboard = anatomy_handlers.get_anatomy_exam_practice_question_list_keyboard
cb_anatomy_exam_practice_section = anatomy_handlers.cb_anatomy_exam_practice_section
get_anatomy_exam_practice_question_text = anatomy_handlers.get_anatomy_exam_practice_question_text
get_anatomy_exam_practice_question_keyboard = anatomy_handlers.get_anatomy_exam_practice_question_keyboard
cb_anatomy_exam_practice_question = anatomy_handlers.cb_anatomy_exam_practice_question
get_anatomy_exam_theory_section = anatomy_handlers.get_anatomy_exam_theory_section
get_anatomy_exam_theory_section_keyboard = anatomy_handlers.get_anatomy_exam_theory_section_keyboard
cb_anatomy_exam_theory = anatomy_handlers.cb_anatomy_exam_theory
get_anatomy_exam_theory_question_list_keyboard = anatomy_handlers.get_anatomy_exam_theory_question_list_keyboard
cb_anatomy_exam_theory_section = anatomy_handlers.cb_anatomy_exam_theory_section
get_anatomy_exam_theory_question_text = anatomy_handlers.get_anatomy_exam_theory_question_text
get_anatomy_exam_theory_question_keyboard = anatomy_handlers.get_anatomy_exam_theory_question_keyboard
cb_anatomy_exam_theory_question = anatomy_handlers.cb_anatomy_exam_theory_question
ANATOMY_EXAM_TEST_SESSIONS = anatomy_handlers.ANATOMY_EXAM_TEST_SESSIONS
ANATOMY_EXAM_TEST_MISTAKES = anatomy_handlers.ANATOMY_EXAM_TEST_MISTAKES
ANATOMY_EXAM_TEST_OPTION_LETTERS = anatomy_handlers.ANATOMY_EXAM_TEST_OPTION_LETTERS
ANATOMY_EXAM_FLASH_SIZE = anatomy_handlers.ANATOMY_EXAM_FLASH_SIZE
ANATOMY_EXAM_TEST_ALL_QUESTIONS = anatomy_handlers.ANATOMY_EXAM_TEST_ALL_QUESTIONS
get_anatomy_exam_test_part = anatomy_handlers.get_anatomy_exam_test_part
get_anatomy_exam_test_mode = anatomy_handlers.get_anatomy_exam_test_mode
set_anatomy_exam_test_mode = anatomy_handlers.set_anatomy_exam_test_mode
record_anatomy_exam_test_score = anatomy_handlers.record_anatomy_exam_test_score
get_anatomy_exam_test_leaderboard_text = anatomy_handlers.get_anatomy_exam_test_leaderboard_text
get_anatomy_exam_test_leaderboard_keyboard = anatomy_handlers.get_anatomy_exam_test_leaderboard_keyboard
cb_anatomy_exam_test_leaderboard = anatomy_handlers.cb_anatomy_exam_test_leaderboard
record_anatomy_exam_flash_score = anatomy_handlers.record_anatomy_exam_flash_score
get_anatomy_exam_flash_leaderboard_text = anatomy_handlers.get_anatomy_exam_flash_leaderboard_text
get_anatomy_exam_flash_leaderboard_keyboard = anatomy_handlers.get_anatomy_exam_flash_leaderboard_keyboard
cb_anatomy_exam_flash_leaderboard = anatomy_handlers.cb_anatomy_exam_flash_leaderboard
get_anatomy_exam_test_menu_keyboard = anatomy_handlers.get_anatomy_exam_test_menu_keyboard
render_anatomy_exam_test_menu = anatomy_handlers.render_anatomy_exam_test_menu
cb_anatomy_exam_test_menu = anatomy_handlers.cb_anatomy_exam_test_menu
cb_anatomy_exam_test_mode_toggle = anatomy_handlers.cb_anatomy_exam_test_mode_toggle
start_anatomy_exam_test_session = anatomy_handlers.start_anatomy_exam_test_session
start_anatomy_exam_flash_session = anatomy_handlers.start_anatomy_exam_flash_session
get_anatomy_exam_test_keyboard = anatomy_handlers.get_anatomy_exam_test_keyboard
render_anatomy_exam_test_question = anatomy_handlers.render_anatomy_exam_test_question
get_anatomy_exam_test_mistake_text = anatomy_handlers.get_anatomy_exam_test_mistake_text
get_anatomy_exam_test_mistake_keyboard = anatomy_handlers.get_anatomy_exam_test_mistake_keyboard
render_anatomy_exam_test_summary = anatomy_handlers.render_anatomy_exam_test_summary
cb_anatomy_exam_test_start = anatomy_handlers.cb_anatomy_exam_test_start
cb_anatomy_exam_test_flash_start = anatomy_handlers.cb_anatomy_exam_test_flash_start
cb_anatomy_exam_test_answer = anatomy_handlers.cb_anatomy_exam_test_answer
cb_anatomy_exam_test_stop = anatomy_handlers.cb_anatomy_exam_test_stop
cb_anatomy_exam_test_mistakes = anatomy_handlers.cb_anatomy_exam_test_mistakes

# ==================== ГИСТОЛОГИЯ ====================
# Хендлеры и вся логика раздела вынесены в handlers/histology.py (свой Router) — здесь только
# регистрация роутера на dp и реэкспорт имён, на которые ссылается остальной код файла (главное
# меню, админ-панель) и тесты, обращающиеся к ним как tb.<имя>. Сам импорт стоит здесь (в самом
# конце файла, как и раньше стояла вся секция), чтобы handlers/histology.py при своём импорте
# уже видел все нужные ему имена (stats, save_stats, DIVIDER, REFERRAL_* и т.д.) определёнными в
# этом модуле.
from handlers import histology as histology_handlers  # noqa: E402 — deliberately late, see above

dp.include_router(histology_handlers.router)

HISTOLOGY_PUBLIC = histology_handlers.HISTOLOGY_PUBLIC
HISTOLOGY_PROMO_SECONDS = histology_handlers.HISTOLOGY_PROMO_SECONDS
HISTOLOGY_WARNING_THRESHOLD = histology_handlers.HISTOLOGY_WARNING_THRESHOLD
HISTOLOGY_WARNING_COOLDOWN_SECONDS = histology_handlers.HISTOLOGY_WARNING_COOLDOWN_SECONDS
HISTOLOGY_GUESS_SESSION_SIZE = histology_handlers.HISTOLOGY_GUESS_SESSION_SIZE
HISTOLOGY_GUESS_SESSIONS = histology_handlers.HISTOLOGY_GUESS_SESSIONS
get_histology_temp_expiry = histology_handlers.get_histology_temp_expiry
has_histology_temp_access = histology_handlers.has_histology_temp_access
histology_permanently_unlocked = histology_handlers.histology_permanently_unlocked
histology_access_ok = histology_handlers.histology_access_ok
histology_gate_ok = histology_handlers.histology_gate_ok
get_histology_specimen = histology_handlers.get_histology_specimen
get_histology_locked_text = histology_handlers.get_histology_locked_text
get_histology_locked_keyboard = histology_handlers.get_histology_locked_keyboard
announce_histology_promo_start = histology_handlers.announce_histology_promo_start
get_histology_menu_keyboard = histology_handlers.get_histology_menu_keyboard
get_histology_topic_text = histology_handlers.get_histology_topic_text
get_histology_topic_keyboard = histology_handlers.get_histology_topic_keyboard
get_histology_specimen_text = histology_handlers.get_histology_specimen_text
get_histology_specimen_keyboard = histology_handlers.get_histology_specimen_keyboard
get_histology_image_keyboard = histology_handlers.get_histology_image_keyboard
render_histology_image = histology_handlers.render_histology_image
get_histology_guess_pool = histology_handlers.get_histology_guess_pool
start_histology_guess_session = histology_handlers.start_histology_guess_session
get_histology_guess_question_keyboard = histology_handlers.get_histology_guess_question_keyboard
get_histology_guess_answer_keyboard = histology_handlers.get_histology_guess_answer_keyboard
get_histology_guess_summary_keyboard = histology_handlers.get_histology_guess_summary_keyboard
render_histology_guess_question = histology_handlers.render_histology_guess_question
render_histology_guess_answer = histology_handlers.render_histology_guess_answer
render_histology_guess_summary = histology_handlers.render_histology_guess_summary
cb_histology_menu = histology_handlers.cb_histology_menu
cb_histology_topic = histology_handlers.cb_histology_topic
cb_histology_specimen = histology_handlers.cb_histology_specimen
cb_histology_img = histology_handlers.cb_histology_img
cb_histology_guess_start = histology_handlers.cb_histology_guess_start
cb_histology_guess_show_answer = histology_handlers.cb_histology_guess_show_answer
cb_histology_guess_answer = histology_handlers.cb_histology_guess_answer
cb_histology_guess_stop = histology_handlers.cb_histology_guess_stop

# ==================== ОПЕРАТИВНАЯ ХИРУРГИЯ ====================
# Тот же паттерн, что ГИСТОЛОГИЯ выше — свой Router, импортируется в самом конце файла, когда
# все нужные оттуда имена (OPERATIVE_SURGERY, DIVIDER, safe_edit_text, OH_SEARCH_PENDING,
# search_operative_surgery) уже определены здесь. Раздел свободен для всех — своего гейта нет,
# поэтому здесь реэкспортированы только клавиатуры/тексты и сами callback-хендлеры, без
# вспомогательных access-предикатов (их для этого раздела попросту нет).
from handlers import operative_surgery as operative_surgery_handlers  # noqa: E402 — deliberately late, see above

dp.include_router(operative_surgery_handlers.router)

get_oh_topic = operative_surgery_handlers.get_oh_topic
get_oh_volume = operative_surgery_handlers.get_oh_volume
get_oh_menu_text = operative_surgery_handlers.get_oh_menu_text
get_oh_menu_keyboard = operative_surgery_handlers.get_oh_menu_keyboard
get_oh_volumes_text = operative_surgery_handlers.get_oh_volumes_text
get_oh_volumes_keyboard = operative_surgery_handlers.get_oh_volumes_keyboard
get_oh_volume_text = operative_surgery_handlers.get_oh_volume_text
get_oh_volume_keyboard = operative_surgery_handlers.get_oh_volume_keyboard
get_oh_volume_control_text = operative_surgery_handlers.get_oh_volume_control_text
get_oh_volume_control_keyboard = operative_surgery_handlers.get_oh_volume_control_keyboard
get_oh_topic_text = operative_surgery_handlers.get_oh_topic_text
get_oh_topic_keyboard = operative_surgery_handlers.get_oh_topic_keyboard
get_oh_material_text = operative_surgery_handlers.get_oh_material_text
get_oh_material_keyboard = operative_surgery_handlers.get_oh_material_keyboard
get_oh_quick_text = operative_surgery_handlers.get_oh_quick_text
get_oh_quick_keyboard = operative_surgery_handlers.get_oh_quick_keyboard
get_oh_topic_control_text = operative_surgery_handlers.get_oh_topic_control_text
get_oh_topic_control_keyboard = operative_surgery_handlers.get_oh_topic_control_keyboard
get_oh_instruments_text = operative_surgery_handlers.get_oh_instruments_text
get_oh_instruments_keyboard = operative_surgery_handlers.get_oh_instruments_keyboard
get_oh_instrument_group_text = operative_surgery_handlers.get_oh_instrument_group_text
get_oh_instrument_group_keyboard = operative_surgery_handlers.get_oh_instrument_group_keyboard
get_oh_projections_text = operative_surgery_handlers.get_oh_projections_text
get_oh_projections_keyboard = operative_surgery_handlers.get_oh_projections_keyboard
get_oh_projection_group_text = operative_surgery_handlers.get_oh_projection_group_text
get_oh_projection_group_keyboard = operative_surgery_handlers.get_oh_projection_group_keyboard
get_oh_stations_text = operative_surgery_handlers.get_oh_stations_text
get_oh_stations_keyboard = operative_surgery_handlers.get_oh_stations_keyboard
get_oh_station_group_text = operative_surgery_handlers.get_oh_station_group_text
get_oh_station_group_keyboard = operative_surgery_handlers.get_oh_station_group_keyboard
cb_oh_menu = operative_surgery_handlers.cb_oh_menu
cb_oh_volumes = operative_surgery_handlers.cb_oh_volumes
cb_oh_volume = operative_surgery_handlers.cb_oh_volume
cb_oh_volume_control = operative_surgery_handlers.cb_oh_volume_control
cb_oh_topic = operative_surgery_handlers.cb_oh_topic
cb_oh_material = operative_surgery_handlers.cb_oh_material
cb_oh_quick = operative_surgery_handlers.cb_oh_quick
cb_oh_topic_control = operative_surgery_handlers.cb_oh_topic_control
cb_oh_instruments = operative_surgery_handlers.cb_oh_instruments
cb_oh_instrument_group = operative_surgery_handlers.cb_oh_instrument_group
cb_oh_projections = operative_surgery_handlers.cb_oh_projections
cb_oh_projection_group = operative_surgery_handlers.cb_oh_projection_group
cb_oh_stations = operative_surgery_handlers.cb_oh_stations
cb_oh_station_group = operative_surgery_handlers.cb_oh_station_group
cb_oh_search_prompt = operative_surgery_handlers.cb_oh_search_prompt

# ==================== НОРМАЛЬНАЯ ФИЗИОЛОГИЯ ====================
# Тот же паттерн, что ОПЕРАТИВНАЯ ХИРУРГИЯ выше — свой Router, импортируется в самом конце файла,
# когда все нужные оттуда имена (PHYSIOLOGY, DIVIDER, safe_edit_text, stats, save_stats,
# PHYS_SEARCH_PENDING) уже определены здесь. Раздел свободен для всех — своего гейта нет.
from handlers import physiology as physiology_handlers  # noqa: E402 — deliberately late, see above

dp.include_router(physiology_handlers.router)

get_phys_topic = physiology_handlers.get_phys_topic
phys_topic_ids_in_order = physiology_handlers.phys_topic_ids_in_order
get_phys_topic_quiz_pool = physiology_handlers.get_phys_topic_quiz_pool
get_phys_progress = physiology_handlers.get_phys_progress
phys_mark_opened = physiology_handlers.phys_mark_opened
phys_mark_card_done = physiology_handlers.phys_mark_card_done
phys_record_quiz_answer = physiology_handlers.phys_record_quiz_answer
phys_record_quiz_session_complete = physiology_handlers.phys_record_quiz_session_complete
phys_topic_status = physiology_handlers.phys_topic_status
phys_favorites = physiology_handlers.phys_favorites
phys_is_favorite = physiology_handlers.phys_is_favorite
phys_toggle_favorite = physiology_handlers.phys_toggle_favorite
build_phys_learn_cards = physiology_handlers.build_phys_learn_cards
render_phys_learn_card = physiology_handlers.render_phys_learn_card
get_phys_learn_keyboard = physiology_handlers.get_phys_learn_keyboard
render_phys_comparison_body = physiology_handlers.render_phys_comparison_body
get_phys_menu_text = physiology_handlers.get_phys_menu_text
get_phys_menu_keyboard = physiology_handlers.get_phys_menu_keyboard
phys_next_topic_for_continue = physiology_handlers.phys_next_topic_for_continue
get_phys_topics_text = physiology_handlers.get_phys_topics_text
get_phys_topics_keyboard = physiology_handlers.get_phys_topics_keyboard
get_phys_topic_text = physiology_handlers.get_phys_topic_text
get_phys_topic_keyboard = physiology_handlers.get_phys_topic_keyboard
get_phys_read_text = physiology_handlers.get_phys_read_text
get_phys_read_keyboard = physiology_handlers.get_phys_read_keyboard
get_phys_quick_text = physiology_handlers.get_phys_quick_text
get_phys_quick_keyboard = physiology_handlers.get_phys_quick_keyboard
build_phys_chains = physiology_handlers.build_phys_chains
get_phys_chain_text = physiology_handlers.get_phys_chain_text
get_phys_chain_keyboard = physiology_handlers.get_phys_chain_keyboard
get_phys_cmp_text = physiology_handlers.get_phys_cmp_text
get_phys_cmp_keyboard = physiology_handlers.get_phys_cmp_keyboard
start_phys_quiz_session = physiology_handlers.start_phys_quiz_session
render_phys_quiz_question = physiology_handlers.render_phys_quiz_question
get_phys_quiz_question_keyboard = physiology_handlers.get_phys_quiz_question_keyboard
render_phys_quiz_answer = physiology_handlers.render_phys_quiz_answer
get_phys_quiz_answer_keyboard = physiology_handlers.get_phys_quiz_answer_keyboard
get_phys_quiz_summary_text = physiology_handlers.get_phys_quiz_summary_text
get_phys_quiz_summary_keyboard = physiology_handlers.get_phys_quiz_summary_keyboard
get_phys_favorites_text = physiology_handlers.get_phys_favorites_text
get_phys_favorites_keyboard = physiology_handlers.get_phys_favorites_keyboard
get_phys_progress_text = physiology_handlers.get_phys_progress_text
get_phys_progress_keyboard = physiology_handlers.get_phys_progress_keyboard
search_physiology = physiology_handlers.search_physiology
cb_phys_menu = physiology_handlers.cb_phys_menu
cb_phys_continue = physiology_handlers.cb_phys_continue
cb_phys_topics = physiology_handlers.cb_phys_topics
cb_phys_qpick = physiology_handlers.cb_phys_qpick
cb_phys_zpick = physiology_handlers.cb_phys_zpick
cb_phys_topic = physiology_handlers.cb_phys_topic
cb_phys_fav_toggle = physiology_handlers.cb_phys_fav_toggle
cb_phys_learn = physiology_handlers.cb_phys_learn
cb_phys_learn_ok = physiology_handlers.cb_phys_learn_ok
cb_phys_read = physiology_handlers.cb_phys_read
cb_phys_quick = physiology_handlers.cb_phys_quick
cb_phys_chains = physiology_handlers.cb_phys_chains
cb_phys_cmp = physiology_handlers.cb_phys_cmp
cb_phys_quiz_start = physiology_handlers.cb_phys_quiz_start
cb_phys_quiz_answer = physiology_handlers.cb_phys_quiz_answer
cb_phys_quiz_next = physiology_handlers.cb_phys_quiz_next
cb_phys_quiz_stop = physiology_handlers.cb_phys_quiz_stop
cb_phys_mini = physiology_handlers.cb_phys_mini
cb_phys_mini_answer = physiology_handlers.cb_phys_mini_answer
cb_phys_favorites = physiology_handlers.cb_phys_favorites
cb_phys_progress = physiology_handlers.cb_phys_progress
cb_phys_search_prompt = physiology_handlers.cb_phys_search_prompt
get_rk_control = physiology_handlers.get_rk_control
rk_control_ids_in_order = physiology_handlers.rk_control_ids_in_order
build_rk_pages = physiology_handlers.build_rk_pages
get_rk_menu_text = physiology_handlers.get_rk_menu_text
get_rk_menu_keyboard = physiology_handlers.get_rk_menu_keyboard
get_rk_page_keyboard = physiology_handlers.get_rk_page_keyboard
cb_phys_rk_menu = physiology_handlers.cb_phys_rk_menu
cb_phys_rk_page = physiology_handlers.cb_phys_rk_page

# ==================== ЗАПУСК ====================
async def setup_bot_commands() -> None:
    default_commands = [
        BotCommand(command="start", description="Начать работу с ботом"),
        BotCommand(command="random", description="Получить случайный билет"),
        BotCommand(command="help", description="Помощь и инструкция"),
    ]
    await bot.set_my_commands(default_commands, scope=BotCommandScopeDefault())

    admin_commands = default_commands + [
        BotCommand(command="admin", description="Админ-панель"),
    ]
    for admin_id in ADMIN_IDS:
        try:
            await bot.set_my_commands(admin_commands, scope=BotCommandScopeChat(chat_id=admin_id))
        except Exception:
            logger.exception("Не удалось установить админ-команды для %s", admin_id)

# AI_BUILD_EMBEDDINGS_ON_START=0 полностью отключает пересчёт эмбеддингов на старте (RAG падает
# на чистый keyword/IDF-поиск, как без ключа OpenAI вообще) — аварийный рубильник на случай, если
# кэш-файл потерян/повреждён/лежит не на постоянном томе или бот попал в crash-loop: без него
# КАЖДЫЙ рестарт пытался бы заново оплатить эмбеддинги всей базы. AI_MAX_EMBEDDING_BUILD_ITEMS_PER_START
# — более мягкая версия той же защиты (см. ai_rag.MAX_EMBEDDING_BUILD_ITEMS_PER_START): не отключает
# пересчёт целиком, а ограничивает бюджет ОДНОГО прогона, размазывая полный пересчёт по нескольким
# рестартам вместо одного большого счёта.
AI_BUILD_EMBEDDINGS_ON_START = os.environ.get("AI_BUILD_EMBEDDINGS_ON_START", "1") != "0"
AI_MAX_EMBEDDING_BUILD_ITEMS_PER_START = int(
    os.environ.get("AI_MAX_EMBEDDING_BUILD_ITEMS_PER_START", str(ai_rag.MAX_EMBEDDING_BUILD_ITEMS_PER_START))
)

async def main():
    logger.info("Бот запускается...")
    logger.info("Загружена статистика: %d пользователей", len(stats["total_users"]))
    await setup_bot_commands()
    resume_battle_timer_if_needed()
    if AI_BUILD_EMBEDDINGS_ON_START:
        # Фоновой задачей, не блокируя polling — на первом прогоне подсчёт эмбеддингов всей базы
        # может занять заметное время (см. ai_rag.build_embeddings); инкрементальный кэш на диске
        # делает все последующие запуски бота почти мгновенными.
        asyncio.create_task(ai_rag.build_embeddings(
            os.path.join(STATS_DIR, "ai_rag_embeddings_cache.json"),
            max_items=AI_MAX_EMBEDDING_BUILD_ITEMS_PER_START,
        ))
    else:
        logger.info("AI_BUILD_EMBEDDINGS_ON_START=0 — пересчёт эмбеддингов RAG на старте пропущен")
    try:
        await dp.start_polling(bot)
    finally:
        _stats_executor.shutdown(wait=True)

if __name__ == "__main__":
    asyncio.run(main())
