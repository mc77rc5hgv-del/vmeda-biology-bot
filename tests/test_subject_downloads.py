# -*- coding: utf-8 -*-
import asyncio
import io
from _bootstrap import tb
from docx import Document as DocxDocument

class FakeUser:
    def __init__(self, uid):
        self.id = uid

class FakeMsg:
    def __init__(self):
        self.edits = []
        self.documents = []
    async def edit_text(self, text, **kwargs):
        self.edits.append((text, kwargs.get("reply_markup")))
        return self
    async def delete(self):
        pass
    async def answer(self, text, **kwargs):
        self.edits.append((text, kwargs.get("reply_markup")))
        return self
    async def answer_document(self, document, **kwargs):
        self.documents.append((document, kwargs.get("caption")))
        return self

class FakeCB:
    def __init__(self, data, uid=111222333):
        self.data = data
        self.from_user = FakeUser(uid)
        self.message = FakeMsg()
    async def answer(self, text=None, show_alert=False):
        pass

def kb_data(markup):
    return [b.callback_data for row in markup.inline_keyboard for b in row]

def docx_text(buffered_input_file) -> str:
    doc = DocxDocument(io.BytesIO(buffered_input_file.data))
    return "\n".join(p.text for p in doc.paragraphs)

async def main():
    # strip_html_tags removes markup without mangling plain text
    assert tb.strip_html_tags("<b>Ядро</b> клетки — <i>nucleus</i>") == "Ядро клетки — nucleus"
    print("strip_html_tags: OK")

    # each subject menu offers the requested download button(s)
    bio_kb_data = kb_data(tb.get_biology_menu())
    assert "download_biology_tickets" in bio_kb_data
    phys_kb_data = kb_data(tb.get_physics_menu())
    assert "download_physics_full" in phys_kb_data and "download_physics_ticket_tasks" in phys_kb_data
    chem_kb_data = kb_data(tb.get_chemistry_menu())
    assert "download_chemistry_labs" in chem_kb_data and "download_chemistry_tasks" in chem_kb_data
    print("menus expose download buttons: OK")

    # biology: all 40 tickets present, HTML stripped, content matches source, real .docx
    cb = FakeCB("download_biology_tickets")
    await tb.cb_download_biology_tickets(cb)
    assert cb.message.documents
    doc, caption = cb.message.documents[0]
    assert doc.data[:2] == b"PK"  # docx is a zip archive
    text = docx_text(doc)
    assert "<b>" not in text and "<i>" not in text
    for ticket in tb.TICKETS:
        assert ticket["title"] in text
    assert tb.strip_html_tags(tb.TICKETS[0]["questions"][0]["answer"]).split("\n\n")[0] in text
    assert caption and f"@{tb.BOT_USERNAME}" in caption
    print(f"biology tickets file: {len(tb.TICKETS)} tickets, HTML-free docx, content matches: OK")

    # physics: 186 questions + all 9 task topics present
    cb2 = FakeCB("download_physics_full")
    await tb.cb_download_physics_full(cb2)
    doc2, caption2 = cb2.message.documents[0]
    text2 = docx_text(doc2)
    assert "<b>" not in text2
    assert len(tb.PHYSICS_QUESTIONS) == 186
    for item in tb.PHYSICS_QUESTIONS.values():
        assert tb.strip_html_tags(item["title"]) in text2
    for topic in tb.PHYSICS_TASKS.values():
        assert tb.strip_html_tags(topic["title"]) in text2
    assert caption2 and f"@{tb.BOT_USERNAME}" in caption2
    print("physics full file: 186 questions + all task topics present, HTML-free: OK")

    # physics: 4-ticket task answers file has all 20 solved problems
    cb3 = FakeCB("download_physics_ticket_tasks")
    await tb.cb_download_physics_ticket_tasks(cb3)
    doc3, caption3 = cb3.message.documents[0]
    text3 = docx_text(doc3)
    assert "<b>" not in text3
    total_tasks = 0
    for ticket in tb.PHYSICS_TEST_TICKETS.values():
        assert ticket["title"] in text3
        for task in ticket.get("tasks", []):
            total_tasks += 1
            assert tb.strip_html_tags(task["title"]) in text3
    assert total_tasks == 20
    assert caption3 and f"@{tb.BOT_USERNAME}" in caption3
    print("physics ticket-tasks file: all 20 solved problems present: OK")

    # chemistry: labs file
    cb4 = FakeCB("download_chemistry_labs")
    await tb.cb_download_chemistry_labs(cb4)
    doc4, caption4 = cb4.message.documents[0]
    text4 = docx_text(doc4)
    assert "<b>" not in text4
    for lab in tb.CHEMISTRY_LABS["labs"]:
        assert tb.strip_html_tags(lab["theme"]) in text4
        for exp in lab.get("experiments", []):
            assert tb.strip_html_tags(exp["name"]) in text4
    assert caption4 and f"@{tb.BOT_USERNAME}" in caption4
    print(f"chemistry labs file: {len(tb.CHEMISTRY_LABS['labs'])} labs present, HTML-free: OK")

    # chemistry: tasks file
    cb5 = FakeCB("download_chemistry_tasks")
    await tb.cb_download_chemistry_tasks(cb5)
    doc5, caption5 = cb5.message.documents[0]
    text5 = docx_text(doc5)
    assert "<b>" not in text5
    for topic in tb.CHEMISTRY_TASKS.values():
        assert tb.strip_html_tags(topic["title"]) in text5
    assert caption5 and f"@{tb.BOT_USERNAME}" in caption5
    print(f"chemistry tasks file: {len(tb.CHEMISTRY_TASKS)} topics present, HTML-free: OK")

    # filenames are sane .docx files, and bold/italic runs actually survived as real formatting
    for doc, _ in (cb.message.documents[0], cb2.message.documents[0], cb3.message.documents[0],
                   cb4.message.documents[0], cb5.message.documents[0]):
        assert doc.filename.endswith(".docx")
    print("all five files have .docx filenames: OK")

    reloaded = DocxDocument(io.BytesIO(doc.data))
    assert any(run.bold for p in reloaded.paragraphs for run in p.runs), "expected at least one bold run"
    print("bold Telegram-HTML markup survived as real docx formatting: OK")

    print("ALL SUBJECT DOWNLOAD TESTS PASSED")

asyncio.run(main())
