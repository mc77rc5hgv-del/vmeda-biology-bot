"""Inventory and extract text from the supplied biochemistry source set."""

from __future__ import annotations

import hashlib
import json
import re
import sys
import zipfile
from pathlib import Path

from docx import Document
from pypdf import PdfReader


def clean(text: str) -> str:
    return re.sub(r"[ \t]+", " ", text.replace("\u00a0", " ")).strip()


def inspect_docx(path: Path) -> dict:
    doc = Document(path)
    paragraphs = [clean(p.text) for p in doc.paragraphs if clean(p.text)]
    tables = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            values = [clean(cell.text) for cell in row.cells]
            if any(values):
                rows.append(" | ".join(values))
        if rows:
            tables.append(rows)
    with zipfile.ZipFile(path) as archive:
        media = [n for n in archive.namelist() if n.startswith("word/media/")]
    return {
        "kind": "docx",
        "paragraph_count": len(paragraphs),
        "table_count": len(tables),
        "table_row_count": sum(map(len, tables)),
        "image_count": len(media),
        "text": "\n".join(paragraphs + [row for table in tables for row in table]),
    }


def inspect_pdf(path: Path) -> dict:
    reader = PdfReader(path)
    pages = []
    image_count = 0
    for index, page in enumerate(reader.pages, 1):
        text = clean(page.extract_text() or "")
        pages.append({"page": index, "chars": len(text), "text": text})
        try:
            image_count += len(page.images)
        except Exception:
            pass
    return {
        "kind": "pdf",
        "page_count": len(pages),
        "empty_page_count": sum(p["chars"] < 30 for p in pages),
        "image_count": image_count,
        "pages": pages,
        "text": "\n\n".join(f"[Страница {p['page']}]\n{p['text']}" for p in pages),
    }


def main(source_dir: Path, output_dir: Path) -> None:
    output_dir.mkdir(parents=True, exist_ok=True)
    manifest = []
    for path in sorted(source_dir.iterdir(), key=lambda p: p.name.casefold()):
        if not path.is_file() or path.suffix.lower() not in {".pdf", ".docx"}:
            continue
        data = inspect_pdf(path) if path.suffix.lower() == ".pdf" else inspect_docx(path)
        digest = hashlib.sha256(path.read_bytes()).hexdigest()
        text = data.pop("text")
        text_path = output_dir / f"{path.name}.txt"
        text_path.write_text(text, encoding="utf-8")
        manifest.append({
            "file": path.name,
            "bytes": path.stat().st_size,
            "sha256": digest,
            "extracted_chars": len(text),
            "text_file": str(text_path.relative_to(output_dir.parent)),
            **data,
        })
        print(path.name, len(text), data.get("page_count", data.get("paragraph_count")))
    (output_dir.parent / "source_manifest.json").write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8"
    )


if __name__ == "__main__":
    main(Path(sys.argv[1]), Path(sys.argv[2]))
