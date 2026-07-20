import json, re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

REPO = "/home/user/vmeda-biology-bot"

def strip_tags(s):
    return re.sub(r'<[^>]+>', '', s)

def clean_rest(rest):
    rest = strip_tags(rest).strip()
    rest = re.sub(r'^[—\-:\s]+', '', rest)
    rest = re.sub(r'\s+', ' ', rest)
    rest = re.sub(r'^\d+\.\s*', '', rest)
    return rest

def condense(answer, max_chars=600, max_lines=9):
    parts = [p.strip() for p in answer.split('\n\n') if p.strip()]
    pieces = []
    total = 0
    for p in parts:
        m = re.match(r'<b>(.*?)</b>\s*(.*)', p, re.S)
        if m:
            header = strip_tags(m.group(1)).strip().rstrip(':')
            rest = clean_rest(m.group(2))
            if rest.startswith('•'):
                first_bullet = rest.split('•')[1] if '•' in rest else rest
                rest = clean_rest(first_bullet)
            cut = re.split(r'[.;]\s|\n', rest, maxsplit=1)[0]
            cut = cut[:150].rstrip(',;:— ')
            if not cut:
                continue
            line = f"{header}: {cut}"
        else:
            line = strip_tags(p)
            line = re.sub(r'\s+', ' ', line).strip().lstrip('•').strip()
            if not line:
                continue
            line = line[:140]
        pieces.append(line)
        total += len(line)
        if total >= max_chars or len(pieces) >= max_lines:
            break
    return pieces

def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), '999999')
        borders.append(el)
    tcPr.append(borders)

def set_cell_margins(cell, top=40, bottom=40, left=60, right=60):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    mar = OxmlElement('w:tcMar')
    for edge, val in (('top', top), ('bottom', bottom), ('left', left), ('right', right)):
        node = OxmlElement(f'w:{edge}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        mar.append(node)
    tcPr.append(mar)

def add_run(paragraph, text, size=6, bold=False, color=None):
    r = paragraph.add_run(text)
    r.font.size = Pt(size)
    r.font.name = 'Arial Narrow'
    r.bold = bold
    if color:
        r.font.color.rgb = RGBColor(*color)
    return r

def fill_ticket_cell(cell, ticket):
    cell.vertical_alignment = 1
    set_cell_border(cell)
    set_cell_margins(cell)
    # clear default paragraph, reuse it as title
    p0 = cell.paragraphs[0]
    p0.paragraph_format.space_after = Pt(1)
    add_run(p0, f"БИЛЕТ {ticket['num']}", size=7, bold=True, color=(0x1a, 0x3c, 0x6e))

    for q in ticket.get('questions', []):
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(0)
        add_run(p, f"{q['num']}. {q.get('title','')}", size=5.5, bold=True)

        for line in condense(q.get('answer', '')):
            lp = cell.add_paragraph()
            lp.paragraph_format.space_before = Pt(0)
            lp.paragraph_format.space_after = Pt(0)
            lp.paragraph_format.left_indent = Cm(0.15)
            add_run(lp, f"• {line}", size=5)

def build():
    tickets = json.load(open(f"{REPO}/tickets.json", encoding="utf-8"))
    tickets.sort(key=lambda t: (int(re.match(r'\d+', str(t.get('num','0'))).group() or 0), str(t.get('num'))))

    doc = Document()
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(0.5)
    section.bottom_margin = Cm(0.5)
    section.left_margin = Cm(0.6)
    section.right_margin = Cm(0.6)

    PER_PAGE = 8  # 2 columns x 4 rows
    COLS = 2

    for i in range(0, len(tickets), PER_PAGE):
        chunk = tickets[i:i+PER_PAGE]
        rows = (len(chunk) + COLS - 1) // COLS
        table = doc.add_table(rows=rows, cols=COLS)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        for col in table.columns:
            for c in col.cells:
                c.width = Cm(9.8)
        for idx, ticket in enumerate(chunk):
            r, c = divmod(idx, COLS)
            fill_ticket_cell(table.cell(r, c), ticket)
        if i + PER_PAGE < len(tickets):
            doc.add_page_break()

    out_path = "/tmp/claude-0/-home-user-vmeda-biology-bot/d28eaeb6-ad1f-5436-88a0-95ec7597ae8d/scratchpad/VMedA_shpora_biology.docx"
    doc.save(out_path)
    print("saved:", out_path, "tickets:", len(tickets), "pages:", (len(tickets)+PER_PAGE-1)//PER_PAGE)

build()
